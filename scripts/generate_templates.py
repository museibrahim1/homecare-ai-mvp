#!/usr/bin/env python3
"""
Generate DOCX templates for visit notes and service contracts.

Run this script to create the initial template files:
    python scripts/generate_templates.py
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


def create_note_template():
    """Create the visit note template."""
    doc = Document()
    
    # Set up styles
    style = doc.styles['Heading 1']
    style.font.size = Pt(16)
    style.font.bold = True
    
    # Title
    title = doc.add_heading('HOME CARE VISIT NOTE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Visit Information Section
    doc.add_heading('Visit Information', level=1)
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    cells = [
        ('Date:', '{{ visit_date }}'),
        ('Client:', '{{ client_name }}'),
        ('Caregiver:', '{{ caregiver_name }}'),
        ('Duration:', '{{ visit_duration }} minutes'),
    ]
    
    for i, (label, value) in enumerate(cells):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # Tasks Performed Section
    doc.add_heading('Tasks Performed', level=1)
    doc.add_paragraph('{% for task in tasks %}')
    doc.add_paragraph('• {{ task.category }}: {{ task.description }} ({{ task.duration_minutes }} min)')
    doc.add_paragraph('{% endfor %}')
    
    doc.add_paragraph()
    
    # Observations Section
    doc.add_heading('Observations', level=1)
    doc.add_paragraph('{{ observations }}')
    
    doc.add_paragraph()
    
    # Client Condition Section
    doc.add_heading('Client Condition', level=1)
    doc.add_paragraph('Status: {{ client_condition }}')
    
    doc.add_paragraph()
    
    # Concerns Section
    doc.add_heading('Risks / Concerns', level=1)
    doc.add_paragraph('{{ concerns }}')
    
    doc.add_paragraph()
    
    # Narrative Section
    doc.add_heading('Narrative Note', level=1)
    doc.add_paragraph('{{ narrative }}')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signatures Section
    doc.add_heading('Signatures', level=1)
    
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.rows[0].cells[0].text = 'Caregiver Signature:'
    sig_table.rows[0].cells[1].text = '_______________________'
    sig_table.rows[1].cells[0].text = 'Date:'
    sig_table.rows[1].cells[1].text = '_______________________'
    
    return doc


def create_contract_template():
    """Create the service contract template."""
    doc = Document()
    
    # Title
    title = doc.add_heading('HOME CARE SERVICE AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Client Information Section
    doc.add_heading('Client Information', level=1)
    
    client_table = doc.add_table(rows=5, cols=2)
    client_table.style = 'Table Grid'
    
    client_cells = [
        ('Name:', '{{ client_name }}'),
        ('Address:', '{{ client_address }}'),
        ('Phone:', '{{ client_phone }}'),
        ('Emergency Contact:', '{{ emergency_contact }}'),
        ('Emergency Phone:', '{{ emergency_phone }}'),
    ]
    
    for i, (label, value) in enumerate(client_cells):
        client_table.rows[i].cells[0].text = label
        client_table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # Services Section
    doc.add_heading('Services Provided', level=1)
    doc.add_paragraph('The following services will be provided under this agreement:')
    doc.add_paragraph()
    doc.add_paragraph('{% for service in services %}')
    doc.add_paragraph('• {{ service.name }}: {{ service.description }}')
    doc.add_paragraph('{% endfor %}')
    
    doc.add_paragraph()
    
    # Schedule Section
    doc.add_heading('Schedule', level=1)
    
    schedule_table = doc.add_table(rows=3, cols=2)
    schedule_table.style = 'Table Grid'
    
    schedule_cells = [
        ('Days of Service:', '{{ schedule_days }}'),
        ('Hours:', '{{ schedule_hours }}'),
        ('Hours per Week:', '{{ weekly_hours }}'),
    ]
    
    for i, (label, value) in enumerate(schedule_cells):
        schedule_table.rows[i].cells[0].text = label
        schedule_table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # Rates Section
    doc.add_heading('Rates and Fees', level=1)
    
    rates_table = doc.add_table(rows=4, cols=2)
    rates_table.style = 'Table Grid'
    
    rates_cells = [
        ('Hourly Rate:', '${{ hourly_rate }}'),
        ('Weekly Hours:', '{{ weekly_hours }}'),
        ('Weekly Cost:', '${{ weekly_cost }}'),
        ('Monthly Cost (est.):', '${{ monthly_cost }}'),
    ]
    
    for i, (label, value) in enumerate(rates_cells):
        rates_table.rows[i].cells[0].text = label
        rates_table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # Terms Section
    doc.add_heading('Terms', level=1)
    
    terms_table = doc.add_table(rows=2, cols=2)
    terms_table.style = 'Table Grid'
    
    terms_cells = [
        ('Start Date:', '{{ start_date }}'),
        ('End Date:', '{{ end_date }}'),
    ]
    
    for i, (label, value) in enumerate(terms_cells):
        terms_table.rows[i].cells[0].text = label
        terms_table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # Cancellation Policy
    doc.add_heading('Cancellation Policy', level=1)
    doc.add_paragraph('{{ cancellation_policy }}')
    
    doc.add_paragraph()
    
    # Liability
    doc.add_heading('Liability', level=1)
    doc.add_paragraph(
        'The Agency maintains comprehensive liability insurance. Caregivers are '
        'employees of the Agency and are covered under workers\' compensation. '
        'The Agency is not responsible for valuables left unsecured in the home.'
    )
    
    doc.add_paragraph()
    
    # Confidentiality
    doc.add_heading('Confidentiality', level=1)
    doc.add_paragraph(
        'All client information is kept strictly confidential in accordance with '
        'HIPAA regulations. Information will only be shared with healthcare providers '
        'directly involved in the client\'s care or as required by law.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signatures Section
    doc.add_heading('Signatures', level=1)
    
    doc.add_paragraph(
        'By signing below, both parties agree to the terms and conditions outlined '
        'in this agreement.'
    )
    
    doc.add_paragraph()
    
    sig_table = doc.add_table(rows=4, cols=2)
    
    sig_table.rows[0].cells[0].text = 'Client Signature:'
    sig_table.rows[0].cells[1].text = '{{ client_signature_line }}'
    sig_table.rows[1].cells[0].text = 'Date:'
    sig_table.rows[1].cells[1].text = '_______________________'
    sig_table.rows[2].cells[0].text = 'Agency Representative:'
    sig_table.rows[2].cells[1].text = '{{ agency_signature_line }}'
    sig_table.rows[3].cells[0].text = 'Date:'
    sig_table.rows[3].cells[1].text = '_______________________'
    
    return doc


def main():
    """Generate all templates."""
    # Ensure template directories exist
    notes_dir = Path(__file__).parent.parent / 'templates' / 'notes'
    contracts_dir = Path(__file__).parent.parent / 'templates' / 'contracts'
    
    notes_dir.mkdir(parents=True, exist_ok=True)
    contracts_dir.mkdir(parents=True, exist_ok=True)
    
    # Generate visit note template
    note_doc = create_note_template()
    note_path = notes_dir / 'visit_note_template.docx'
    note_doc.save(str(note_path))
    print(f"✓ Created: {note_path}")
    
    # Generate contract template
    contract_doc = create_contract_template()
    contract_path = contracts_dir / 'service_contract_template.docx'
    contract_doc.save(str(contract_path))
    print(f"✓ Created: {contract_path}")
    
    print("\nTemplates generated successfully!")


if __name__ == '__main__':
    main()
