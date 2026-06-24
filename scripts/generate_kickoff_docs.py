#!/usr/bin/env python3
"""Generate the 6 AI-build kickoff documents for PalmCare AI as .docx files.

Docs (per the "6 docs before you build" framework):
  1. PRD  - Product Requirements Document
  2. TRD  - Technical Requirements Document
  3. App Flow
  4. UI/UX Design Brief
  5. Backend Schema
  6. Implementation Plan

Run: python3 scripts/generate_kickoff_docs.py
Output: docs/kickoff/*.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs", "kickoff")
os.makedirs(OUT_DIR, exist_ok=True)

TEAL = RGBColor(0x0D, 0x94, 0x88)   # palmPrimary
DARK = RGBColor(0x1F, 0x29, 0x37)
GRAY = RGBColor(0x6B, 0x72, 0x80)


def new_doc(title, subtitle):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = t.add_run("PALM \u2014 Home Care Contracts (PalmCare AI)")
    run.font.size = Pt(10)
    run.font.color.rgb = TEAL
    run.bold = True

    h = doc.add_paragraph()
    run = h.add_run(title)
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = DARK

    s = doc.add_paragraph()
    run = s.add_run(subtitle)
    run.font.size = Pt(12)
    run.italic = True
    run.font.color.rgb = GRAY

    meta = doc.add_paragraph()
    run = meta.add_run("Owner: Muse Ibrahim / Palm Technologies Inc.    \u2022    Status: Living document    \u2022    Last updated: 2026-06")
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    doc.add_paragraph()
    return doc


def h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = TEAL
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = DARK
    return p


def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def numbered(doc, text):
    doc.add_paragraph(text, style="List Number")


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, htext in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    doc.add_paragraph()
    return t


def save(doc, name):
    path = os.path.join(OUT_DIR, name)
    doc.save(path)
    print("wrote", path)


# ---------------------------------------------------------------- 1. PRD
def build_prd():
    d = new_doc("Product Requirements Document (PRD)", "What we're building and why")

    h1(d, "1. Overview")
    body(d, "PalmCare AI turns a recorded home-care client assessment into finished paperwork. A "
            "caregiver or agency admin records (or uploads) the assessment conversation; the app "
            "transcribes it, separates speakers, extracts billable services, drafts SOAP clinical "
            "notes, and generates a state-compliant home care service contract \u2014 reducing hours "
            "of manual documentation to minutes.")

    h1(d, "2. Problem")
    bullet(d, "Home-care agencies spend hours per client turning assessment visits into notes, billing, and contracts.")
    bullet(d, "Manual contracts risk non-compliance across 50 states + DC, each with different rules.")
    bullet(d, "Billable services are missed or mis-coded, causing revenue loss.")

    h1(d, "3. Target Users")
    table(d, ["Persona", "Need"], [
        ["Agency owner/admin", "Fast, compliant contracts + accurate billing across many clients"],
        ["Caregiver / nurse", "Hands-free capture of an assessment; auto SOAP notes"],
        ["Back-office biller", "Clean list of billable services to submit"],
    ])

    h1(d, "4. Core Features (v1)")
    table(d, ["Feature", "Description", "Priority"], [
        ["Client management", "7-section client profile (Personal, Contact, Emergency, Medical, Care Plan, Insurance, Scheduling)", "P0"],
        ["Visit recording", "Record live or upload audio of an assessment visit", "P0"],
        ["Transcription", "Speaker-labeled transcript (Deepgram Nova-3)", "P0"],
        ["Billables extraction", "AI-detected billable services with accept/deny", "P0"],
        ["SOAP notes", "Auto-generated clinical documentation", "P0"],
        ["Contract generation", "State-compliant contract, 3 selectable templates", "P0"],
        ["AI data consent", "Explicit consent at registration + before processing", "P0"],
        ["Subscriptions", "Paid plans gating usage", "P1"],
    ])

    h1(d, "5. User Stories")
    bullet(d, "As an admin, I can add a client with all intake details so the contract is accurate.")
    bullet(d, "As a caregiver, I can record a visit and get a speaker-labeled transcript automatically.")
    bullet(d, "As a biller, I can review AI-extracted billables and accept/deny each one.")
    bullet(d, "As an admin, I can generate a state-compliant contract and preview 3 templates.")
    bullet(d, "As a user, I must explicitly consent to AI data processing before any audio is sent.")

    h1(d, "6. Success Metrics")
    bullet(d, "Time from recording to finished contract < 10 minutes.")
    bullet(d, "Billables accuracy (accepted without edit) > 80%.")
    bullet(d, "Trial-to-paid conversion (see Growth playbook: onboarding wow-moment, day-3 push).")

    h1(d, "7. Out of Scope (v1)")
    bullet(d, "Direct clearinghouse claim submission.")
    bullet(d, "Multi-language transcription beyond English.")
    bullet(d, "EHR integrations.")
    save(d, "1_PRD.docx")


# ---------------------------------------------------------------- 2. TRD
def build_trd():
    d = new_doc("Technical Requirements Document (TRD)", "The blueprint: stack, services, and decisions")

    h1(d, "1. Architecture Summary")
    body(d, "Native iOS + web clients talk to a FastAPI backend on Railway. The backend orchestrates "
            "an AI pipeline (Deepgram for transcription, Anthropic Claude for reasoning) and persists "
            "data in Supabase (Postgres + Auth + Storage). State-specific contract rules are encoded in "
            "the worker library.")

    h1(d, "2. Stack")
    table(d, ["Layer", "Choice", "Notes"], [
        ["iOS app", "SwiftUI (iOS 16+), XcodeGen", "Bundle com.palmcareai.app; Keychain token storage"],
        ["Web app", "Next.js / React", "apps/web; marketing + app"],
        ["Backend API", "FastAPI (Python)", "Railway: api-production-a0a2.up.railway.app"],
        ["Database/Auth", "Supabase (Postgres)", "RLS-enforced; Supabase Auth"],
        ["Storage", "Supabase Storage", "Audio uploads; lifecycle cleanup required"],
        ["Transcription", "Deepgram Nova-3", "Speaker diarization + smart format"],
        ["LLM", "Anthropic Claude Sonnet 4.6", "claude-sonnet-4-6"],
        ["Email", "Resend", "send.palmtai.com (verified)"],
        ["Billing", "App Store subscriptions (+ RevenueCat/Superwall planned)", "Pricing experiments per Growth playbook"],
    ])

    h1(d, "3. AI Pipeline")
    numbered(d, "Transcribe audio (Deepgram Nova-3).")
    numbered(d, "Speaker separation / diarization.")
    numbered(d, "Billable service extraction (Claude).")
    numbered(d, "SOAP note generation (Claude).")
    numbered(d, "State-compliant contract generation (Claude + state_rules.py).")

    h1(d, "4. Key APIs / Endpoints (representative)")
    bullet(d, "Auth: business registration (captures accepted_terms + timestamp), login.")
    bullet(d, "Clients: CRUD with full intake schema.")
    bullet(d, "Visits: create, upload audio, fetch transcript/billables/notes/contract.")
    bullet(d, "Processing: trigger pipeline; poll status.")

    h1(d, "5. Non-Functional Requirements")
    bullet(d, "Security: RLS on every table; server-side validation (Pydantic); rate limiting on auth + AI endpoints; tokens in Keychain (iOS) / httpOnly cookies (web).")
    bullet(d, "Privacy/Compliance: explicit AI-processing consent; PHI handled per HIPAA-minded practices; no model training on user data.")
    bullet(d, "Cost: scoped DB queries (no SELECT *), indexes on filter columns, connection pooling, storage cleanup, log monitoring.")
    bullet(d, "Reliability: pipeline retries capped with dead-letter + alerting.")

    h1(d, "6. Environments & Secrets")
    bullet(d, "Secrets in .env / Railway dashboard \u2014 never committed.")
    bullet(d, "Keys: DEEPGRAM_API_KEY, ANTHROPIC, SUPABASE, RESEND_API_KEY.")
    save(d, "2_TRD.docx")


# ---------------------------------------------------------------- 3. App Flow
def build_flow():
    d = new_doc("App Flow", "Navigation, screens, and the user journey")

    h1(d, "1. Primary Journey")
    numbered(d, "Onboarding \u2192 reach a value/'wow' moment before signup (sample processed assessment).")
    numbered(d, "Registration \u2192 enter business details \u2192 review & agree to Terms, Privacy, and AI Data Processing (mandatory consent) \u2192 account created.")
    numbered(d, "Home / Clients list \u2192 add a client (7-section form).")
    numbered(d, "Open client \u2192 start a Visit \u2192 record or upload audio (AI consent required before send).")
    numbered(d, "Processing \u2192 transcript \u2192 billables (accept/deny) \u2192 SOAP notes \u2192 contract (pick 1 of 3 templates).")
    numbered(d, "Review/export finished contract.")

    h1(d, "2. Screen Map")
    table(d, ["Area", "Screens"], [
        ["Auth", "Landing, Login, Register, Registration Consent"],
        ["Clients", "Clients list, Add/Edit Client (7 tabs), Client Detail"],
        ["Visits", "Visit list, Record, Upload, Visit Detail (Overview/Transcript/Billables/Notes/Contract)"],
        ["Settings", "Profile, Subscription, Terms & Privacy, AI Data Consent"],
    ])

    h1(d, "3. Client Form Tabs")
    bullet(d, "Personal \u2022 Contact \u2022 Emergency \u2022 Medical \u2022 Care Plan \u2022 Insurance \u2022 Scheduling")

    h1(d, "4. Visit Detail Tabs")
    bullet(d, "Overview (eager-loads all tab data) \u2022 Transcript (speaker-labeled) \u2022 Billables (accept/deny) \u2022 Notes (SOAP) \u2022 Contract (formatted, 3 templates)")

    h1(d, "5. Key State / Conversion Moments")
    bullet(d, "Consent gate: no audio leaves the device until AI consent is accepted.")
    bullet(d, "Onboarding wow-moment before paywall (Growth playbook 3.5).")
    bullet(d, "Cancellation save-offer; day-3 lifecycle push (Growth playbook 3.3/3.4).")
    save(d, "3_App_Flow.docx")


# ---------------------------------------------------------------- 4. UI/UX
def build_uiux():
    d = new_doc("UI/UX Design Brief", "Look, feel, and interaction standards")

    h1(d, "1. Brand")
    table(d, ["Token", "Value", "Use"], [
        ["palmPrimary", "Teal #0D9488", "Primary actions, active states, accents"],
        ["palmText", "Dark slate", "Primary text"],
        ["Backgrounds", "Adaptive (secondarySystemGroupedBackground)", "Cards/sections, light + dark mode"],
    ])
    body(d, "Adaptive light/dark theme. Never hardcode colors \u2014 always use theme tokens (Color.palmPrimary, Color.palmText).")

    h1(d, "2. Typography")
    bullet(d, "iOS: system font (SF Pro), Dynamic Type supported.")
    bullet(d, "Clear hierarchy: large bold titles, medium section headers, readable body.")

    h1(d, "3. Components")
    bullet(d, "PalmAlert: custom branded alert \u2014 use .palmErrorAlert() / .palmConfirmAlert(), never native .alert().")
    bullet(d, "Recording orb: mic.fill icon (not waveform).")
    bullet(d, "Horizontal scrollable tab bar with palmPrimary active state.")
    bullet(d, "Structured cards (e.g., emergency contact card with icons).")

    h1(d, "4. Interaction & Accessibility")
    bullet(d, "Every interactive element has an .accessibilityLabel().")
    bullet(d, "Target WCAG 2.2 AA contrast; respect Dynamic Type and reduce-motion.")
    bullet(d, "Loading states for all async pipeline steps.")

    h1(d, "5. Screen Look (rough)")
    bullet(d, "Clients: list of cards, prominent 'Add Client' CTA.")
    bullet(d, "Visit Detail: segmented tabs; contract rendered as clean formatted document, not raw markup.")
    bullet(d, "Consent screens: clear summary blocks + checkbox; primary CTA disabled until agreed.")

    h1(d, "6. App Store Presentation (ASO)")
    bullet(d, "Screenshots show how the app works (capture flow \u2192 contract), not abstract art.")
    bullet(d, "Title/subtitle packed with searched keywords (home care, contracts, AI notes, billing).")
    save(d, "4_UIUX_Brief.docx")


# ---------------------------------------------------------------- 5. Schema
def build_schema():
    d = new_doc("Backend Schema", "Data model, auth flow, tables & relationships")

    h1(d, "1. Auth Flow")
    bullet(d, "Supabase Auth for identity. Business registration captures accepted_terms + accepted_terms_at (audit log).")
    bullet(d, "iOS stores tokens in Keychain; web uses httpOnly cookies. Email verification + password rules + leaked-password check enforced.")

    h1(d, "2. Core Tables (representative)")
    table(d, ["Table", "Key Columns", "Relationships"], [
        ["businesses", "id, name, state, owner_user_id, accepted_terms, accepted_terms_at, created_at", "1 owner (auth user)"],
        ["users", "id, business_id, email, role", "belongs to business"],
        ["clients", "id, business_id, personal, contact, emergency, medical, care_plan, insurance, scheduling (JSON sections)", "belongs to business"],
        ["visits", "id, client_id, business_id, audio_url, status, created_at", "belongs to client"],
        ["transcripts", "id, visit_id, segments(JSON w/ speaker labels)", "1:1 visit"],
        ["billables", "id, visit_id, code, description, status(accepted/denied)", "many per visit"],
        ["notes", "id, visit_id, soap(JSON)", "1:1 visit"],
        ["contracts", "id, visit_id, state, template, body, created_at", "1:1 visit"],
        ["consents", "id, user_id, type(ai_processing/terms/privacy), accepted_at", "many per user"],
        ["audit_log", "id, business_id, action, metadata(JSON), created_at", "many per business"],
    ])

    h1(d, "3. Security (RLS) \u2014 mandatory")
    bullet(d, "RLS enabled on EVERY table; policies scope rows to the user's business_id.")
    bullet(d, "Admin/role checks enforced server-side, never on the client.")
    bullet(d, "Validate get_advisors (security + performance) before release.")

    h1(d, "4. Cost & Performance")
    bullet(d, "Indexes on business_id, client_id, visit_id, status (all filter/join columns).")
    bullet(d, "Select only needed columns (no SELECT *). Paginate lists.")
    bullet(d, "Storage: delete orphaned/failed audio uploads; lifecycle rules on the audio bucket.")

    h1(d, "5. State Rules")
    bullet(d, "Contract compliance encoded in apps/worker/libs/state_rules.py (50 states + DC).")
    bullet(d, "scripts/test_all_states.py validates all 51 jurisdictions.")
    save(d, "5_Backend_Schema.docx")


# ---------------------------------------------------------------- 6. Plan
def build_plan():
    d = new_doc("Implementation Plan", "Step-by-step build sequence")

    h1(d, "Build Order (security & data foundations first)")
    numbered(d, "Foundations: Supabase project, schema + RLS policies, Auth (email verify, password rules, leaked-password check).")
    numbered(d, "Backend skeleton: FastAPI on Railway, Pydantic validation, rate limiting, secrets via env.")
    numbered(d, "Client management: schema + CRUD endpoints + iOS/web 7-section form.")
    numbered(d, "Recording/upload: capture audio, store in Supabase Storage, AI consent gate before any send.")
    numbered(d, "AI pipeline: transcribe \u2192 speakers \u2192 billables \u2192 SOAP \u2192 contract (with state_rules).")
    numbered(d, "Review UX: Visit Detail tabs (transcript, billables accept/deny, notes, contract w/ 3 templates).")
    numbered(d, "Compliance: consent screens (registration + pre-processing), privacy policy, audit logging.")
    numbered(d, "Monetization: subscriptions + RevenueCat/Superwall; pricing experiments, cancellation save-offer, day-3 push.")
    numbered(d, "Growth/ASO: onboarding wow-moment, keyword-optimized listing, how-it-works screenshots.")
    numbered(d, "Hardening: get_advisors clean, indexes, storage cleanup, log/alerting on pipeline retries.")

    h1(d, "Per-Feature Definition of Done")
    bullet(d, "Security ship-gate passed (RLS, server-side authz, rate limits, validation, token storage).")
    bullet(d, "Cost ship-gate passed (scoped queries, indexes, no leaked connections, storage clean).")
    bullet(d, "Accessibility labels + loading/error states present.")
    bullet(d, "Committed to git (auto-commit rule).")

    h1(d, "Release Checklist (iOS)")
    bullet(d, "Bump CURRENT_PROJECT_VERSION, archive, upload; wait for VALID.")
    bullet(d, "Set 'What to Test', attach build to external group, invite testers.")
    bullet(d, "Submit for review with working demo account.")
    save(d, "6_Implementation_Plan.docx")


if __name__ == "__main__":
    build_prd()
    build_trd()
    build_flow()
    build_uiux()
    build_schema()
    build_plan()
    print("\nAll 6 kickoff docs written to:", OUT_DIR)
