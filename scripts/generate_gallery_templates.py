#!/usr/bin/env python3
"""
Generate professional DOCX gallery templates for the contract template system.
Places them in apps/api/templates/contracts/ with matching _meta.json files.
"""

import json
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


OUTPUT_DIR = Path(__file__).parent.parent / "apps" / "api" / "templates" / "contracts"


def _set_cell_shading(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _add_styled_table(doc, rows_data, header_color="4F46E5"):
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = "Table Grid"
    table.autofit = True

    for i, (label, value) in enumerate(rows_data):
        label_cell = table.rows[i].cells[0]
        value_cell = table.rows[i].cells[1]
        label_cell.text = label
        value_cell.text = value

        for p in label_cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        for p in value_cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

    return table


def create_standard_agreement():
    """PalmCare Standard Service Agreement — full-featured template."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("HOME CARE SERVICE AGREEMENT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Contract #: {{contract_id}}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    doc.add_heading("AGENCY INFORMATION", level=1)
    _add_styled_table(doc, [
        ("Agency Name:", "{{agency_name}}"),
        ("Address:", "{{agency_address}}"),
        ("City:", "{{agency_city}}"),
        ("State:", "{{agency_state}}"),
        ("ZIP Code:", "{{agency_zip}}"),
        ("Phone:", "{{agency_phone}}"),
        ("Email:", "{{agency_email}}"),
    ])

    doc.add_paragraph()
    doc.add_heading("CLIENT INFORMATION", level=1)
    _add_styled_table(doc, [
        ("Full Name:", "{{client_name}}"),
        ("Date of Birth:", "{{date_of_birth}}"),
        ("Address:", "{{client_address}}"),
        ("City:", "{{client_city}}"),
        ("State:", "{{client_state}}"),
        ("ZIP Code:", "{{client_zip}}"),
        ("Phone:", "{{client_phone}}"),
        ("Email:", "{{client_email}}"),
        ("Emergency Contact:", "{{emergency_contact}}"),
        ("Emergency Phone:", "{{emergency_phone}}"),
    ])

    doc.add_paragraph()
    doc.add_heading("CARE ASSESSMENT", level=1)
    _add_styled_table(doc, [
        ("Care Need Level:", "{{care_level}}"),
        ("Primary Diagnosis:", "{{primary_diagnosis}}"),
        ("Mobility Status:", "{{mobility_status}}"),
        ("Cognitive Status:", "{{cognitive_status}}"),
        ("Living Situation:", "{{living_situation}}"),
    ])

    doc.add_paragraph()
    doc.add_heading("SERVICES PROVIDED", level=1)
    doc.add_paragraph(
        "The following services will be provided under this agreement:"
    )
    doc.add_paragraph("{{services_list}}")

    doc.add_paragraph()
    doc.add_heading("SCHEDULE", level=1)
    _add_styled_table(doc, [
        ("Days of Service:", "{{schedule_days}}"),
        ("Start Time:", "{{start_time}}"),
        ("End Time:", "{{end_time}}"),
        ("Hours per Week:", "{{weekly_hours}}"),
        ("Frequency:", "{{frequency}}"),
    ])

    doc.add_paragraph()
    doc.add_heading("RATES AND FEES", level=1)
    _add_styled_table(doc, [
        ("Hourly Rate:", "{{hourly_rate}}"),
        ("Weekend Rate:", "{{weekend_rate}}"),
        ("Holiday Rate:", "{{holiday_rate}}"),
        ("Estimated Weekly Cost:", "{{weekly_cost}}"),
        ("Estimated Monthly Cost:", "{{monthly_cost}}"),
    ])

    doc.add_paragraph()
    doc.add_heading("CONTRACT TERMS", level=1)
    _add_styled_table(doc, [
        ("Effective Date:", "{{contract_date}}"),
        ("Start Date:", "{{start_date}}"),
        ("End Date:", "{{end_date}}"),
    ])

    doc.add_paragraph()
    doc.add_heading("SPECIAL REQUIREMENTS", level=1)
    doc.add_paragraph("{{special_requirements}}")

    doc.add_paragraph()
    doc.add_heading("SAFETY CONSIDERATIONS", level=1)
    doc.add_paragraph("{{safety_concerns}}")

    doc.add_paragraph()
    doc.add_heading("CANCELLATION POLICY", level=1)
    doc.add_paragraph("{{cancellation_policy}}")

    doc.add_paragraph()
    doc.add_heading("TERMS AND CONDITIONS", level=1)
    doc.add_paragraph("{{terms_and_conditions}}")

    doc.add_paragraph()
    doc.add_heading("LIABILITY", level=1)
    doc.add_paragraph(
        "The Agency maintains comprehensive liability insurance. Caregivers are "
        "employees of the Agency and are covered under workers' compensation. "
        "The Agency is not responsible for valuables left unsecured in the home."
    )

    doc.add_paragraph()
    doc.add_heading("CONFIDENTIALITY", level=1)
    doc.add_paragraph(
        "All client information is kept strictly confidential in accordance with "
        "HIPAA regulations. Information will only be shared with healthcare providers "
        "directly involved in the client's care or as required by law."
    )

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_heading("SIGNATURES", level=1)
    doc.add_paragraph(
        "By signing below, both parties agree to the terms and conditions "
        "outlined in this agreement."
    )
    doc.add_paragraph()

    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.style = "Table Grid"
    sig_table.rows[0].cells[0].text = "Client Signature:"
    sig_table.rows[0].cells[1].text = "______________________________"
    sig_table.rows[1].cells[0].text = "Date:"
    sig_table.rows[1].cells[1].text = "{{client_signature_date}}"
    sig_table.rows[2].cells[0].text = "Agency Representative:"
    sig_table.rows[2].cells[1].text = "______________________________"
    sig_table.rows[3].cells[0].text = "Date:"
    sig_table.rows[3].cells[1].text = "{{agency_signature_date}}"

    return doc


def create_basic_agreement():
    """Basic Home Care Agreement — simpler template for quick setup."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("HOME CARE SERVICES CONTRACT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    doc.add_heading("Provider", level=1)
    _add_styled_table(doc, [
        ("Agency:", "{{agency_name}}"),
        ("Phone:", "{{agency_phone}}"),
        ("Email:", "{{agency_email}}"),
    ])

    doc.add_paragraph()
    doc.add_heading("Client", level=1)
    _add_styled_table(doc, [
        ("Name:", "{{client_name}}"),
        ("Address:", "{{client_address}}"),
        ("Phone:", "{{client_phone}}"),
        ("Email:", "{{client_email}}"),
        ("Emergency Contact:", "{{emergency_contact}}"),
        ("Emergency Phone:", "{{emergency_phone}}"),
    ])

    doc.add_paragraph()
    doc.add_heading("Services & Schedule", level=1)
    doc.add_paragraph("{{services_list}}")
    doc.add_paragraph()
    _add_styled_table(doc, [
        ("Days:", "{{schedule_days}}"),
        ("Hours/Week:", "{{weekly_hours}}"),
        ("Hourly Rate:", "{{hourly_rate}}"),
        ("Monthly Estimate:", "{{monthly_cost}}"),
    ])

    doc.add_paragraph()
    doc.add_heading("Dates", level=1)
    _add_styled_table(doc, [
        ("Start Date:", "{{start_date}}"),
        ("End Date:", "{{end_date}}"),
    ])

    doc.add_paragraph()
    doc.add_heading("Agreement", level=1)
    doc.add_paragraph(
        "Both parties agree to the services, schedule, and rates described above. "
        "Either party may cancel this agreement with 30 days written notice. "
        "All client information is kept strictly confidential per HIPAA regulations."
    )
    doc.add_paragraph()

    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.style = "Table Grid"
    sig_table.rows[0].cells[0].text = "Client Signature:"
    sig_table.rows[0].cells[1].text = "______________________________"
    sig_table.rows[1].cells[0].text = "Date:"
    sig_table.rows[1].cells[1].text = "{{client_signature_date}}"
    sig_table.rows[2].cells[0].text = "Agency Representative:"
    sig_table.rows[2].cells[1].text = "______________________________"
    sig_table.rows[3].cells[0].text = "Date:"
    sig_table.rows[3].cells[1].text = "{{agency_signature_date}}"

    return doc


def create_skilled_nursing_agreement():
    """Skilled Nursing / Medical Home Care Agreement."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("SKILLED NURSING CARE AGREEMENT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Medical Home Care Services")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(79, 70, 229)

    doc.add_paragraph()

    doc.add_heading("CARE PROVIDER", level=1)
    _add_styled_table(doc, [
        ("Agency Name:", "{{agency_name}}"),
        ("Address:", "{{agency_address}}"),
        ("City/State/ZIP:", "{{agency_city}}, {{agency_state}} {{agency_zip}}"),
        ("Phone:", "{{agency_phone}}"),
        ("Email:", "{{agency_email}}"),
    ])

    doc.add_paragraph()
    doc.add_heading("PATIENT INFORMATION", level=1)
    _add_styled_table(doc, [
        ("Patient Name:", "{{client_name}}"),
        ("Date of Birth:", "{{date_of_birth}}"),
        ("Address:", "{{client_address}}"),
        ("City/State/ZIP:", "{{client_city}}, {{client_state}} {{client_zip}}"),
        ("Phone:", "{{client_phone}}"),
        ("Email:", "{{client_email}}"),
        ("Emergency Contact:", "{{emergency_contact}}"),
        ("Emergency Phone:", "{{emergency_phone}}"),
    ])

    doc.add_paragraph()
    doc.add_heading("MEDICAL ASSESSMENT", level=1)
    _add_styled_table(doc, [
        ("Care Level:", "{{care_level}}"),
        ("Primary Diagnosis:", "{{primary_diagnosis}}"),
        ("Mobility Status:", "{{mobility_status}}"),
        ("Cognitive Status:", "{{cognitive_status}}"),
        ("Living Situation:", "{{living_situation}}"),
    ])

    doc.add_paragraph()
    doc.add_heading("NURSING SERVICES", level=1)
    doc.add_paragraph(
        "The following skilled nursing and care services will be provided:"
    )
    doc.add_paragraph("{{services_list}}")

    doc.add_paragraph()
    doc.add_heading("CARE SCHEDULE", level=1)
    _add_styled_table(doc, [
        ("Days of Service:", "{{schedule_days}}"),
        ("Start Time:", "{{start_time}}"),
        ("End Time:", "{{end_time}}"),
        ("Hours per Week:", "{{weekly_hours}}"),
    ])

    doc.add_paragraph()
    doc.add_heading("RATES", level=1)
    _add_styled_table(doc, [
        ("Hourly Rate:", "{{hourly_rate}}"),
        ("Weekend Rate:", "{{weekend_rate}}"),
        ("Holiday Rate:", "{{holiday_rate}}"),
        ("Weekly Estimate:", "{{weekly_cost}}"),
        ("Monthly Estimate:", "{{monthly_cost}}"),
    ])

    doc.add_paragraph()
    doc.add_heading("SPECIAL REQUIREMENTS & SAFETY", level=1)
    doc.add_paragraph("Special Requirements:")
    doc.add_paragraph("{{special_requirements}}")
    doc.add_paragraph("Safety Considerations:")
    doc.add_paragraph("{{safety_concerns}}")

    doc.add_paragraph()
    doc.add_heading("CONTRACT PERIOD", level=1)
    _add_styled_table(doc, [
        ("Contract Date:", "{{contract_date}}"),
        ("Start Date:", "{{start_date}}"),
        ("End Date:", "{{end_date}}"),
    ])

    doc.add_paragraph()
    doc.add_heading("TERMS & CONDITIONS", level=1)
    doc.add_paragraph(
        "1. All care is provided under the supervision of a licensed nurse or "
        "physician as required by state regulations."
    )
    doc.add_paragraph(
        "2. The Agency maintains professional liability and malpractice insurance."
    )
    doc.add_paragraph(
        "3. All patient information is protected under HIPAA regulations."
    )
    doc.add_paragraph(
        "4. Either party may terminate with 14 days written notice."
    )
    doc.add_paragraph(
        "5. In case of a medical emergency, 911 will be called immediately."
    )

    doc.add_paragraph()
    doc.add_heading("SIGNATURES", level=1)
    doc.add_paragraph()

    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.style = "Table Grid"
    sig_table.rows[0].cells[0].text = "Patient/Guardian Signature:"
    sig_table.rows[0].cells[1].text = "______________________________"
    sig_table.rows[1].cells[0].text = "Date:"
    sig_table.rows[1].cells[1].text = "{{client_signature_date}}"
    sig_table.rows[2].cells[0].text = "Agency Representative:"
    sig_table.rows[2].cells[1].text = "______________________________"
    sig_table.rows[3].cells[0].text = "Date:"
    sig_table.rows[3].cells[1].text = "{{agency_signature_date}}"

    return doc


# ---- Meta JSON builders ----

STANDARD_META = {
    "name": "PalmCare Standard Service Agreement",
    "version": 1,
    "description": "Comprehensive home care service agreement with 45 fields covering agency info, client details, care assessment, services, schedule, rates, and legal terms.",
    "file_type": "docx",
    "detected_fields": [
        {"field_id": "agency_name", "label": "Agency Name", "type": "text", "required": True, "section": "agency_info", "mapped_to": "agency_name", "is_filled": False},
        {"field_id": "agency_address", "label": "Agency Address", "type": "text", "required": False, "section": "agency_info", "mapped_to": "agency_address", "is_filled": False},
        {"field_id": "agency_city", "label": "Agency City", "type": "text", "required": False, "section": "agency_info", "mapped_to": "agency_city", "is_filled": False},
        {"field_id": "agency_state", "label": "Agency State", "type": "text", "required": False, "section": "agency_info", "mapped_to": "agency_state", "is_filled": False},
        {"field_id": "agency_zip", "label": "Agency ZIP", "type": "text", "required": False, "section": "agency_info", "mapped_to": "agency_zip", "is_filled": False},
        {"field_id": "agency_phone", "label": "Agency Phone", "type": "phone", "required": False, "section": "agency_info", "mapped_to": "agency_phone", "is_filled": False},
        {"field_id": "agency_email", "label": "Agency Email", "type": "email", "required": False, "section": "agency_info", "mapped_to": "agency_email", "is_filled": False},
        {"field_id": "client_name", "label": "Client Full Name", "type": "text", "required": True, "section": "client_info", "mapped_to": "client_name", "is_filled": False},
        {"field_id": "date_of_birth", "label": "Date of Birth", "type": "date", "required": False, "section": "client_info", "mapped_to": "date_of_birth", "is_filled": False},
        {"field_id": "client_address", "label": "Client Address", "type": "text", "required": False, "section": "client_info", "mapped_to": "client_address", "is_filled": False},
        {"field_id": "client_city", "label": "Client City", "type": "text", "required": False, "section": "client_info", "mapped_to": "client_city", "is_filled": False},
        {"field_id": "client_state", "label": "Client State", "type": "text", "required": False, "section": "client_info", "mapped_to": "client_state", "is_filled": False},
        {"field_id": "client_zip", "label": "Client ZIP", "type": "text", "required": False, "section": "client_info", "mapped_to": "client_zip", "is_filled": False},
        {"field_id": "client_phone", "label": "Client Phone", "type": "phone", "required": False, "section": "client_info", "mapped_to": "client_phone", "is_filled": False},
        {"field_id": "client_email", "label": "Client Email", "type": "email", "required": False, "section": "client_info", "mapped_to": "client_email", "is_filled": False},
        {"field_id": "emergency_contact", "label": "Emergency Contact", "type": "text", "required": False, "section": "client_info", "mapped_to": "emergency_contact", "is_filled": False},
        {"field_id": "emergency_phone", "label": "Emergency Phone", "type": "phone", "required": False, "section": "client_info", "mapped_to": "emergency_phone", "is_filled": False},
        {"field_id": "care_level", "label": "Care Need Level", "type": "text", "required": True, "section": "assessment", "mapped_to": "care_level", "is_filled": False},
        {"field_id": "primary_diagnosis", "label": "Primary Diagnosis", "type": "text", "required": False, "section": "assessment", "mapped_to": "primary_diagnosis", "is_filled": False},
        {"field_id": "mobility_status", "label": "Mobility Status", "type": "text", "required": False, "section": "assessment", "mapped_to": "mobility_status", "is_filled": False},
        {"field_id": "cognitive_status", "label": "Cognitive Status", "type": "text", "required": False, "section": "assessment", "mapped_to": "cognitive_status", "is_filled": False},
        {"field_id": "living_situation", "label": "Living Situation", "type": "text", "required": False, "section": "assessment", "mapped_to": "living_situation", "is_filled": False},
        {"field_id": "services_list", "label": "Services List", "type": "list", "required": True, "section": "services", "mapped_to": "services_list", "is_filled": False},
        {"field_id": "schedule_days", "label": "Days of Service", "type": "text", "required": False, "section": "schedule", "mapped_to": "schedule_days", "is_filled": False},
        {"field_id": "start_time", "label": "Start Time", "type": "time", "required": False, "section": "schedule", "mapped_to": "start_time", "is_filled": False},
        {"field_id": "end_time", "label": "End Time", "type": "time", "required": False, "section": "schedule", "mapped_to": "end_time", "is_filled": False},
        {"field_id": "weekly_hours", "label": "Hours per Week", "type": "number", "required": False, "section": "schedule", "mapped_to": "weekly_hours", "is_filled": False},
        {"field_id": "frequency", "label": "Service Frequency", "type": "text", "required": False, "section": "schedule", "mapped_to": None, "is_filled": False},
        {"field_id": "hourly_rate", "label": "Hourly Rate", "type": "currency", "required": True, "section": "rates", "mapped_to": "hourly_rate", "is_filled": False},
        {"field_id": "weekend_rate", "label": "Weekend Rate", "type": "currency", "required": False, "section": "rates", "mapped_to": None, "is_filled": False},
        {"field_id": "holiday_rate", "label": "Holiday Rate", "type": "currency", "required": False, "section": "rates", "mapped_to": None, "is_filled": False},
        {"field_id": "weekly_cost", "label": "Weekly Cost", "type": "currency", "required": False, "section": "rates", "mapped_to": "weekly_cost", "is_filled": False},
        {"field_id": "monthly_cost", "label": "Monthly Cost", "type": "currency", "required": False, "section": "rates", "mapped_to": "monthly_cost", "is_filled": False},
        {"field_id": "contract_date", "label": "Contract Date", "type": "date", "required": True, "section": "contract", "mapped_to": "contract_date", "is_filled": False},
        {"field_id": "start_date", "label": "Start Date", "type": "date", "required": False, "section": "contract", "mapped_to": "start_date", "is_filled": False},
        {"field_id": "end_date", "label": "End Date", "type": "date", "required": False, "section": "contract", "mapped_to": "end_date", "is_filled": False},
        {"field_id": "contract_id", "label": "Contract ID", "type": "text", "required": False, "section": "contract", "mapped_to": None, "is_filled": False},
        {"field_id": "special_requirements", "label": "Special Requirements", "type": "list", "required": False, "section": "terms", "mapped_to": "special_requirements", "is_filled": False},
        {"field_id": "safety_concerns", "label": "Safety Considerations", "type": "list", "required": False, "section": "terms", "mapped_to": "safety_concerns", "is_filled": False},
        {"field_id": "cancellation_policy", "label": "Cancellation Policy", "type": "text", "required": False, "section": "terms", "mapped_to": "cancellation_policy", "is_filled": False},
        {"field_id": "terms_and_conditions", "label": "Terms and Conditions", "type": "text", "required": False, "section": "terms", "mapped_to": "terms_and_conditions", "is_filled": False},
        {"field_id": "client_signature", "label": "Client Signature", "type": "signature", "required": True, "section": "signatures", "mapped_to": None, "is_filled": False},
        {"field_id": "agency_signature", "label": "Agency Signature", "type": "signature", "required": True, "section": "signatures", "mapped_to": None, "is_filled": False},
        {"field_id": "client_signature_date", "label": "Client Signature Date", "type": "date", "required": False, "section": "signatures", "mapped_to": None, "is_filled": False},
        {"field_id": "agency_signature_date", "label": "Agency Signature Date", "type": "date", "required": False, "section": "signatures", "mapped_to": None, "is_filled": False},
    ],
    "field_mapping": {
        "agency_name": "agency.name", "agency_address": "agency.address",
        "agency_city": "agency.city", "agency_state": "agency.state",
        "agency_zip": "agency.zip_code", "agency_phone": "agency.phone",
        "agency_email": "agency.email", "client_name": "client.full_name",
        "date_of_birth": "client.date_of_birth", "client_address": "client.address",
        "client_city": "client.city", "client_state": "client.state",
        "client_zip": "client.zip_code", "client_phone": "client.phone",
        "client_email": "client.email", "emergency_contact": "client.emergency_contact_name",
        "emergency_phone": "client.emergency_contact_phone",
        "care_level": "contract.schedule.care_need_level",
        "primary_diagnosis": "client.primary_diagnosis",
        "mobility_status": "client.mobility_status",
        "cognitive_status": "client.cognitive_status",
        "living_situation": "client.living_situation",
        "services_list": "contract.services",
        "schedule_days": "contract.schedule.days",
        "start_time": "contract.schedule.start_time",
        "end_time": "contract.schedule.end_time",
        "weekly_hours": "contract.weekly_hours",
        "hourly_rate": "contract.hourly_rate",
        "weekly_cost": "computed.weekly_cost",
        "monthly_cost": "computed.monthly_cost",
        "contract_date": "computed.today",
        "start_date": "contract.start_date",
        "end_date": "contract.end_date",
        "special_requirements": "contract.schedule.special_requirements",
        "safety_concerns": "contract.schedule.safety_concerns",
        "cancellation_policy": "contract.cancellation_policy",
        "terms_and_conditions": "contract.terms_and_conditions",
    },
    "unmapped_fields": [
        {"field_id": "frequency", "label": "Service Frequency", "type": "text", "section": "schedule"},
        {"field_id": "weekend_rate", "label": "Weekend Rate", "type": "currency", "section": "rates"},
        {"field_id": "holiday_rate", "label": "Holiday Rate", "type": "currency", "section": "rates"},
        {"field_id": "contract_id", "label": "Contract ID", "type": "text", "section": "contract"},
        {"field_id": "client_signature", "label": "Client Signature", "type": "signature", "section": "signatures"},
        {"field_id": "agency_signature", "label": "Agency Signature", "type": "signature", "section": "signatures"},
        {"field_id": "client_signature_date", "label": "Client Signature Date", "type": "date", "section": "signatures"},
        {"field_id": "agency_signature_date", "label": "Agency Signature Date", "type": "date", "section": "signatures"},
    ],
}

BASIC_META = {
    "name": "Basic Home Care Contract",
    "version": 1,
    "description": "Simple one-page home care contract. Great for quick setup with essential fields only.",
    "file_type": "docx",
    "detected_fields": [
        {"field_id": "agency_name", "label": "Agency Name", "type": "text", "required": True, "section": "agency_info", "mapped_to": "agency_name", "is_filled": False},
        {"field_id": "agency_phone", "label": "Agency Phone", "type": "phone", "required": False, "section": "agency_info", "mapped_to": "agency_phone", "is_filled": False},
        {"field_id": "agency_email", "label": "Agency Email", "type": "email", "required": False, "section": "agency_info", "mapped_to": "agency_email", "is_filled": False},
        {"field_id": "client_name", "label": "Client Name", "type": "text", "required": True, "section": "client_info", "mapped_to": "client_name", "is_filled": False},
        {"field_id": "client_address", "label": "Client Address", "type": "text", "required": False, "section": "client_info", "mapped_to": "client_address", "is_filled": False},
        {"field_id": "client_phone", "label": "Client Phone", "type": "phone", "required": False, "section": "client_info", "mapped_to": "client_phone", "is_filled": False},
        {"field_id": "client_email", "label": "Client Email", "type": "email", "required": False, "section": "client_info", "mapped_to": "client_email", "is_filled": False},
        {"field_id": "emergency_contact", "label": "Emergency Contact", "type": "text", "required": False, "section": "client_info", "mapped_to": "emergency_contact", "is_filled": False},
        {"field_id": "emergency_phone", "label": "Emergency Phone", "type": "phone", "required": False, "section": "client_info", "mapped_to": "emergency_phone", "is_filled": False},
        {"field_id": "services_list", "label": "Services", "type": "list", "required": True, "section": "services", "mapped_to": "services_list", "is_filled": False},
        {"field_id": "schedule_days", "label": "Days", "type": "text", "required": False, "section": "schedule", "mapped_to": "schedule_days", "is_filled": False},
        {"field_id": "weekly_hours", "label": "Hours/Week", "type": "number", "required": False, "section": "schedule", "mapped_to": "weekly_hours", "is_filled": False},
        {"field_id": "hourly_rate", "label": "Hourly Rate", "type": "currency", "required": True, "section": "rates", "mapped_to": "hourly_rate", "is_filled": False},
        {"field_id": "monthly_cost", "label": "Monthly Estimate", "type": "currency", "required": False, "section": "rates", "mapped_to": "monthly_cost", "is_filled": False},
        {"field_id": "start_date", "label": "Start Date", "type": "date", "required": False, "section": "contract", "mapped_to": "start_date", "is_filled": False},
        {"field_id": "end_date", "label": "End Date", "type": "date", "required": False, "section": "contract", "mapped_to": "end_date", "is_filled": False},
        {"field_id": "client_signature", "label": "Client Signature", "type": "signature", "required": True, "section": "signatures", "mapped_to": None, "is_filled": False},
        {"field_id": "agency_signature", "label": "Agency Signature", "type": "signature", "required": True, "section": "signatures", "mapped_to": None, "is_filled": False},
        {"field_id": "client_signature_date", "label": "Client Signature Date", "type": "date", "required": False, "section": "signatures", "mapped_to": None, "is_filled": False},
        {"field_id": "agency_signature_date", "label": "Agency Signature Date", "type": "date", "required": False, "section": "signatures", "mapped_to": None, "is_filled": False},
    ],
    "field_mapping": {
        "agency_name": "agency.name", "agency_phone": "agency.phone",
        "agency_email": "agency.email", "client_name": "client.full_name",
        "client_address": "client.address", "client_phone": "client.phone",
        "client_email": "client.email", "emergency_contact": "client.emergency_contact_name",
        "emergency_phone": "client.emergency_contact_phone",
        "services_list": "contract.services", "schedule_days": "contract.schedule.days",
        "weekly_hours": "contract.weekly_hours", "hourly_rate": "contract.hourly_rate",
        "monthly_cost": "computed.monthly_cost", "start_date": "contract.start_date",
        "end_date": "contract.end_date",
    },
    "unmapped_fields": [
        {"field_id": "client_signature", "label": "Client Signature", "type": "signature", "section": "signatures"},
        {"field_id": "agency_signature", "label": "Agency Signature", "type": "signature", "section": "signatures"},
        {"field_id": "client_signature_date", "label": "Client Signature Date", "type": "date", "section": "signatures"},
        {"field_id": "agency_signature_date", "label": "Agency Signature Date", "type": "date", "section": "signatures"},
    ],
}

SKILLED_NURSING_META = {
    "name": "Skilled Nursing Care Agreement",
    "version": 1,
    "description": "Medical-grade home care agreement for skilled nursing services. Includes diagnosis, medical assessment, and clinical terms.",
    "file_type": "docx",
    "detected_fields": [
        {"field_id": "agency_name", "label": "Agency Name", "type": "text", "required": True, "section": "agency_info", "mapped_to": "agency_name", "is_filled": False},
        {"field_id": "agency_address", "label": "Agency Address", "type": "text", "required": False, "section": "agency_info", "mapped_to": "agency_address", "is_filled": False},
        {"field_id": "agency_city", "label": "Agency City", "type": "text", "required": False, "section": "agency_info", "mapped_to": "agency_city", "is_filled": False},
        {"field_id": "agency_state", "label": "Agency State", "type": "text", "required": False, "section": "agency_info", "mapped_to": "agency_state", "is_filled": False},
        {"field_id": "agency_zip", "label": "Agency ZIP", "type": "text", "required": False, "section": "agency_info", "mapped_to": "agency_zip", "is_filled": False},
        {"field_id": "agency_phone", "label": "Agency Phone", "type": "phone", "required": False, "section": "agency_info", "mapped_to": "agency_phone", "is_filled": False},
        {"field_id": "agency_email", "label": "Agency Email", "type": "email", "required": False, "section": "agency_info", "mapped_to": "agency_email", "is_filled": False},
        {"field_id": "client_name", "label": "Patient Name", "type": "text", "required": True, "section": "client_info", "mapped_to": "client_name", "is_filled": False},
        {"field_id": "date_of_birth", "label": "Date of Birth", "type": "date", "required": False, "section": "client_info", "mapped_to": "date_of_birth", "is_filled": False},
        {"field_id": "client_address", "label": "Patient Address", "type": "text", "required": False, "section": "client_info", "mapped_to": "client_address", "is_filled": False},
        {"field_id": "client_city", "label": "Patient City", "type": "text", "required": False, "section": "client_info", "mapped_to": "client_city", "is_filled": False},
        {"field_id": "client_state", "label": "Patient State", "type": "text", "required": False, "section": "client_info", "mapped_to": "client_state", "is_filled": False},
        {"field_id": "client_zip", "label": "Patient ZIP", "type": "text", "required": False, "section": "client_info", "mapped_to": "client_zip", "is_filled": False},
        {"field_id": "client_phone", "label": "Patient Phone", "type": "phone", "required": False, "section": "client_info", "mapped_to": "client_phone", "is_filled": False},
        {"field_id": "client_email", "label": "Patient Email", "type": "email", "required": False, "section": "client_info", "mapped_to": "client_email", "is_filled": False},
        {"field_id": "emergency_contact", "label": "Emergency Contact", "type": "text", "required": False, "section": "client_info", "mapped_to": "emergency_contact", "is_filled": False},
        {"field_id": "emergency_phone", "label": "Emergency Phone", "type": "phone", "required": False, "section": "client_info", "mapped_to": "emergency_phone", "is_filled": False},
        {"field_id": "care_level", "label": "Care Level", "type": "text", "required": True, "section": "assessment", "mapped_to": "care_level", "is_filled": False},
        {"field_id": "primary_diagnosis", "label": "Primary Diagnosis", "type": "text", "required": True, "section": "assessment", "mapped_to": "primary_diagnosis", "is_filled": False},
        {"field_id": "mobility_status", "label": "Mobility Status", "type": "text", "required": False, "section": "assessment", "mapped_to": "mobility_status", "is_filled": False},
        {"field_id": "cognitive_status", "label": "Cognitive Status", "type": "text", "required": False, "section": "assessment", "mapped_to": "cognitive_status", "is_filled": False},
        {"field_id": "living_situation", "label": "Living Situation", "type": "text", "required": False, "section": "assessment", "mapped_to": "living_situation", "is_filled": False},
        {"field_id": "services_list", "label": "Nursing Services", "type": "list", "required": True, "section": "services", "mapped_to": "services_list", "is_filled": False},
        {"field_id": "schedule_days", "label": "Days of Service", "type": "text", "required": False, "section": "schedule", "mapped_to": "schedule_days", "is_filled": False},
        {"field_id": "start_time", "label": "Start Time", "type": "time", "required": False, "section": "schedule", "mapped_to": "start_time", "is_filled": False},
        {"field_id": "end_time", "label": "End Time", "type": "time", "required": False, "section": "schedule", "mapped_to": "end_time", "is_filled": False},
        {"field_id": "weekly_hours", "label": "Hours per Week", "type": "number", "required": False, "section": "schedule", "mapped_to": "weekly_hours", "is_filled": False},
        {"field_id": "hourly_rate", "label": "Hourly Rate", "type": "currency", "required": True, "section": "rates", "mapped_to": "hourly_rate", "is_filled": False},
        {"field_id": "weekend_rate", "label": "Weekend Rate", "type": "currency", "required": False, "section": "rates", "mapped_to": None, "is_filled": False},
        {"field_id": "holiday_rate", "label": "Holiday Rate", "type": "currency", "required": False, "section": "rates", "mapped_to": None, "is_filled": False},
        {"field_id": "weekly_cost", "label": "Weekly Estimate", "type": "currency", "required": False, "section": "rates", "mapped_to": "weekly_cost", "is_filled": False},
        {"field_id": "monthly_cost", "label": "Monthly Estimate", "type": "currency", "required": False, "section": "rates", "mapped_to": "monthly_cost", "is_filled": False},
        {"field_id": "contract_date", "label": "Contract Date", "type": "date", "required": True, "section": "contract", "mapped_to": "contract_date", "is_filled": False},
        {"field_id": "start_date", "label": "Start Date", "type": "date", "required": False, "section": "contract", "mapped_to": "start_date", "is_filled": False},
        {"field_id": "end_date", "label": "End Date", "type": "date", "required": False, "section": "contract", "mapped_to": "end_date", "is_filled": False},
        {"field_id": "special_requirements", "label": "Special Requirements", "type": "list", "required": False, "section": "terms", "mapped_to": "special_requirements", "is_filled": False},
        {"field_id": "safety_concerns", "label": "Safety Considerations", "type": "list", "required": False, "section": "terms", "mapped_to": "safety_concerns", "is_filled": False},
        {"field_id": "client_signature", "label": "Patient/Guardian Signature", "type": "signature", "required": True, "section": "signatures", "mapped_to": None, "is_filled": False},
        {"field_id": "agency_signature", "label": "Agency Signature", "type": "signature", "required": True, "section": "signatures", "mapped_to": None, "is_filled": False},
        {"field_id": "client_signature_date", "label": "Patient Signature Date", "type": "date", "required": False, "section": "signatures", "mapped_to": None, "is_filled": False},
        {"field_id": "agency_signature_date", "label": "Agency Signature Date", "type": "date", "required": False, "section": "signatures", "mapped_to": None, "is_filled": False},
    ],
    "field_mapping": {
        "agency_name": "agency.name", "agency_address": "agency.address",
        "agency_city": "agency.city", "agency_state": "agency.state",
        "agency_zip": "agency.zip_code", "agency_phone": "agency.phone",
        "agency_email": "agency.email", "client_name": "client.full_name",
        "date_of_birth": "client.date_of_birth", "client_address": "client.address",
        "client_city": "client.city", "client_state": "client.state",
        "client_zip": "client.zip_code", "client_phone": "client.phone",
        "client_email": "client.email", "emergency_contact": "client.emergency_contact_name",
        "emergency_phone": "client.emergency_contact_phone",
        "care_level": "contract.schedule.care_need_level",
        "primary_diagnosis": "client.primary_diagnosis",
        "mobility_status": "client.mobility_status",
        "cognitive_status": "client.cognitive_status",
        "living_situation": "client.living_situation",
        "services_list": "contract.services",
        "schedule_days": "contract.schedule.days",
        "start_time": "contract.schedule.start_time",
        "end_time": "contract.schedule.end_time",
        "weekly_hours": "contract.weekly_hours",
        "hourly_rate": "contract.hourly_rate",
        "weekly_cost": "computed.weekly_cost",
        "monthly_cost": "computed.monthly_cost",
        "contract_date": "computed.today",
        "start_date": "contract.start_date",
        "end_date": "contract.end_date",
        "special_requirements": "contract.schedule.special_requirements",
        "safety_concerns": "contract.schedule.safety_concerns",
    },
    "unmapped_fields": [
        {"field_id": "weekend_rate", "label": "Weekend Rate", "type": "currency", "section": "rates"},
        {"field_id": "holiday_rate", "label": "Holiday Rate", "type": "currency", "section": "rates"},
        {"field_id": "client_signature", "label": "Patient/Guardian Signature", "type": "signature", "section": "signatures"},
        {"field_id": "agency_signature", "label": "Agency Signature", "type": "signature", "section": "signatures"},
        {"field_id": "client_signature_date", "label": "Patient Signature Date", "type": "date", "section": "signatures"},
        {"field_id": "agency_signature_date", "label": "Agency Signature Date", "type": "date", "section": "signatures"},
    ],
}


TEMPLATES = [
    ("palmcare_service_agreement", create_standard_agreement, STANDARD_META),
    ("basic_home_care_contract", create_basic_agreement, BASIC_META),
    ("skilled_nursing_agreement", create_skilled_nursing_agreement, SKILLED_NURSING_META),
]


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    for slug, builder, meta in TEMPLATES:
        docx_path = OUTPUT_DIR / f"{slug}.docx"
        meta_path = OUTPUT_DIR / f"{slug}_meta.json"

        doc = builder()
        doc.save(str(docx_path))
        print(f"  DOCX  {docx_path}")

        with open(meta_path, "w") as f:
            json.dump(meta, f, indent=2)
        print(f"  META  {meta_path}")

    print(f"\nGenerated {len(TEMPLATES)} gallery templates in {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
