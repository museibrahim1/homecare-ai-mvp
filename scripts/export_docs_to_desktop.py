#!/usr/bin/env python3
"""
Export project documents to the Desktop deliverables folder as .docx.

Converts the markdown plans and docs in docs/ and marketing/ into Word
documents and copies other final artifacts (decks, brochures, images)
into /Users/musaibrahim/Desktop/PalmCare Documents/.

Rerunnable: overwrites existing exports with the latest content.
"""

import re
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DEST = Path.home() / "Desktop" / "PalmCare Documents"

TEAL = RGBColor(0x0F, 0x76, 0x6E)
INK = RGBColor(0x1A, 0x1A, 0x1A)

DOCS = {
    "Company Docs": [
        "docs/INVESTOR_MEMO.md",
        "docs/API_COST_MODEL.md",
        "docs/ARCHITECTURE.md",
        "docs/DEPLOYMENT.md",
        "docs/data_flows.md",
        "docs/STRESS_TEST_REPORT.md",
        "docs/state-assessment-requirements.md",
        "docs/PALMCARE_APP_BRIEF.txt",
        "docs/PALMCARE_ROADMAP.txt",
    ],
    "Plans": [
        "docs/APP_STORE_LAUNCH_PLAN.md",
        "docs/APP_STORE_SUBMISSION.md",
        "docs/voice-agent-mobile-app-plan.md",
        "docs/POSTHOG_WAREHOUSE_RUNBOOK.md",
        "docs/marketing-research.md",
    ],
    "Marketing": [
        "marketing/PALM_BRAND_VOICE.md",
        "marketing/appstore-metadata.md",
        "marketing/asset-overview.md",
        "marketing/content-calendar-month-july.md",
        "marketing/content-calendar-week.md",
        "marketing/sage-style-campaign.md",
        "marketing/social-media-copy.md",
        "marketing/social-media-graphics.md",
    ],
}

COPIES = [
    ("marketing/pitch-deck-v5/PalmCare_Deck_v5.pdf", "PalmCare_Deck_v5.pdf"),
    ("marketing/pitch-deck-v5/PalmCare_Deck_v5.pptx", "PalmCare_Deck_v5.pptx"),
    ("apps/web/public/brochure/PalmCare-AI-Brochure.pdf", "PalmCare-AI-Brochure.pdf"),
]

INLINE_MD = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`|\[.+?\]\(.+?\))")


def add_inline(paragraph, text):
    """Render **bold**, *italic*, `code`, and [links](url) inside a paragraph."""
    pos = 0
    for m in INLINE_MD.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            paragraph.add_run(tok[2:-2]).bold = True
        elif tok.startswith("*"):
            paragraph.add_run(tok[1:-1]).italic = True
        elif tok.startswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Menlo"
            run.font.size = Pt(9.5)
        else:
            label = tok[1:tok.index("]")]
            url = tok[tok.index("(") + 1:-1]
            run = paragraph.add_run(f"{label} ({url})")
            run.font.color.rgb = TEAL
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def md_to_docx(src: Path, out: Path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Helvetica Neue"
    style.font.size = Pt(10.5)
    style.font.color.rgb = INK

    lines = src.read_text(encoding="utf-8", errors="replace").splitlines()
    i = 0
    in_code = False
    code_buf = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Menlo"
                run.font.size = Pt(8.5)
                code_buf = []
            in_code = not in_code
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # tables
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|?$", lines[i + 1].strip()):
            rows = [stripped]
            j = i + 2
            while j < len(lines) and lines[j].strip().startswith("|"):
                rows.append(lines[j].strip())
                j += 1
            cells = [
                [c.strip() for c in r.strip("|").split("|")]
                for r in rows
            ]
            ncols = max(len(r) for r in cells)
            table = doc.add_table(rows=len(cells), cols=ncols)
            table.style = "Light Grid Accent 3"
            for ri, row in enumerate(cells):
                for ci in range(ncols):
                    text = row[ci] if ci < len(row) else ""
                    cell_p = table.cell(ri, ci).paragraphs[0]
                    add_inline(cell_p, text)
                    if ri == 0:
                        for run in cell_p.runs:
                            run.bold = True
            i = j
            continue

        m = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if m:
            level = min(len(m.group(1)), 4)
            h = doc.add_heading("", level=level)
            add_inline(h, m.group(2))
            for run in h.runs:
                run.font.color.rgb = TEAL if level <= 2 else INK
        elif re.match(r"^[-*+]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, re.sub(r"^[-*+]\s+", "", stripped))
        elif re.match(r"^\d+[.)]\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+[.)]\s+", "", stripped))
        elif stripped.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            add_inline(p, stripped.lstrip("> "))
        elif re.match(r"^(-{3,}|\*{3,}|_{3,})$", stripped):
            pass  # horizontal rule, skip
        elif stripped:
            p = doc.add_paragraph()
            add_inline(p, stripped)
        i += 1

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)


def main():
    DEST.mkdir(parents=True, exist_ok=True)
    count = 0
    for folder, paths in DOCS.items():
        for rel in paths:
            src = PROJECT_ROOT / rel
            if not src.exists():
                print(f"SKIP (missing): {rel}")
                continue
            name = src.stem.replace("_", " ").replace("-", " ").title().replace(" ", "_")
            out = DEST / folder / f"{name}.docx"
            md_to_docx(src, out)
            print(f"docx: {folder}/{out.name}")
            count += 1

    for rel, name in COPIES:
        src = PROJECT_ROOT / rel
        if src.exists():
            shutil.copy2(src, DEST / name)
            print(f"copy: {name}")
            count += 1
        else:
            print(f"SKIP (missing): {rel}")

    print(f"\n{count} files exported to {DEST}")


if __name__ == "__main__":
    main()
