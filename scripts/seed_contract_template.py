#!/usr/bin/env python3
"""
Generate the master contract template DOCX with all known fields.
This template is used by the OCR system and seeded into the database.

Run:
    python scripts/seed_contract_template.py
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path
import json


def set_cell_shading(cell, color_hex: str):
    """Set background shading on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shd)


def add_field_row(table, label: str, placeholder: str, row_idx: int = None):
    """Add a label/value row to a table."""
    if row_idx is not None:
        row = table.rows[row_idx]
    else:
        row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = placeholder
    for cell in row.cells:
        for para in cell.paragraphs:
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
    row.cells[0].paragraphs[0].runs[0].bold = True if row.cells[0].paragraphs[0].runs else None
    return row


def create_master_contract_template() -> Document:
    """Create a comprehensive DOCX contract template with all form fields."""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    # ================================================================
    # HEADER / LETTERHEAD
    # ================================================================
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run('{agency_name}')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(30, 58, 138)

    addr_para = doc.add_paragraph()
    addr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    addr_para.add_run('{agency_address}').font.size = Pt(9)

    city_para = doc.add_paragraph()
    city_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = city_para.add_run('{agency_city}, {agency_state} {agency_zip}')
    run.font.size = Pt(9)

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact_para.add_run('Phone: {agency_phone}  |  Email: {agency_email}')
    run.font.size = Pt(9)

    doc.add_paragraph()

    # ================================================================
    # TITLE
    # ================================================================
    title = doc.add_heading('HOME CARE SERVICE AGREEMENT', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run('Effective Date: {contract_date}').font.size = Pt(10)

    doc.add_paragraph()

    # ================================================================
    # 1. PARTIES
    # ================================================================
    doc.add_heading('1. PARTIES', level=1)

    doc.add_paragraph().add_run('SERVICE PROVIDER:').bold = True
    parties_agency = doc.add_table(rows=4, cols=2)
    parties_agency.style = 'Table Grid'
    add_field_row(parties_agency, 'Agency Name:', '{agency_name}', 0)
    add_field_row(parties_agency, 'Address:', '{agency_address}', 1)
    add_field_row(parties_agency, 'Phone:', '{agency_phone}', 2)
    add_field_row(parties_agency, 'Email:', '{agency_email}', 3)

    doc.add_paragraph()
    doc.add_paragraph().add_run('CLIENT (Person to Receive Services):').bold = True

    parties_client = doc.add_table(rows=10, cols=2)
    parties_client.style = 'Table Grid'
    add_field_row(parties_client, 'Full Name:', '{client_name}', 0)
    add_field_row(parties_client, 'Date of Birth:', '{date_of_birth}', 1)
    add_field_row(parties_client, 'Address:', '{client_address}', 2)
    add_field_row(parties_client, 'City:', '{client_city}', 3)
    add_field_row(parties_client, 'State:', '{client_state}', 4)
    add_field_row(parties_client, 'ZIP Code:', '{client_zip}', 5)
    add_field_row(parties_client, 'Phone:', '{client_phone}', 6)
    add_field_row(parties_client, 'Email:', '{client_email}', 7)
    add_field_row(parties_client, 'Emergency Contact:', '{emergency_contact}', 8)
    add_field_row(parties_client, 'Emergency Phone:', '{emergency_phone}', 9)

    doc.add_paragraph()

    # ================================================================
    # 2. CARE ASSESSMENT
    # ================================================================
    doc.add_heading('2. CARE ASSESSMENT SUMMARY', level=1)

    assess_table = doc.add_table(rows=5, cols=2)
    assess_table.style = 'Table Grid'
    add_field_row(assess_table, 'Care Need Level:', '{care_level}', 0)
    add_field_row(assess_table, 'Primary Diagnosis:', '{primary_diagnosis}', 1)
    add_field_row(assess_table, 'Mobility Status:', '{mobility_status}', 2)
    add_field_row(assess_table, 'Cognitive Status:', '{cognitive_status}', 3)
    add_field_row(assess_table, 'Living Situation:', '{living_situation}', 4)

    doc.add_paragraph()

    # ================================================================
    # 3. SERVICES
    # ================================================================
    doc.add_heading('3. SERVICES TO BE PROVIDED', level=1)

    doc.add_paragraph('The following home care services will be provided under this Agreement:')
    doc.add_paragraph()
    doc.add_paragraph('{services_list}')

    doc.add_paragraph()

    # ================================================================
    # 4. SCHEDULE
    # ================================================================
    doc.add_heading('4. SCHEDULE OF SERVICES', level=1)

    sched_table = doc.add_table(rows=5, cols=2)
    sched_table.style = 'Table Grid'
    add_field_row(sched_table, 'Days of Service:', '{schedule_days}', 0)
    add_field_row(sched_table, 'Start Time:', '{start_time}', 1)
    add_field_row(sched_table, 'End Time:', '{end_time}', 2)
    add_field_row(sched_table, 'Hours per Week:', '{weekly_hours}', 3)
    add_field_row(sched_table, 'Frequency:', '{frequency}', 4)

    doc.add_paragraph()

    # ================================================================
    # 5. RATES
    # ================================================================
    doc.add_heading('5. SERVICE RATES AND PAYMENT', level=1)

    rates_table = doc.add_table(rows=6, cols=2)
    rates_table.style = 'Table Grid'
    add_field_row(rates_table, 'Hourly Rate:', '${hourly_rate}/hour', 0)
    add_field_row(rates_table, 'Weekend Rate:', '${weekend_rate}/hour', 1)
    add_field_row(rates_table, 'Holiday Rate:', '${holiday_rate}/hour', 2)
    add_field_row(rates_table, 'Estimated Weekly Cost:', '${weekly_cost}', 3)
    add_field_row(rates_table, 'Estimated Monthly Cost:', '${monthly_cost}', 4)
    add_field_row(rates_table, 'Payment Terms:', 'Due within 7 days of invoice', 5)

    doc.add_paragraph()

    # ================================================================
    # 6. CONTRACT DATES
    # ================================================================
    doc.add_heading('6. CONTRACT TERM', level=1)

    term_table = doc.add_table(rows=2, cols=2)
    term_table.style = 'Table Grid'
    add_field_row(term_table, 'Start Date:', '{start_date}', 0)
    add_field_row(term_table, 'End Date:', '{end_date}', 1)

    doc.add_paragraph()

    # ================================================================
    # 7. SPECIAL REQUIREMENTS
    # ================================================================
    doc.add_heading('7. SPECIAL REQUIREMENTS', level=1)
    doc.add_paragraph('{special_requirements}')

    doc.add_paragraph()

    # ================================================================
    # 8. SAFETY CONSIDERATIONS
    # ================================================================
    doc.add_heading('8. SAFETY CONSIDERATIONS', level=1)
    doc.add_paragraph('{safety_concerns}')

    doc.add_paragraph()

    # ================================================================
    # 9. CANCELLATION POLICY
    # ================================================================
    doc.add_heading('9. CANCELLATION POLICY', level=1)
    doc.add_paragraph('{cancellation_policy}')

    doc.add_paragraph()

    # ================================================================
    # 10. TERMS AND CONDITIONS
    # ================================================================
    doc.add_heading('10. TERMS AND CONDITIONS', level=1)
    doc.add_paragraph('{terms_and_conditions}')

    doc.add_paragraph()

    # ================================================================
    # 11. SIGNATURES
    # ================================================================
    doc.add_heading('11. SIGNATURES', level=1)

    doc.add_paragraph(
        'By signing below, both parties agree to the terms and conditions '
        'of this Agreement.'
    )
    doc.add_paragraph()

    sig_table = doc.add_table(rows=6, cols=2)

    sig_table.rows[0].cells[0].text = 'CLIENT / AUTHORIZED REPRESENTATIVE'
    sig_table.rows[0].cells[1].text = 'AGENCY REPRESENTATIVE'
    for cell in sig_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True

    sig_table.rows[1].cells[0].text = 'Signature: _________________________________'
    sig_table.rows[1].cells[1].text = 'Signature: _________________________________'

    sig_table.rows[2].cells[0].text = 'Printed Name: {client_name}'
    sig_table.rows[2].cells[1].text = 'Printed Name: _________________________'

    sig_table.rows[3].cells[0].text = 'Date: _______________'
    sig_table.rows[3].cells[1].text = 'Date: _______________'

    sig_table.rows[4].cells[0].text = ''
    sig_table.rows[4].cells[1].text = 'Title: _________________________'

    sig_table.rows[5].cells[0].text = ''
    sig_table.rows[5].cells[1].text = ''

    doc.add_paragraph()

    # ================================================================
    # FOOTER
    # ================================================================
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run('Contract ID: {contract_id}')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

    return doc


def get_all_detected_fields() -> list:
    """Return the full list of detected fields with their metadata."""
    return [
        # Agency fields
        {"field_id": "agency_name", "label": "Agency Name", "type": "text", "required": True, "section": "agency_info", "mapped_to": "agency_name", "is_filled": False},
        {"field_id": "agency_address", "label": "Agency Address", "type": "text", "required": False, "section": "agency_info", "mapped_to": "agency_address", "is_filled": False},
        {"field_id": "agency_city", "label": "Agency City", "type": "text", "required": False, "section": "agency_info", "mapped_to": "agency_city", "is_filled": False},
        {"field_id": "agency_state", "label": "Agency State", "type": "text", "required": False, "section": "agency_info", "mapped_to": "agency_state", "is_filled": False},
        {"field_id": "agency_zip", "label": "Agency ZIP", "type": "text", "required": False, "section": "agency_info", "mapped_to": "agency_zip", "is_filled": False},
        {"field_id": "agency_phone", "label": "Agency Phone", "type": "phone", "required": False, "section": "agency_info", "mapped_to": "agency_phone", "is_filled": False},
        {"field_id": "agency_email", "label": "Agency Email", "type": "email", "required": False, "section": "agency_info", "mapped_to": "agency_email", "is_filled": False},
        # Client fields
        {"field_id": "client_name", "label": "Client Full Name", "type": "text", "required": True, "section": "client_info", "mapped_to": "client_name", "is_filled": False},
        {"field_id": "date_of_birth", "label": "Date of Birth", "type": "date", "required": False, "section": "client_info", "mapped_to": "date_of_birth", "is_filled": False},
        {"field_id": "client_address", "label": "Client Address", "type": "text", "required": False, "section": "client_info", "mapped_to": "client_address", "is_filled": False},
        {"field_id": "client_city", "label": "Client City", "type": "text", "required": False, "section": "client_info", "mapped_to": "client_city", "is_filled": False},
        {"field_id": "client_state", "label": "Client State", "type": "text", "required": False, "section": "client_info", "mapped_to": "client_state", "is_filled": False},
        {"field_id": "client_zip", "label": "Client ZIP", "type": "text", "required": False, "section": "client_info", "mapped_to": "client_zip", "is_filled": False},
        {"field_id": "client_phone", "label": "Client Phone", "type": "phone", "required": False, "section": "client_info", "mapped_to": "client_phone", "is_filled": False},
        {"field_id": "client_email", "label": "Client Email", "type": "email", "required": False, "section": "client_info", "mapped_to": "client_email", "is_filled": False},
        {"field_id": "emergency_contact", "label": "Emergency Contact", "type": "text", "required": False, "section": "client_info", "mapped_to": "emergency_contact", "is_filled": False},
        {"field_id": "emergency_phone", "label": "Emergency Phone", "type": "phone", "required": False, "section": "client_info", "mapped_to": "emergency_phone", "is_filled": False},
        # Assessment fields
        {"field_id": "care_level", "label": "Care Need Level", "type": "text", "required": True, "section": "assessment", "mapped_to": "care_level", "is_filled": False},
        {"field_id": "primary_diagnosis", "label": "Primary Diagnosis", "type": "text", "required": False, "section": "assessment", "mapped_to": "primary_diagnosis", "is_filled": False},
        {"field_id": "mobility_status", "label": "Mobility Status", "type": "text", "required": False, "section": "assessment", "mapped_to": "mobility_status", "is_filled": False},
        {"field_id": "cognitive_status", "label": "Cognitive Status", "type": "text", "required": False, "section": "assessment", "mapped_to": "cognitive_status", "is_filled": False},
        {"field_id": "living_situation", "label": "Living Situation", "type": "text", "required": False, "section": "assessment", "mapped_to": "living_situation", "is_filled": False},
        # Services
        {"field_id": "services_list", "label": "Services List", "type": "list", "required": True, "section": "services", "mapped_to": "services_list", "is_filled": False},
        # Schedule
        {"field_id": "schedule_days", "label": "Days of Service", "type": "text", "required": False, "section": "schedule", "mapped_to": "schedule_days", "is_filled": False},
        {"field_id": "start_time", "label": "Start Time", "type": "time", "required": False, "section": "schedule", "mapped_to": "start_time", "is_filled": False},
        {"field_id": "end_time", "label": "End Time", "type": "time", "required": False, "section": "schedule", "mapped_to": "end_time", "is_filled": False},
        {"field_id": "weekly_hours", "label": "Hours per Week", "type": "number", "required": False, "section": "schedule", "mapped_to": "weekly_hours", "is_filled": False},
        {"field_id": "frequency", "label": "Service Frequency", "type": "text", "required": False, "section": "schedule", "mapped_to": None, "is_filled": False},
        # Rates
        {"field_id": "hourly_rate", "label": "Hourly Rate", "type": "currency", "required": True, "section": "rates", "mapped_to": "hourly_rate", "is_filled": False},
        {"field_id": "weekend_rate", "label": "Weekend Rate", "type": "currency", "required": False, "section": "rates", "mapped_to": None, "is_filled": False},
        {"field_id": "holiday_rate", "label": "Holiday Rate", "type": "currency", "required": False, "section": "rates", "mapped_to": None, "is_filled": False},
        {"field_id": "weekly_cost", "label": "Weekly Cost", "type": "currency", "required": False, "section": "rates", "mapped_to": "weekly_cost", "is_filled": False},
        {"field_id": "monthly_cost", "label": "Monthly Cost", "type": "currency", "required": False, "section": "rates", "mapped_to": "monthly_cost", "is_filled": False},
        # Contract dates
        {"field_id": "contract_date", "label": "Contract Date", "type": "date", "required": True, "section": "contract", "mapped_to": "contract_date", "is_filled": False},
        {"field_id": "start_date", "label": "Start Date", "type": "date", "required": False, "section": "contract", "mapped_to": "start_date", "is_filled": False},
        {"field_id": "end_date", "label": "End Date", "type": "date", "required": False, "section": "contract", "mapped_to": "end_date", "is_filled": False},
        {"field_id": "contract_id", "label": "Contract ID", "type": "text", "required": False, "section": "contract", "mapped_to": None, "is_filled": False},
        # Terms
        {"field_id": "special_requirements", "label": "Special Requirements", "type": "list", "required": False, "section": "terms", "mapped_to": "special_requirements", "is_filled": False},
        {"field_id": "safety_concerns", "label": "Safety Considerations", "type": "list", "required": False, "section": "terms", "mapped_to": "safety_concerns", "is_filled": False},
        {"field_id": "cancellation_policy", "label": "Cancellation Policy", "type": "text", "required": False, "section": "terms", "mapped_to": "cancellation_policy", "is_filled": False},
        {"field_id": "terms_and_conditions", "label": "Terms and Conditions", "type": "text", "required": False, "section": "terms", "mapped_to": "terms_and_conditions", "is_filled": False},
        # Signatures
        {"field_id": "client_signature", "label": "Client Signature", "type": "signature", "required": True, "section": "signatures", "mapped_to": None, "is_filled": False},
        {"field_id": "agency_signature", "label": "Agency Signature", "type": "signature", "required": True, "section": "signatures", "mapped_to": None, "is_filled": False},
        {"field_id": "client_signature_date", "label": "Client Signature Date", "type": "date", "required": False, "section": "signatures", "mapped_to": None, "is_filled": False},
        {"field_id": "agency_signature_date", "label": "Agency Signature Date", "type": "date", "required": False, "section": "signatures", "mapped_to": None, "is_filled": False},
    ]


def get_field_mapping() -> dict:
    """Return the pre-built field mapping (field_id → db_path)."""
    from_registry = {
        "agency_name": "agency.name",
        "agency_address": "agency.address",
        "agency_city": "agency.city",
        "agency_state": "agency.state",
        "agency_zip": "agency.zip_code",
        "agency_phone": "agency.phone",
        "agency_email": "agency.email",
        "client_name": "client.full_name",
        "date_of_birth": "client.date_of_birth",
        "client_address": "client.address",
        "client_city": "client.city",
        "client_state": "client.state",
        "client_zip": "client.zip_code",
        "client_phone": "client.phone",
        "client_email": "client.email",
        "emergency_contact": "client.emergency_contact_name",
        "emergency_phone": "client.emergency_contact_phone",
        "care_level": "contract.schedule.care_need_level",
        "primary_diagnosis": "client.primary_diagnosis",
        "mobility_status": "client.mobility_status",
        "cognitive_status": "client.cognitive_status",
        "living_situation": "client.living_situation",
        "services_list": "contract.services",
        "schedule_days": "contract.schedule.days",
        "start_time": "contract.schedule.start_time",
        "end_time": "contract.schedule.end_time",
        "weekly_hours": "contract.weekly_hours",
        "hourly_rate": "contract.hourly_rate",
        "weekly_cost": "computed.weekly_cost",
        "monthly_cost": "computed.monthly_cost",
        "contract_date": "computed.today",
        "start_date": "contract.start_date",
        "end_date": "contract.end_date",
        "special_requirements": "contract.schedule.special_requirements",
        "safety_concerns": "contract.schedule.safety_concerns",
        "cancellation_policy": "contract.cancellation_policy",
        "terms_and_conditions": "contract.terms_and_conditions",
    }
    return from_registry


def get_unmapped_fields() -> list:
    """Return fields not mapped to DB (signatures, computed rates, etc.)."""
    return [
        {"field_id": "frequency", "label": "Service Frequency", "type": "text", "section": "schedule"},
        {"field_id": "weekend_rate", "label": "Weekend Rate", "type": "currency", "section": "rates"},
        {"field_id": "holiday_rate", "label": "Holiday Rate", "type": "currency", "section": "rates"},
        {"field_id": "contract_id", "label": "Contract ID", "type": "text", "section": "contract"},
        {"field_id": "client_signature", "label": "Client Signature", "type": "signature", "section": "signatures"},
        {"field_id": "agency_signature", "label": "Agency Signature", "type": "signature", "section": "signatures"},
        {"field_id": "client_signature_date", "label": "Client Signature Date", "type": "date", "section": "signatures"},
        {"field_id": "agency_signature_date", "label": "Agency Signature Date", "type": "date", "section": "signatures"},
    ]


def main():
    """Generate the master contract template DOCX and metadata JSON."""
    contracts_dir = Path(__file__).parent.parent / 'templates' / 'contracts'
    contracts_dir.mkdir(parents=True, exist_ok=True)

    # Generate DOCX template
    doc = create_master_contract_template()
    template_path = contracts_dir / 'palmcare_service_agreement.docx'
    doc.save(str(template_path))
    print(f"  Created: {template_path}")

    # Generate metadata JSON for seeding
    metadata = {
        "name": "PalmCare Service Agreement",
        "version": 1,
        "description": "Standard home care service agreement template with all fields for OCR extraction",
        "file_type": "docx",
        "detected_fields": get_all_detected_fields(),
        "field_mapping": get_field_mapping(),
        "unmapped_fields": get_unmapped_fields(),
    }

    meta_path = contracts_dir / 'palmcare_service_agreement_meta.json'
    with open(meta_path, 'w') as f:
        json.dump(metadata, f, indent=2)
    print(f"  Created: {meta_path}")

    total_fields = len(metadata["detected_fields"])
    mapped = len(metadata["field_mapping"])
    unmapped = len(metadata["unmapped_fields"])
    print(f"\n  Template: {total_fields} fields detected, {mapped} mapped, {unmapped} unmapped")
    print("  Done!")


if __name__ == '__main__':
    main()
