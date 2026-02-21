"""
OCR Template Scanner Service

Uses Stirling-PDF for OCR text extraction from contract PDFs,
then uses AI (Claude/GPT) to identify form fields and build a JSON schema
mapping template fields to database columns.
"""

import os
import io
import json
import hashlib
import logging
from typing import Any

import httpx

logger = logging.getLogger(__name__)

STIRLING_PDF_URL = os.getenv("STIRLING_PDF_URL", "http://stirling-pdf:8080")

# All known contract database fields with their source paths
DB_FIELD_REGISTRY: dict[str, dict[str, str]] = {
    # Client fields
    "client_name": {"path": "client.full_name", "type": "text", "category": "client"},
    "name": {"path": "client.full_name", "type": "text", "category": "client"},
    "client_first_name": {"path": "client.full_name.split()[0]", "type": "text", "category": "client"},
    "client_last_name": {"path": "client.full_name.split()[-1]", "type": "text", "category": "client"},
    "client_address": {"path": "client.address", "type": "text", "category": "client"},
    "address": {"path": "client.address", "type": "text", "category": "client"},
    "client_city": {"path": "client.city", "type": "text", "category": "client"},
    "city": {"path": "client.city", "type": "text", "category": "client"},
    "client_state": {"path": "client.state", "type": "text", "category": "client"},
    "state": {"path": "client.state", "type": "text", "category": "client"},
    "client_zip": {"path": "client.zip_code", "type": "text", "category": "client"},
    "zip": {"path": "client.zip_code", "type": "text", "category": "client"},
    "zip_code": {"path": "client.zip_code", "type": "text", "category": "client"},
    "client_phone": {"path": "client.phone", "type": "phone", "category": "client"},
    "phone": {"path": "client.phone", "type": "phone", "category": "client"},
    "home_phone": {"path": "client.phone", "type": "phone", "category": "client"},
    "work_phone": {"path": "client.work_phone", "type": "phone", "category": "client"},
    "cell_phone": {"path": "client.phone", "type": "phone", "category": "client"},
    "client_email": {"path": "client.email", "type": "email", "category": "client"},
    "email": {"path": "client.email", "type": "email", "category": "client"},
    "date_of_birth": {"path": "client.date_of_birth", "type": "date", "category": "client"},
    "dob": {"path": "client.date_of_birth", "type": "date", "category": "client"},
    "emergency_contact": {"path": "client.emergency_contact_name", "type": "text", "category": "client"},
    "emergency_contact_name": {"path": "client.emergency_contact_name", "type": "text", "category": "client"},
    "emergency_phone": {"path": "client.emergency_contact_phone", "type": "phone", "category": "client"},
    "ssn": {"path": "client.ssn", "type": "text", "category": "client"},
    "social_security": {"path": "client.ssn", "type": "text", "category": "client"},
    "medicaid_number": {"path": "client.medicaid_number", "type": "text", "category": "client"},
    "medicare_number": {"path": "client.medicare_number", "type": "text", "category": "client"},
    "physician_name": {"path": "client.physician_name", "type": "text", "category": "client"},
    "physician_phone": {"path": "client.physician_phone", "type": "phone", "category": "client"},

    # Bill-To fields (alias to client — same person unless specified otherwise)
    "bill_to_name": {"path": "client.full_name", "type": "text", "category": "billing"},
    "bill_to_address": {"path": "client.address", "type": "text", "category": "billing"},
    "bill_to_city": {"path": "client.city", "type": "text", "category": "billing"},
    "bill_to_state": {"path": "client.state", "type": "text", "category": "billing"},
    "bill_to_zip": {"path": "client.zip_code", "type": "text", "category": "billing"},
    "bill_to_phone": {"path": "client.phone", "type": "phone", "category": "billing"},
    "bill_to_home_phone": {"path": "client.phone", "type": "phone", "category": "billing"},
    "bill_to_work_phone": {"path": "client.work_phone", "type": "phone", "category": "billing"},

    # Agency fields
    "agency_name": {"path": "agency.name", "type": "text", "category": "agency"},
    "agency_address": {"path": "agency.address", "type": "text", "category": "agency"},
    "agency_city": {"path": "agency.city", "type": "text", "category": "agency"},
    "agency_state": {"path": "agency.state", "type": "text", "category": "agency"},
    "agency_zip": {"path": "agency.zip_code", "type": "text", "category": "agency"},
    "agency_phone": {"path": "agency.phone", "type": "phone", "category": "agency"},
    "agency_email": {"path": "agency.email", "type": "email", "category": "agency"},

    # Rates & billing
    "hourly_rate": {"path": "contract.hourly_rate", "type": "currency", "category": "rates"},
    "weekday_rate": {"path": "contract.hourly_rate", "type": "currency", "category": "rates"},
    "weekend_rate": {"path": "computed.weekend_rate", "type": "currency", "category": "rates"},
    "holiday_rate": {"path": "computed.holiday_rate", "type": "currency", "category": "rates"},
    "client_rate": {"path": "contract.hourly_rate", "type": "currency", "category": "rates"},
    "weekly_cost": {"path": "computed.weekly_cost", "type": "currency", "category": "rates"},
    "monthly_cost": {"path": "computed.monthly_cost", "type": "currency", "category": "rates"},
    "monthly_package": {"path": "computed.monthly_cost", "type": "currency", "category": "rates"},
    "admin_fee": {"path": "computed.admin_fee", "type": "currency", "category": "rates"},
    "administrative_fee": {"path": "computed.admin_fee", "type": "currency", "category": "rates"},
    "deposit": {"path": "computed.deposit", "type": "currency", "category": "rates"},
    "prepayment": {"path": "computed.prepayment", "type": "currency", "category": "rates"},
    "total": {"path": "computed.monthly_cost", "type": "currency", "category": "rates"},
    "total_cost": {"path": "computed.monthly_cost", "type": "currency", "category": "rates"},

    # Schedule
    "weekly_hours": {"path": "contract.weekly_hours", "type": "number", "category": "schedule"},
    "hours_per_week": {"path": "contract.weekly_hours", "type": "number", "category": "schedule"},
    "schedule_days": {"path": "contract.schedule.days", "type": "text", "category": "schedule"},
    "days_of_service": {"path": "contract.schedule.days", "type": "text", "category": "schedule"},
    "start_time": {"path": "contract.schedule.start_time", "type": "time", "category": "schedule"},
    "end_time": {"path": "contract.schedule.end_time", "type": "time", "category": "schedule"},
    "frequency": {"path": "contract.schedule.frequency", "type": "text", "category": "schedule"},

    # Assessment
    "care_level": {"path": "contract.schedule.care_need_level", "type": "text", "category": "assessment"},
    "care_need_level": {"path": "contract.schedule.care_need_level", "type": "text", "category": "assessment"},
    "primary_diagnosis": {"path": "client.primary_diagnosis", "type": "text", "category": "assessment"},
    "diagnosis": {"path": "client.primary_diagnosis", "type": "text", "category": "assessment"},
    "mobility_status": {"path": "client.mobility_status", "type": "text", "category": "assessment"},
    "cognitive_status": {"path": "client.cognitive_status", "type": "text", "category": "assessment"},
    "living_situation": {"path": "client.living_situation", "type": "text", "category": "assessment"},

    # Services
    "services": {"path": "contract.services", "type": "list", "category": "services"},
    "services_list": {"path": "contract.services", "type": "list", "category": "services"},
    "services_provided": {"path": "contract.services", "type": "list", "category": "services"},

    # Contract dates
    "start_date": {"path": "contract.start_date", "type": "date", "category": "contract"},
    "effective_date": {"path": "contract.start_date", "type": "date", "category": "contract"},
    "end_date": {"path": "contract.end_date", "type": "date", "category": "contract"},
    "contract_date": {"path": "computed.today", "type": "date", "category": "contract"},
    "date": {"path": "computed.today", "type": "date", "category": "contract"},
    "agreement_date": {"path": "computed.today", "type": "date", "category": "contract"},
    "contract_id": {"path": "contract.id", "type": "text", "category": "contract"},

    # Terms / policies
    "cancellation_policy": {"path": "contract.cancellation_policy", "type": "text", "category": "terms"},
    "termination_policy": {"path": "contract.cancellation_policy", "type": "text", "category": "terms"},
    "terms_and_conditions": {"path": "contract.terms_and_conditions", "type": "text", "category": "terms"},
    "policies_and_procedures": {"path": "contract.policies_and_procedures", "type": "text", "category": "terms"},
    "special_requirements": {"path": "contract.schedule.special_requirements", "type": "list", "category": "assessment"},
    "safety_concerns": {"path": "contract.schedule.safety_concerns", "type": "list", "category": "assessment"},

    # Signatures
    "signature": {"path": "signature.client", "type": "signature", "category": "signatures"},
    "client_signature": {"path": "signature.client", "type": "signature", "category": "signatures"},
    "agency_signature": {"path": "signature.agency", "type": "signature", "category": "signatures"},
    "client_signature_date": {"path": "computed.today", "type": "date", "category": "signatures"},
    "agency_signature_date": {"path": "computed.today", "type": "date", "category": "signatures"},
    "signature_date": {"path": "computed.today", "type": "date", "category": "signatures"},
}


def compute_file_hash(file_bytes: bytes) -> str:
    return hashlib.sha256(file_bytes).hexdigest()


async def ocr_with_stirling(file_bytes: bytes, filename: str = "document.pdf") -> str:
    """
    Send a PDF to Stirling-PDF's OCR endpoint and get extracted text back.
    Falls back to local PyPDF extraction if Stirling is unavailable.
    """
    try:
        async with httpx.AsyncClient(timeout=120.0) as client:
            # Stirling-PDF v1 API: /api/v1/misc/ocr-pdf
            files = {"fileInput": (filename, io.BytesIO(file_bytes), "application/pdf")}
            data = {
                "languages": "eng",
                "sidecar": "true",
                "deskew": "true",
                "clean": "true",
                "cleanFinal": "true",
                "ocrType": "skip-text",
            }

            response = await client.post(
                f"{STIRLING_PDF_URL}/api/v1/misc/ocr-pdf",
                files=files,
                data=data,
            )

            if response.status_code == 200:
                # Stirling returns the OCR'd PDF; extract text from it
                ocr_pdf_bytes = response.content
                text = _extract_text_from_pdf_bytes(ocr_pdf_bytes)
                if text and text.strip():
                    logger.info(f"Stirling OCR extracted {len(text)} chars from {filename}")
                    return text

            logger.warning(f"Stirling OCR returned status {response.status_code}, falling back to local extraction")

    except httpx.ConnectError:
        logger.warning("Stirling-PDF not available, falling back to local extraction")
    except Exception as e:
        logger.warning(f"Stirling OCR failed: {e}, falling back to local extraction")

    # Fallback: local text extraction
    return _extract_text_from_pdf_bytes(file_bytes)


def _extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    """Extract text from PDF using PyPDF2 or pdfplumber."""
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n\n".join(text_parts)
    except ImportError:
        pass

    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            text_parts = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return "\n\n".join(text_parts)
    except ImportError:
        pass

    return "[PDF text extraction unavailable - install PyPDF2 or pdfplumber]"


def _extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    """Extract text from DOCX."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(docx_bytes))
        parts = []
        for section in doc.sections:
            try:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        parts.append(f"[HEADER] {para.text.strip()}")
            except Exception:
                pass
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return ""


async def extract_text(file_bytes: bytes, filename: str, content_type: str) -> str:
    """Extract text from a file, dispatching to the right extractor."""
    if "pdf" in content_type.lower() or filename.lower().endswith(".pdf"):
        return await ocr_with_stirling(file_bytes, filename)
    elif "word" in content_type.lower() or filename.lower().endswith((".docx", ".doc")):
        return _extract_text_from_docx_bytes(file_bytes)
    else:
        try:
            return file_bytes.decode("utf-8")
        except Exception:
            return ""


async def detect_fields_with_ai(ocr_text: str, filename: str) -> list[dict[str, Any]]:
    """
    Use Claude or GPT to analyze OCR'd text and identify all form fields,
    including their labels, types, and positions in the document.
    """
    anthropic_key = os.getenv("ANTHROPIC_API_KEY")
    openai_key = os.getenv("OPENAI_API_KEY")

    known_fields_json = json.dumps(list(DB_FIELD_REGISTRY.keys()), indent=2)

    system_prompt = f"""You are an expert at analyzing home care contract templates and identifying form fields.

Given the OCR-extracted text from a contract template, identify EVERY field/input that should be populated with data.

For each field found, return:
- "field_id": a snake_case identifier (use known IDs when matching)
- "label": the label text as it appears in the document
- "type": one of "text", "date", "phone", "email", "currency", "number", "time", "list", "signature", "checkbox"
- "required": true/false
- "section": which section of the contract this belongs to (e.g., "client_info", "agency_info", "services", "schedule", "rates", "terms", "signatures")
- "mapped_to": if this matches a known database field, set to the field_id from the known list. Otherwise null.
- "is_filled": true if the field already has a value in the template, false if blank/empty/underscores

KNOWN DATABASE FIELDS:
{known_fields_json}

IMPORTANT:
- Detect ALL blank lines, underscores (____), empty table cells, and form inputs
- Fields with "Label: _______" pattern are form fields
- Signature lines ("Signature: ____") count as signature fields
- Detect fields from BOTH old and new template versions
- If you see a field that doesn't match any known database field, still include it with mapped_to: null
- These unmapped fields are crucial — they represent NEW fields the database doesn't have yet

Return ONLY a JSON array of field objects. No other text."""

    user_prompt = f"""Template filename: {filename}

OCR-extracted text:
{ocr_text[:12000]}

Identify all form fields and return as a JSON array."""

    try:
        if anthropic_key:
            async with httpx.AsyncClient(timeout=60.0) as client:
                resp = await client.post(
                    "https://api.anthropic.com/v1/messages",
                    headers={
                        "x-api-key": anthropic_key,
                        "anthropic-version": "2023-06-01",
                        "content-type": "application/json",
                    },
                    json={
                        "model": "claude-3-haiku-20240307",
                        "max_tokens": 4096,
                        "temperature": 0,
                        "system": system_prompt,
                        "messages": [{"role": "user", "content": user_prompt}],
                    },
                )
                if resp.status_code == 200:
                    content = resp.json()["content"][0]["text"]
                    return _parse_json_array(content)

        if openai_key:
            async with httpx.AsyncClient(timeout=60.0) as client:
                resp = await client.post(
                    "https://api.openai.com/v1/chat/completions",
                    headers={
                        "Authorization": f"Bearer {openai_key}",
                        "Content-Type": "application/json",
                    },
                    json={
                        "model": "gpt-4o-mini",
                        "temperature": 0,
                        "messages": [
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_prompt},
                        ],
                    },
                )
                if resp.status_code == 200:
                    content = resp.json()["choices"][0]["message"]["content"]
                    return _parse_json_array(content)

    except Exception as e:
        logger.error(f"AI field detection failed: {e}")

    return []


def _parse_json_array(text: str) -> list[dict]:
    """Extract a JSON array from LLM response text."""
    text = text.strip()
    # Strip markdown code fences
    if text.startswith("```"):
        lines = text.split("\n")
        text = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])

    try:
        start = text.index("[")
        end = text.rindex("]") + 1
        return json.loads(text[start:end])
    except (ValueError, json.JSONDecodeError) as e:
        logger.error(f"Failed to parse AI response as JSON array: {e}")
        return []


def _fuzzy_match_registry(field_id: str) -> str | None:
    """
    Try to find a matching DB_FIELD_REGISTRY key for an unrecognized field_id.
    Handles common patterns like bill_to_name -> client_name, etc.
    """
    # Normalize: lowercase, strip whitespace
    fid = field_id.lower().strip()

    # Direct match
    if fid in DB_FIELD_REGISTRY:
        return fid

    # Strip "bill_to_" prefix and map to client equivalent
    if fid.startswith("bill_to_"):
        base = fid.replace("bill_to_", "")
        candidate = f"client_{base}" if not base.startswith("client_") else base
        if candidate in DB_FIELD_REGISTRY:
            return candidate
        if base in DB_FIELD_REGISTRY:
            return base

    # Try adding "client_" prefix
    if f"client_{fid}" in DB_FIELD_REGISTRY:
        return f"client_{fid}"

    # Try removing "client_" prefix
    if fid.startswith("client_") and fid[7:] in DB_FIELD_REGISTRY:
        return fid[7:]

    return None


def build_field_mapping(detected_fields: list[dict]) -> tuple[dict, list[dict]]:
    """
    Build a mapping from detected template fields to database columns.
    Returns (field_mapping, unmapped_fields).

    field_mapping: {field_id: db_path}
    unmapped_fields: [{field_id, label, type, section}]
    """
    field_mapping: dict[str, str] = {}
    unmapped: list[dict] = []

    for field in detected_fields:
        field_id = field.get("field_id", "")
        mapped_to = field.get("mapped_to")

        # Strategy 1: AI provided a mapped_to that exists in registry
        if mapped_to and mapped_to in DB_FIELD_REGISTRY:
            field_mapping[field_id] = DB_FIELD_REGISTRY[mapped_to]["path"]
        # Strategy 2: field_id itself exists in registry
        elif field_id in DB_FIELD_REGISTRY:
            field_mapping[field_id] = DB_FIELD_REGISTRY[field_id]["path"]
        else:
            # Strategy 3: fuzzy match
            matched_key = _fuzzy_match_registry(field_id)
            if not matched_key and mapped_to:
                matched_key = _fuzzy_match_registry(mapped_to)
            if matched_key:
                field_mapping[field_id] = DB_FIELD_REGISTRY[matched_key]["path"]
            else:
                unmapped.append({
                    "field_id": field_id,
                    "label": field.get("label", ""),
                    "type": field.get("type", "text"),
                    "section": field.get("section", "unknown"),
                })

    return field_mapping, unmapped


def reconcile_versions(
    old_fields: list[dict], new_fields: list[dict]
) -> dict[str, Any]:
    """
    Compare old template version fields with new template version fields.
    Returns a reconciliation report showing added, removed, and unchanged fields.
    """
    old_ids = {f["field_id"] for f in old_fields}
    new_ids = {f["field_id"] for f in new_fields}

    added = new_ids - old_ids
    removed = old_ids - new_ids
    unchanged = old_ids & new_ids

    new_by_id = {f["field_id"]: f for f in new_fields}
    old_by_id = {f["field_id"]: f for f in old_fields}

    return {
        "added_fields": [new_by_id[fid] for fid in added],
        "removed_fields": [old_by_id[fid] for fid in removed],
        "unchanged_fields": list(unchanged),
        "total_old": len(old_ids),
        "total_new": len(new_ids),
        "summary": (
            f"Template updated: {len(added)} new fields, "
            f"{len(removed)} removed, {len(unchanged)} unchanged"
        ),
    }
