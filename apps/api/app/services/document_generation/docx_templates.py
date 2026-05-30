"""User-uploaded DOCX template filling and HTML preview rendering."""

import io
import re
import base64
import logging
from typing import Any, List, Dict, Optional
from datetime import date

from .placeholders import get_template_placeholders

logger = logging.getLogger(__name__)


def fill_docx_template(template_bytes: bytes, placeholders: Dict[str, str], template_field_mapping: Optional[Dict[str, str]] = None) -> bytes:
    """
    Fill a DOCX template by replacing placeholders with actual values.
    
    Supports multiple placeholder formats:
    - {placeholder}, {{placeholder}}, [placeholder], [[placeholder]]
    - "Label:" format (e.g., "Name:" followed by blank space)
    - Case insensitive matching
    - AI-detected field_mapping from OCR template scan
    
    Args:
        template_bytes: The DOCX file bytes
        placeholders: Dict of placeholder key -> value
        template_field_mapping: Optional AI-detected field mapping from OCR scan
                                {field_id: db_path} e.g. {"client_name": "client.full_name"}
    """
    try:
        from docx import Document

        doc = Document(io.BytesIO(template_bytes))

        # If we have AI-detected field mappings, inject them as additional
        # placeholder aliases so the fill logic can resolve them
        if template_field_mapping:
            for field_id, db_path in template_field_mapping.items():
                # db_path is like "client.full_name" or "contract.hourly_rate"
                # Convert to placeholder key format
                path_key = db_path.replace(".", "_") if db_path else field_id
                # Try to find the value from existing placeholders
                value = placeholders.get(path_key, "")
                if not value:
                    # Try the field_id directly
                    value = placeholders.get(field_id, "")
                if not value:
                    # Try stripping prefix like "client." and matching
                    short_key = db_path.split(".")[-1] if "." in db_path else db_path
                    value = placeholders.get(short_key, "")
                if value and field_id not in placeholders:
                    placeholders[field_id] = value

        # Ordered list of (label_pattern, placeholder_key).
        # Longer/more-specific labels come first so they match before shorter ones.
        LABEL_MAPPINGS = [
            # Client / Patient fields (specific first)
            ('client name:', 'client_name'),
            ('patient name:', 'client_name'),
            ('client address:', 'client_address'),
            ('client city:', 'client_city'),
            ('client state:', 'client_state'),
            ('client zip:', 'client_zip'),
            ('client phone:', 'client_phone'),
            ('client email:', 'client_email'),
            ('emergency contact name:', 'emergency_contact'),
            ('emergency contact:', 'emergency_contact'),
            ('emergency phone:', 'emergency_phone'),
            ('date of birth:', 'date_of_birth'),
            ('dob:', 'date_of_birth'),
            ('home phone:', 'client_phone'),
            ('work phone:', 'work_phone'),
            ('cell phone:', 'client_phone'),
            ('social security:', 'ssn'),
            ('ssn:', 'ssn'),
            ('medicaid #:', 'medicaid_number'),
            ('medicaid number:', 'medicaid_number'),
            ('medicare #:', 'medicare_number'),
            ('medicare number:', 'medicare_number'),
            ('physician name:', 'physician_name'),
            ('physician phone:', 'physician_phone'),
            ('physician:', 'physician_name'),

            # Agency fields
            ('agency name:', 'agency_name'),
            ('agency address:', 'agency_address'),
            ('agency phone:', 'agency_phone'),
            ('agency email:', 'agency_email'),
            ('agency rep name:', 'agency_name'),
            ('agency rep signature:', 'agency_name'),

            # Rate & billing fields (specific first)
            ('hourly rate:', 'hourly_rate_value'),
            ('hourly:', 'hourly_rate_value'),
            ('weekday rate:', 'hourly_rate_value'),
            ('weekday:', 'hourly_rate_value'),
            ('weekend rate:', 'weekend_rate'),
            ('weekend:', 'weekend_rate'),
            ('holiday rate:', 'holiday_rate'),
            ('holiday:', 'holiday_rate'),
            ('client rate:', 'hourly_rate_value'),
            ('monthly package:', 'monthly_cost'),
            ('monthly cost:', 'monthly_cost'),
            ('monthly estimate:', 'monthly_cost'),
            ('weekly cost:', 'weekly_cost'),
            ('weekly estimate:', 'weekly_cost'),
            ('administrative fee:', 'admin_fee'),
            ('admin fee:', 'admin_fee'),
            ('deposit:', 'deposit'),
            ('prepayment:', 'prepayment'),
            ('total:', 'monthly_cost'),

            # Schedule fields
            ('weekly hours:', 'weekly_hours'),
            ('hours per week:', 'weekly_hours'),
            ('days of service:', 'schedule_days'),
            ('schedule:', 'schedule_days'),
            ('frequency:', 'frequency'),
            ('start time:', 'time'),
            ('end time:', 'time'),

            # Date fields
            ('effective date:', 'effective_date'),
            ('contract date:', 'contract_date'),
            ('start date:', 'effective_date'),
            ('end date:', 'effective_date'),
            ('agreement date:', 'contract_date'),

            # Services
            ('services to be provided:', 'services'),
            ('services provided:', 'services'),
            ('services:', 'services'),

            # Policy / Procedure / Terms fields
            ('cancellation policy:', 'cancellation_policy'),
            ('termination policy:', 'cancellation_policy'),
            ('policies and procedures:', 'policies_and_procedures'),
            ('policy and procedures:', 'policies_and_procedures'),
            ('terms and conditions:', 'terms_and_conditions'),
            ('terms of service:', 'terms_and_conditions'),
            ('special requirements:', 'special_requirements'),
            ('special instructions:', 'special_requirements'),
            ('safety considerations:', 'safety_concerns'),
            ('safety concerns:', 'safety_concerns'),
            ('living situation:', 'living_situation'),
            ('bill to:', 'client_name'),

            # Generic fallbacks (must be last)
            ('zip code:', 'client_zip'),
            ('zip:', 'client_zip'),
            ('name:', 'client_name'),
            ('address:', 'client_address'),
            ('city:', 'client_city'),
            ('state:', 'client_state'),
            ('phone:', 'client_phone'),
            ('email:', 'client_email'),
            ('date:', 'date'),
        ]

        def replace_placeholder_text(text: str, ph: Dict[str, str]) -> str:
            """Replace {{key}}, {key}, [[key]], [key], (key) patterns."""
            if not text:
                return text
            result = text
            for key, value in ph.items():
                for pat in [
                    rf'\{{\{{\s*{re.escape(key)}\s*\}}\}}',
                    rf'\{{\s*{re.escape(key)}\s*\}}',
                    rf'\[\[\s*{re.escape(key)}\s*\]\]',
                    rf'\[\s*{re.escape(key)}\s*\]',
                ]:
                    result = re.sub(pat, str(value or ''), result, flags=re.IGNORECASE)

            # Replace (Label) patterns like $ (Client Rate) -> $29.00
            # Also remove the leading "$ " to avoid "$ $29.00"
            result = re.sub(r'\$\s*\(\s*Client\s+Rate\s*\)', ph.get('hourly_rate', ''), result, flags=re.IGNORECASE)
            result = re.sub(r'\(\s*Client\s+Rate\s*\)', ph.get('hourly_rate_value', ''), result, flags=re.IGNORECASE)
            result = re.sub(r'\(\s*Client\s+Name\s*\)', ph.get('client_name', ''), result, flags=re.IGNORECASE)
            result = re.sub(r'\(\s*Agency\s+Name\s*\)', ph.get('agency_name', ''), result, flags=re.IGNORECASE)
            return result

        def fill_labels_in_text(text: str, ph: Dict[str, str]) -> str:
            """
            Replace "Label:" patterns with values.  Handles:
            - "Label: ___________"  (underscores)
            - "Label:  "           (just spaces)
            - "Label:"             (nothing after, end of line)
            - "City: State: Zip:"  (multiple labels on one line)
            """
            if not text:
                return text

            result = text
            text_lower = text.lower()
            # Track replaced ranges [start, end) to prevent overlapping fills
            replaced_ranges: list = []

            def overlaps(start: int, end: int) -> bool:
                for rs, re_ in replaced_ranges:
                    if start < re_ and end > rs:
                        return True
                return False

            for label, pk in LABEL_MAPPINGS:
                idx = text_lower.find(label)
                if idx == -1:
                    continue

                # Word boundary check: if char before match is a letter, skip
                # (prevents "services:" matching inside "Person to Receive Services:")
                if idx > 0 and result[idx - 1].isalpha():
                    continue

                # Check this position hasn't already been filled
                label_end = idx + len(label)
                if overlaps(idx, label_end):
                    continue

                value = ph.get(pk, '')
                if not value:
                    continue

                # Find the span of underscores/spaces/$ after the label
                span_end = label_end
                while span_end < len(result) and result[span_end] in '_ \t$':
                    span_end += 1

                span_content = result[label_end:span_end]

                should_fill = False
                if span_content and any(c == '_' for c in span_content):
                    should_fill = True
                elif span_end >= len(result):
                    should_fill = True
                elif span_end == label_end:
                    rest = result[label_end:].strip()
                    if not rest:
                        should_fill = True
                else:
                    if not span_content.strip() or span_content.strip() == '$':
                        should_fill = True

                if should_fill:
                    insertion = ' ' + str(value)
                    new_result = result[:label_end] + insertion + result[span_end:]
                    replaced_ranges.append((idx, label_end + len(insertion)))
                    shift = len(new_result) - len(result)
                    # Adjust existing ranges after this point
                    result = new_result
                    text_lower = result.lower()

            return result

        def set_paragraph_text(paragraph, new_text: str):
            """Overwrite a paragraph's text while keeping first run's formatting."""
            for run in paragraph.runs:
                run.text = ''
            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)

        def process_paragraph(paragraph, ph: Dict[str, str]):
            full_text = paragraph.text
            if not full_text:
                return
            new_text = replace_placeholder_text(full_text, ph)
            new_text = fill_labels_in_text(new_text, ph)
            if new_text != full_text:
                set_paragraph_text(paragraph, new_text)

        def process_cell_inline(cell, ph: Dict[str, str]):
            """
            Fill placeholders and label patterns within a cell.
            Returns the placeholder_key if this cell is a pure label cell
            (so the caller can fill the adjacent cell).
            """
            cell_text = cell.text.strip()
            if not cell_text:
                return None

            # Strip trailing asterisks/required markers for matching
            cell_clean = re.sub(r'[\s*]+$', '', cell_text)
            cell_lower = cell_clean.lower().strip()

            # Check if this cell is a pure label like "Name:", "Address", "Date of Birth:"
            for label, pk in LABEL_MAPPINGS:
                label_bare = label.rstrip(':').strip()
                if cell_lower == label or cell_lower == label_bare or cell_lower == label_bare + ':':
                    return pk

            # Also check without trailing colon for labels like "Home Phone" (no colon)
            if cell_lower.endswith(':'):
                cell_no_colon = cell_lower[:-1].strip()
            else:
                cell_no_colon = cell_lower
            for label, pk in LABEL_MAPPINGS:
                label_bare = label.rstrip(':').strip()
                if cell_no_colon == label_bare:
                    return pk

            # Fill {{placeholder}} patterns
            new_text = replace_placeholder_text(cell_text, ph)

            # Fill "Label: ___" patterns inside the cell
            new_text = fill_labels_in_text(new_text, ph)

            if new_text != cell_text:
                for paragraph in cell.paragraphs:
                    old = paragraph.text
                    if old.strip():
                        p_new = replace_placeholder_text(old, ph)
                        p_new = fill_labels_in_text(p_new, ph)
                        if p_new != old:
                            set_paragraph_text(paragraph, p_new)

            return None

        def process_table(table, ph: Dict[str, str]):
            """Process all rows/cells in a table, including nested tables."""
            for row in table.rows:
                cells = list(row.cells)
                pending_pk = None

                for i, cell in enumerate(cells):
                    if pending_pk:
                        value = ph.get(pending_pk, '')
                        if value:
                            cell_text = cell.text.strip()
                            # Fill if cell is blank, has only underscores/spaces, or has placeholder patterns
                            is_blank = (
                                not cell_text
                                or all(c in '_ \t\n*' for c in cell_text)
                                or cell_text in ('*', '**')
                            )
                            if is_blank:
                                for paragraph in cell.paragraphs:
                                    set_paragraph_text(paragraph, str(value))
                                    break
                            else:
                                # Cell has content — still try placeholder/label fill
                                process_cell_inline(cell, ph)
                        pending_pk = None
                    else:
                        pending_pk = process_cell_inline(cell, ph)

                    # Process nested tables inside this cell
                    for nested_table in cell.tables:
                        process_table(nested_table, ph)

        # --- Process document body paragraphs ---
        for paragraph in doc.paragraphs:
            process_paragraph(paragraph, placeholders)

        # --- Process tables (including nested) ---
        for table in doc.tables:
            process_table(table, placeholders)

        # --- Process headers / footers ---
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    process_paragraph(paragraph, placeholders)
                for table in section.header.tables:
                    process_table(table, placeholders)
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    process_paragraph(paragraph, placeholders)
                for table in section.footer.tables:
                    process_table(table, placeholders)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    except Exception as e:
        logger.error(f"Failed to fill DOCX template: {e}")
        raise


def docx_to_html(template_bytes: bytes, placeholders: Dict[str, str]) -> str:
    """
    Convert a DOCX template to styled HTML with placeholders filled in.
    Renders the full document for in-browser preview.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        filled_bytes = fill_docx_template(template_bytes, placeholders)
        doc = Document(io.BytesIO(filled_bytes))

        # Label -> placeholder key for second-pass smart fill in the HTML
        LABEL_MAP = {}
        for key in placeholders:
            nice = key.replace('_', ' ').title()
            LABEL_MAP[key.lower()] = key
            LABEL_MAP[nice.lower()] = key
            LABEL_MAP[nice.lower() + ':'] = key
        # Explicit common labels
        for lbl, pk in [
            ('name', 'client_name'), ('address', 'client_address'),
            ('city', 'client_city'), ('state', 'client_state'),
            ('zip', 'client_zip'), ('zip code', 'client_zip'),
            ('phone', 'client_phone'), ('home phone', 'client_phone'),
            ('work phone', 'work_phone'), ('cell phone', 'client_phone'),
            ('email', 'client_email'), ('date of birth', 'date_of_birth'),
            ('dob', 'date_of_birth'), ('date', 'date'),
            ('hourly rate', 'hourly_rate'), ('hourly', 'hourly_rate_value'),
            ('weekday', 'hourly_rate_value'), ('weekend', 'weekend_rate'),
            ('holiday', 'holiday_rate'), ('deposit', 'deposit'),
            ('prepayment', 'prepayment'), ('total', 'monthly_cost'),
            ('monthly package', 'monthly_cost'), ('administrative fee', 'admin_fee'),
            ('admin fee', 'admin_fee'), ('weekly hours', 'weekly_hours'),
            ('hours per week', 'weekly_hours'), ('effective date', 'effective_date'),
            ('start date', 'effective_date'), ('contract date', 'contract_date'),
            ('emergency contact', 'emergency_contact'), ('emergency phone', 'emergency_phone'),
            ('signature', ''), ('client signature', ''), ('agency representative', ''),
        ]:
            LABEL_MAP[lbl] = pk
            LABEL_MAP[lbl + ':'] = pk

        def smart_fill_cell(text: str) -> str:
            """Try to resolve a cell's content — if it's blank/underscores, keep as-is.
            If it contains a label with blanks, fill the blanks."""
            if not text:
                return text
            clean = text.strip()
            # Already has meaningful content (not just underscores)
            if clean and not all(c in '_ \t\n*$' for c in clean):
                return text
            return text

        def resolve_label(label_text: str) -> str:
            """Given a label like 'Name:' or 'Home Phone *', find the value."""
            clean = re.sub(r'[\s:*]+$', '', label_text).strip().lower()
            pk = LABEL_MAP.get(clean, '') or LABEL_MAP.get(clean + ':', '')
            if pk:
                return placeholders.get(pk, '')
            # Try partial
            for key, pkey in LABEL_MAP.items():
                if key and clean.endswith(key):
                    val = placeholders.get(pkey, '')
                    if val:
                        return val
            return ''

        def esc(text: str) -> str:
            return (text or '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

        html = []
        html.append(
            '<div style="font-family:Calibri,Arial,sans-serif; color:#1a1a1a; '
            'line-height:1.6; max-width:100%;">'
        )

        def runs_to_html(para) -> str:
            parts = []
            for run in para.runs:
                t = run.text
                if not t:
                    continue
                t = esc(t)
                if run.bold and run.italic:
                    t = f'<strong><em>{t}</em></strong>'
                elif run.bold:
                    t = f'<strong>{t}</strong>'
                elif run.italic:
                    t = f'<em>{t}</em>'
                if run.underline:
                    t = f'<u>{t}</u>'
                parts.append(t)
            return ''.join(parts) or esc(para.text)

        def get_align(para) -> str:
            try:
                if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    return 'text-align:center;'
                elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    return 'text-align:right;'
                elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                    return 'text-align:justify;'
            except Exception:
                pass
            return ''

        # Known multi-word labels for paragraph smart-fill (ordered long-first)
        PARA_LABELS = [
            ('person to receive services', None),  # skip — not a fillable field
            ('date of birth', 'date_of_birth'),
            ('home phone', 'client_phone'),
            ('work phone', 'work_phone'),
            ('cell phone', 'client_phone'),
            ('emergency contact', 'emergency_contact'),
            ('emergency phone', 'emergency_phone'),
            ('zip code', 'client_zip'),
            ('administrative fee', 'admin_fee'),
            ('admin fee', 'admin_fee'),
            ('monthly package', 'monthly_cost'),
            ('agency visits', None),
            ('hourly rate', 'hourly_rate'),
            ('weekly hours', 'weekly_hours'),
            ('bill to', None),  # section header
            ('signature', None),  # render as line
            ('name', 'client_name'),
            ('address', 'client_address'),
            ('city', 'client_city'),
            ('state', 'client_state'),
            ('zip', 'client_zip'),
            ('phone', 'client_phone'),
            ('email', 'client_email'),
            ('date', 'date'),
            ('hourly', 'hourly_rate'),
            ('weekday', 'hourly_rate_value'),
            ('weekend', 'weekend_rate'),
            ('holiday', 'holiday_rate'),
            ('prepayment', 'prepayment'),
            ('deposit', 'deposit'),
            ('total', 'monthly_cost'),
        ]

        def smart_fill_paragraph(text: str) -> str:
            """
            Parse paragraph with 'Label:' patterns using known label list.
            Returns HTML with bold labels and filled/editable values.
            """
            if not text or ':' not in text:
                return esc(text)

            # Find all label positions using the known label list
            text_lower = text.lower()
            found_labels = []  # [(start, end_of_label_colon, label_text, pk)]

            for lbl, pk in PARA_LABELS:
                search_from = 0
                while True:
                    idx = text_lower.find(lbl, search_from)
                    if idx == -1:
                        break
                    # Must be preceded by start-of-string or non-alpha
                    if idx > 0 and text_lower[idx - 1].isalpha():
                        search_from = idx + 1
                        continue
                    # Find the colon after the label
                    colon_pos = idx + len(lbl)
                    while colon_pos < len(text) and text[colon_pos] in ' \t':
                        colon_pos += 1
                    if colon_pos < len(text) and text[colon_pos] == ':':
                        found_labels.append((idx, colon_pos + 1, lbl, pk))
                    search_from = idx + 1

            if not found_labels:
                return esc(text)

            # Sort by position, remove overlapping
            found_labels.sort(key=lambda x: x[0])
            filtered = []
            last_end = -1
            for start, end, lbl, pk in found_labels:
                if start >= last_end:
                    filtered.append((start, end, lbl, pk))
                    last_end = end

            # Build HTML segments
            parts = []
            for i, (start, end, lbl, pk) in enumerate(filtered):
                # Get value text between this label's colon and next label's start
                val_start = end
                val_end = filtered[i + 1][0] if i + 1 < len(filtered) else len(text)
                raw_value = text[val_start:val_end].strip()

                # Clean value: strip leading $, underscores, asterisks
                clean_val = raw_value.strip(' \t_*$')

                # Try to fill if empty
                if not clean_val and pk:
                    clean_val = placeholders.get(pk, '')

                label_display = text[start:end - 1].strip()  # original case
                label_esc = esc(label_display)
                val_esc = esc(clean_val)

                if pk is None and lbl == 'signature':
                    parts.append(
                        f'<strong>{label_esc}:</strong> '
                        f'<span style="display:inline-block; width:200px; '
                        f'border-bottom:1px solid #374151; margin:0 8px;">&nbsp;</span>'
                    )
                elif pk is None:
                    # Non-fillable label (section header like "Bill To", "Person to Receive Services")
                    if clean_val:
                        parts.append(
                            f'<strong style="color:#374151;">{label_esc}:</strong> '
                            f'<span style="color:#111827;">{val_esc}</span>'
                        )
                    else:
                        parts.append(f'<strong style="color:#374151;">{label_esc}:</strong>')
                elif clean_val:
                    parts.append(
                        f'<strong style="color:#374151;">{label_esc}:</strong> '
                        f'<span contenteditable="true" '
                        f'style="color:#1e40af; border-bottom:1px dashed #93c5fd; '
                        f'padding:1px 4px;">{val_esc}</span>'
                    )
                else:
                    parts.append(
                        f'<strong style="color:#374151;">{label_esc}:</strong> '
                        f'<span contenteditable="true" '
                        f'style="color:#9ca3af; border-bottom:1px dashed #d1d5db; '
                        f'padding:1px 4px; min-width:80px; display:inline-block;">'
                        f'&nbsp;</span>'
                    )

            return '&emsp;'.join(parts) if parts else esc(text)

        def para_to_html(para) -> str:
            text = para.text.strip()
            if not text:
                return '<div style="height:6px;"></div>'

            style_name = (para.style.name or '').lower() if para.style else ''
            align = get_align(para)
            inline = runs_to_html(para)

            # --- Hide lines that are only underscores ---
            if all(c in '_ \t\n' for c in text):
                return '<hr style="border:none; border-top:1px solid #e5e7eb; margin:4px 0;"/>'

            # --- Hide standalone "$" or "$ 25" lines (duplicate rate echoes) ---
            stripped = text.strip()
            if stripped == '$' or (stripped.startswith('$') and len(stripped) < 10
                                  and all(c in '$ 0123456789.,\t' for c in stripped)):
                return ''

            # --- Section heading detection ---
            is_heading = 'heading' in style_name or 'title' in style_name
            if not is_heading and text.isupper() and len(text) > 3:
                is_heading = True

            if is_heading or 'title' in style_name:
                level = '1' if 'title' in style_name or 'heading 1' in style_name else '2'
                return (
                    f'<div style="margin-top:28px; margin-bottom:10px; padding:8px 12px; '
                    f'background:linear-gradient(135deg,#1e3a8a 0%,#3b82f6 100%); '
                    f'color:white; font-size:{"16px" if level == "1" else "14px"}; '
                    f'font-weight:700; letter-spacing:0.5px; border-radius:4px; {align}">'
                    f'{inline}</div>'
                )
            elif 'heading 3' in style_name:
                return (
                    f'<div style="margin-top:16px; margin-bottom:6px; font-weight:600; '
                    f'color:#374151; font-size:13px; {align}">{inline}</div>'
                )

            # --- Determine if this is a short form-field line vs. a content paragraph ---
            colon_count = text.count(':')
            text_after_last_colon = text.rsplit(':', 1)[-1].strip() if ':' in text else ''
            # It's a form line if: has colons, is short, and content after last colon is short
            is_form_line = (
                colon_count > 0
                and len(text) < 120
                and len(text_after_last_colon) < 60
                and not text.startswith('(')
            )

            if is_form_line:
                filled = smart_fill_paragraph(text)
                if filled and filled != esc(text):
                    return (
                        f'<div style="margin:4px 0; padding:6px 8px; font-size:12px; '
                        f'background:#fafafa; border-radius:4px; border-left:3px solid #e5e7eb; '
                        f'{align}">{filled}</div>'
                    )

            # --- Default: regular paragraph ---
            return (
                f'<p style="margin:3px 0; font-size:12px; {align}">{inline}</p>'
            )

        def table_to_html(table) -> str:
            rows_html = []
            num_cols = max((len(row.cells) for row in table.rows), default=2)
            is_form_table = num_cols == 2

            for row_idx, row in enumerate(table.rows):
                cells = list(row.cells)
                cells_html = []

                if is_form_table and len(cells) == 2:
                    label_text = cells[0].text.strip()
                    value_text = cells[1].text.strip()

                    value_is_blank = (
                        not value_text
                        or all(c in '_ \t\n*$' for c in value_text)
                    )
                    if value_is_blank:
                        resolved = resolve_label(label_text)
                        if resolved:
                            value_text = resolved

                    # Signature rows get a line, not "click to edit"
                    is_signature = 'signature' in label_text.lower()

                    label_esc = esc(label_text)
                    value_esc = esc(value_text)

                    cells_html.append(
                        f'<td style="padding:8px 14px; border:1px solid #e5e7eb; '
                        f'font-weight:600; font-size:12px; color:#374151; '
                        f'background:#f9fafb; width:40%; vertical-align:top;">'
                        f'{label_esc}</td>'
                    )
                    if is_signature and not value_text:
                        cells_html.append(
                            f'<td style="padding:8px 14px; border:1px solid #e5e7eb; font-size:12px;">'
                            f'<span style="display:inline-block; width:200px; '
                            f'border-bottom:1px solid #374151;">&nbsp;</span></td>'
                        )
                    elif value_text:
                        cells_html.append(
                            f'<td contenteditable="true" style="padding:8px 14px; border:1px solid #e5e7eb; '
                            f'font-size:12px; color:#1e40af; font-weight:500;">'
                            f'{value_esc}</td>'
                        )
                    else:
                        cells_html.append(
                            f'<td contenteditable="true" style="padding:8px 14px; border:1px solid #e5e7eb; '
                            f'font-size:12px; color:#9ca3af;">'
                            f'&nbsp;</td>'
                        )
                else:
                    for ci, cell in enumerate(cells):
                        ct = cell.text.strip()
                        # Skip cells that are only $ or underscores
                        if all(c in '_ \t\n$' for c in ct) and ct.strip() in ('$', ''):
                            ct = ''
                        ct_esc = esc(ct)
                        bg = '#f9fafb' if row_idx == 0 else ('#ffffff' if row_idx % 2 == 1 else '#fafafa')
                        fw = '600' if row_idx == 0 else '400'
                        cells_html.append(
                            f'<td contenteditable="true" style="padding:8px 14px; border:1px solid #e5e7eb; '
                            f'font-size:12px; font-weight:{fw}; background:{bg};">'
                            f'{ct_esc or "&nbsp;"}</td>'
                        )

                rows_html.append(f'<tr>{"".join(cells_html)}</tr>')

            return (
                '<table style="width:100%; border-collapse:collapse; '
                'margin:8px 0; border-radius:6px; overflow:hidden; '
                'border:1px solid #e5e7eb;">'
                f'{"".join(rows_html)}</table>'
            )

        body = doc.element.body
        for child in body:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'p':
                for para in doc.paragraphs:
                    if para._element is child:
                        html.append(para_to_html(para))
                        break
            elif tag == 'tbl':
                for table in doc.tables:
                    if table._tbl is child:
                        html.append(table_to_html(table))
                        break

        html.append('</div>')
        return '\n'.join(html)

    except Exception as e:
        logger.error(f"DOCX to HTML conversion failed: {e}")
        return ""


def generate_contract_from_uploaded_template(
    client: Any, 
    contract: Any, 
    template_base64: str,
    agency_settings: Optional[Any] = None
) -> bytes:
    """
    Generate a contract DOCX by filling in an uploaded template.
    
    Args:
        client: Client model with name, address, etc.
        contract: Contract model with services, schedule, rates
        template_base64: Base64-encoded DOCX template
        agency_settings: Optional agency settings for agency info
    
    Returns:
        DOCX file as bytes with placeholders replaced
    """
    # Decode the template
    template_bytes = base64.b64decode(template_base64)
    
    # Get all placeholder values
    placeholders = get_template_placeholders(client, contract, agency_settings)
    
    # Fill the template
    filled_doc = fill_docx_template(template_bytes, placeholders)
    
    return filled_doc


