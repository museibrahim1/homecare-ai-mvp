"""Native python-docx generation for notes and contracts (legacy DOCX output)."""

import io
import re
import base64
import logging
from typing import Any, List, Dict, Optional
from datetime import date

logger = logging.getLogger(__name__)


def generate_note_docx(visit: Any, note: Any) -> bytes:
    """Legacy DOCX generation - use generate_note_pdf instead."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    title = doc.add_heading("Visit Note", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading("Visit Information", level=1)
    doc.add_paragraph(f"Client: {visit.client.full_name if visit.client else 'N/A'}")
    doc.add_paragraph(f"Caregiver: {visit.caregiver.full_name if visit.caregiver else 'N/A'}")
    doc.add_paragraph(f"Date: {str(visit.actual_start or visit.scheduled_start or 'N/A')}")
    
    if note.narrative:
        doc.add_heading("Narrative", level=1)
        doc.add_paragraph(note.narrative)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_contract_docx(client: Any, contract: Any) -> bytes:
    """
    Generate a comprehensive DOCX contract using all available assessment data.
    This is used as fallback when no custom template is uploaded.
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # ==================== HEADER ====================
    title = doc.add_heading("HOME CARE SERVICE AGREEMENT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Effective Date: {date.today().strftime('%B %d, %Y')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # ==================== PARTIES ====================
    doc.add_heading("1. PARTIES", level=1)
    
    # Client info
    p = doc.add_paragraph()
    p.add_run("CLIENT:\n").bold = True
    p.add_run(f"Name: {client.full_name or 'N/A'}\n")
    p.add_run(f"Address: {client.address or 'N/A'}\n")
    if hasattr(client, 'city') and client.city:
        p.add_run(f"City/State/Zip: {client.city or ''}, {client.state or ''} {client.zip_code or ''}\n")
    p.add_run(f"Phone: {client.phone or 'N/A'}\n")
    if client.emergency_contact_name:
        p.add_run(f"Emergency Contact: {client.emergency_contact_name} ({client.emergency_contact_phone or 'N/A'})\n")
    
    doc.add_paragraph()
    
    # ==================== CARE ASSESSMENT ====================
    schedule = contract.schedule or {}
    client_profile = schedule.get('client_profile', {})
    care_need_level = schedule.get('care_need_level', 'MODERATE')
    
    doc.add_heading("2. CARE ASSESSMENT SUMMARY", level=1)
    
    p = doc.add_paragraph()
    p.add_run(f"Care Need Level: ").bold = True
    p.add_run(f"{care_need_level}\n")
    
    if client_profile.get('primary_diagnosis'):
        p.add_run(f"Primary Diagnosis: ").bold = True
        p.add_run(f"{client_profile['primary_diagnosis']}\n")
    
    if client_profile.get('secondary_conditions'):
        conditions = client_profile['secondary_conditions']
        if isinstance(conditions, list):
            conditions = ', '.join(conditions)
        p.add_run(f"Secondary Conditions: ").bold = True
        p.add_run(f"{conditions}\n")
    
    if client_profile.get('mobility_status'):
        p.add_run(f"Mobility Status: ").bold = True
        p.add_run(f"{client_profile['mobility_status']}\n")
    
    if client_profile.get('cognitive_status'):
        p.add_run(f"Cognitive Status: ").bold = True
        p.add_run(f"{client_profile['cognitive_status']}\n")
    
    if client_profile.get('living_situation'):
        p.add_run(f"Living Situation: ").bold = True
        p.add_run(f"{client_profile['living_situation']}\n")
    
    # ==================== SERVICES ====================
    doc.add_heading("3. SERVICES TO BE PROVIDED", level=1)
    
    services = contract.services or []
    if services:
        for i, service in enumerate(services, 1):
            if isinstance(service, str):
                doc.add_paragraph(f"{i}. {service}", style='List Number')
            else:
                name = service.get('name', 'Service')
                desc = service.get('description', '')
                freq = service.get('frequency', '')
                priority = service.get('priority', '')
                
                p = doc.add_paragraph(style='List Number')
                p.add_run(f"{name}").bold = True
                if desc:
                    p.add_run(f"\n   Description: {desc}")
                if freq:
                    p.add_run(f"\n   Frequency: {freq}")
                if priority:
                    p.add_run(f"\n   Priority: {priority}")
    else:
        doc.add_paragraph("Services to be determined based on care plan.")
    
    # ==================== SCHEDULE ====================
    doc.add_heading("4. SCHEDULE", level=1)
    
    p = doc.add_paragraph()
    
    frequency = schedule.get('frequency', 'As scheduled')
    p.add_run(f"Frequency: ").bold = True
    p.add_run(f"{frequency}\n")
    
    days = schedule.get('preferred_days', [])
    if isinstance(days, list):
        days_str = ', '.join([d if isinstance(d, str) else d.get('day', str(d)) for d in days]) if days else 'To be determined'
    else:
        days_str = str(days) if days else 'To be determined'
    p.add_run(f"Days: ").bold = True
    p.add_run(f"{days_str}\n")
    
    times = schedule.get('preferred_times', 'Flexible')
    p.add_run(f"Preferred Time: ").bold = True
    p.add_run(f"{times}\n")
    
    weekly_hours = float(contract.weekly_hours or 0)
    p.add_run(f"Hours per Week: ").bold = True
    p.add_run(f"{weekly_hours:.0f}\n")
    
    # ==================== RATES ====================
    doc.add_heading("5. SERVICE RATES", level=1)
    
    hourly_rate = float(contract.hourly_rate or 0)
    weekly_cost = hourly_rate * weekly_hours
    monthly_cost = weekly_cost * 4.33
    
    p = doc.add_paragraph()
    p.add_run(f"Hourly Rate: ").bold = True
    p.add_run(f"${hourly_rate:.2f}/hour\n")
    p.add_run(f"Estimated Weekly Cost: ").bold = True
    p.add_run(f"${weekly_cost:.2f}\n")
    p.add_run(f"Estimated Monthly Cost: ").bold = True
    p.add_run(f"${monthly_cost:.2f}\n")
    p.add_run(f"\nPayment Terms: ").bold = True
    p.add_run("Payment due within 7 days of invoice.\n")
    
    # ==================== SPECIAL REQUIREMENTS ====================
    special_reqs = schedule.get('special_requirements', [])
    if special_reqs:
        doc.add_heading("6. SPECIAL REQUIREMENTS", level=1)
        for req in special_reqs:
            if isinstance(req, str):
                doc.add_paragraph(f"• {req}", style='List Bullet')
            elif isinstance(req, dict):
                req_text = req.get('requirement') or req.get('name', str(req))
                doc.add_paragraph(f"• {req_text}", style='List Bullet')
    
    # ==================== SAFETY CONCERNS ====================
    safety = schedule.get('safety_concerns', [])
    if safety:
        doc.add_heading("7. SAFETY CONSIDERATIONS", level=1)
        for concern in safety:
            if isinstance(concern, str):
                doc.add_paragraph(f"• {concern}", style='List Bullet')
            elif isinstance(concern, dict):
                text = concern.get('concern', str(concern))
                severity = concern.get('severity', '')
                if severity:
                    doc.add_paragraph(f"• {text} [{severity}]", style='List Bullet')
                else:
                    doc.add_paragraph(f"• {text}", style='List Bullet')
    
    # ==================== CARE PLAN GOALS ====================
    goals = schedule.get('care_plan_goals', {})
    if goals:
        doc.add_heading("8. CARE PLAN GOALS", level=1)
        
        short_term = goals.get('short_term', [])
        if short_term:
            p = doc.add_paragraph()
            p.add_run("Short-Term Goals (30 days):").bold = True
            for goal in short_term:
                doc.add_paragraph(f"• {goal}", style='List Bullet')
        
        long_term = goals.get('long_term', [])
        if long_term:
            p = doc.add_paragraph()
            p.add_run("Long-Term Goals (90+ days):").bold = True
            for goal in long_term:
                doc.add_paragraph(f"• {goal}", style='List Bullet')
    
    # ==================== TERMS & CONDITIONS ====================
    doc.add_heading("9. TERMS AND CONDITIONS", level=1)
    
    doc.add_paragraph("""CANCELLATION POLICY: Cancellations made more than 24 hours in advance incur no charge. Cancellations within 24 hours may incur 50% of scheduled visit fee. Emergency cancellations will be evaluated on a case-by-case basis.

CONFIDENTIALITY: All client information is kept strictly confidential in accordance with HIPAA regulations. Information will only be shared with healthcare providers directly involved in the client's care or as required by law.

TERMINATION: Either party may terminate this Agreement with 14 days written notice. Immediate termination may occur if there is a safety concern or non-payment of services.

LIABILITY: The Agency maintains comprehensive general liability and professional liability insurance. Caregivers are employees of the Agency and covered under workers' compensation insurance.""")
    
    # ==================== SIGNATURES ====================
    doc.add_heading("10. SIGNATURES", level=1)
    
    doc.add_paragraph("By signing below, both parties agree to the terms and conditions of this Agreement.\n")
    
    # Client signature
    p = doc.add_paragraph()
    p.add_run("CLIENT / AUTHORIZED REPRESENTATIVE:\n\n").bold = True
    p.add_run("Signature: _________________________________  Date: _______________\n\n")
    p.add_run(f"Printed Name: {client.full_name or '________________________________'}\n")
    
    doc.add_paragraph()
    
    # Agency signature
    p = doc.add_paragraph()
    p.add_run("AGENCY REPRESENTATIVE:\n\n").bold = True
    p.add_run("Signature: _________________________________  Date: _______________\n\n")
    p.add_run("Printed Name: ________________________________\n\n")
    p.add_run("Title: ________________________________\n")
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph(f"Contract ID: {contract.id}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
