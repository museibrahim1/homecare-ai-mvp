"""
Document Generation Service

Generates professional PDF documents from templates for home care contracts and notes.
Supports user-uploaded DOCX templates with placeholder replacement.
"""

import io
import re
import base64
import logging
from typing import Any, List, Dict, Optional
from datetime import date

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, 
    PageBreak, HRFlowable
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

logger = logging.getLogger(__name__)


def get_template_placeholders(client: Any, contract: Any, agency_settings: Optional[Any] = None) -> Dict[str, str]:
    """
    Build a dictionary of all template placeholders and their values.
    
    Supports placeholders like:
    - {client_name}, {{client_name}}, [client_name], [[client_name]]
    - {CLIENT_NAME} (case-insensitive)
    """
    schedule = contract.schedule or {}
    services = contract.services or []
    client_profile = schedule.get("client_profile", {})
    
    # Calculate costs
    hourly_rate = float(contract.hourly_rate or 0)
    weekly_hours = float(contract.weekly_hours or 0)
    weekly_cost = hourly_rate * weekly_hours
    monthly_cost = weekly_cost * 4.33
    
    # Format schedule days
    days = schedule.get("preferred_days", [])
    if isinstance(days, list):
        days_str = ", ".join([
            d if isinstance(d, str) else (d.get('day', str(d)) if isinstance(d, dict) else str(d)) 
            for d in days
        ]) if days else "To be determined"
    else:
        days_str = str(days) if days else "To be determined"
    
    # Format services list
    services_text = ""
    for i, svc in enumerate(services, 1):
        if isinstance(svc, str):
            services_text += f"{i}. {svc}\n"
        else:
            name = svc.get('name', 'Service')
            desc = svc.get('description', '')
            freq = svc.get('frequency', '')
            line = f"{i}. {name}"
            if desc:
                line += f": {desc}"
            if freq:
                line += f" ({freq})"
            services_text += line + "\n"
    
    # Format special requirements
    special_reqs = schedule.get("special_requirements", [])
    reqs_text = ""
    for req in special_reqs:
        if isinstance(req, str):
            reqs_text += f"• {req}\n"
        elif isinstance(req, dict):
            reqs_text += f"• {req.get('name', req.get('requirement', str(req)))}\n"
    
    # Format safety concerns
    safety = schedule.get("safety_concerns", [])
    safety_text = ""
    for concern in safety:
        if isinstance(concern, str):
            safety_text += f"• {concern}\n"
        elif isinstance(concern, dict):
            safety_text += f"• {concern.get('concern', str(concern))}\n"
    
    # Agency info
    agency_name = "Home Care Services Agency"
    agency_address = ""
    agency_city = ""
    agency_state = ""
    agency_zip = ""
    agency_phone = ""
    agency_email = ""
    
    if agency_settings:
        agency_name = agency_settings.name or agency_name
        agency_address = agency_settings.address or ""
        agency_city = agency_settings.city or ""
        agency_state = agency_settings.state or ""
        agency_zip = agency_settings.zip_code or ""
        agency_phone = agency_settings.phone or ""
        agency_email = agency_settings.email or ""
    
    # Build full agency address
    agency_full_address = agency_address
    city_state_zip = ", ".join(filter(None, [agency_city, agency_state])) 
    if agency_zip:
        city_state_zip += f" {agency_zip}"
    if city_state_zip:
        agency_full_address += f"\n{city_state_zip}" if agency_full_address else city_state_zip
    
    # All placeholders (multiple variations for flexibility)
    placeholders = {
        # Dates
        'date': date.today().strftime('%B %d, %Y'),
        'contract_date': date.today().strftime('%B %d, %Y'),
        'effective_date': date.today().strftime('%B %d, %Y'),
        'today': date.today().strftime('%B %d, %Y'),
        'current_date': date.today().strftime('%B %d, %Y'),
        
        # Agency info
        'agency_name': agency_name,
        'agency': agency_name,
        'company_name': agency_name,
        'provider_name': agency_name,
        'agency_address': agency_full_address,
        'agency_street': agency_address,
        'agency_city': agency_city,
        'agency_state': agency_state,
        'agency_zip': agency_zip,
        'agency_zip_code': agency_zip,
        'agency_phone': agency_phone,
        'agency_email': agency_email,
        
        # Client info
        'client_name': client.full_name or '',
        'client': client.full_name or '',
        'patient_name': client.full_name or '',
        'patient': client.full_name or '',
        'client_first_name': (client.full_name or '').split()[0] if client.full_name else '',
        'client_last_name': (client.full_name or '').split()[-1] if client.full_name and len((client.full_name or '').split()) > 1 else '',
        'client_address': client.address or '',
        'client_street': client.address or '',
        'client_city': client.city or '',
        'client_state': client.state or '',
        'client_zip': client.zip_code or '',
        'client_zip_code': client.zip_code or '',
        'client_phone': client.phone or '',
        'home_phone': client.phone or '',
        'work_phone': getattr(client, 'work_phone', '') or '',
        'cell_phone': client.phone or '',
        'client_email': client.email or '',
        'date_of_birth': getattr(client, 'date_of_birth', '') or '',
        'dob': getattr(client, 'date_of_birth', '') or '',
        'emergency_contact': client.emergency_contact_name or '',
        'emergency_contact_name': client.emergency_contact_name or '',
        'emergency_phone': client.emergency_contact_phone or '',
        'emergency_contact_phone': client.emergency_contact_phone or '',
        
        # Care assessment
        'care_level': schedule.get('care_need_level', 'MODERATE'),
        'care_need_level': schedule.get('care_need_level', 'MODERATE'),
        'primary_diagnosis': client_profile.get('primary_diagnosis', 'See medical records'),
        'diagnosis': client_profile.get('primary_diagnosis', 'See medical records'),
        'mobility_status': client_profile.get('mobility_status', 'N/A'),
        'mobility': client_profile.get('mobility_status', 'N/A'),
        'cognitive_status': client_profile.get('cognitive_status', 'N/A'),
        'living_situation': client_profile.get('living_situation', 'N/A'),
        
        # Services
        'services': services_text.strip() or 'As determined by care plan',
        'services_list': services_text.strip() or 'As determined by care plan',
        
        # Schedule
        'schedule_days': days_str,
        'days': days_str,
        'preferred_days': days_str,
        'schedule_time': schedule.get('preferred_times', 'Flexible'),
        'preferred_time': schedule.get('preferred_times', 'Flexible'),
        'time': schedule.get('preferred_times', 'Flexible'),
        'frequency': schedule.get('frequency', 'As scheduled'),
        'weekly_hours': f"{weekly_hours:.0f}",
        'hours_per_week': f"{weekly_hours:.0f}",
        
        # Rates
        'hourly_rate': f"${hourly_rate:.2f}",
        'hourly_rate_value': f"{hourly_rate:.2f}",  # Without $ sign
        'rate': f"${hourly_rate:.2f}",
        'rate_value': f"{hourly_rate:.2f}",
        'weekly_cost': f"${weekly_cost:.2f}",
        'weekly_estimate': f"${weekly_cost:.2f}",
        'monthly_cost': f"${monthly_cost:.2f}",
        'monthly_estimate': f"${monthly_cost:.2f}",
        'weekday_rate': f"{hourly_rate:.2f}",
        'weekend_rate': f"{hourly_rate * 1.25:.2f}",  # 25% premium
        'holiday_rate': f"{hourly_rate * 1.5:.2f}",  # Time and a half
        
        # Requirements and safety
        'special_requirements': reqs_text.strip() or 'None specified',
        'requirements': reqs_text.strip() or 'None specified',
        'safety_concerns': safety_text.strip() or 'None noted',
        'safety': safety_text.strip() or 'None noted',

        # Contract ID
        'contract_id': str(contract.id) if contract.id else '',

        # Additional billing fields for custom templates
        'admin_fee': '25',
        'deposit': f"${weekly_cost:.2f}" if weekly_cost else '',
        'prepayment': '',
        'ssn': '',
        'medicaid_number': '',
        'medicare_number': '',
        'physician_name': '',
        'physician_phone': '',

        # Cancellation / Terms / Policies (long text)
        'cancellation_policy': getattr(contract, 'cancellation_policy', '') or 'Either party may terminate this agreement with 30 days written notice.',
        'terms_and_conditions': getattr(contract, 'terms_and_conditions', '') or '',
        'policies_and_procedures': getattr(contract, 'policies_and_procedures', '') or '',
    }
    
    return placeholders


def fill_docx_template(template_bytes: bytes, placeholders: Dict[str, str]) -> bytes:
    """
    Fill a DOCX template by replacing placeholders with actual values.
    
    Supports multiple placeholder formats:
    - {placeholder}, {{placeholder}}, [placeholder], [[placeholder]]
    - "Label:" format (e.g., "Name:" followed by blank space)
    - Case insensitive matching
    """
    try:
        from docx import Document

        doc = Document(io.BytesIO(template_bytes))

        # Ordered list of (label_pattern, placeholder_key).
        # Longer/more-specific labels come first so they match before shorter ones.
        LABEL_MAPPINGS = [
            # Client / Patient fields (specific first)
            ('client name:', 'client_name'),
            ('patient name:', 'client_name'),
            ('client address:', 'client_address'),
            ('client city:', 'client_city'),
            ('client state:', 'client_state'),
            ('client zip:', 'client_zip'),
            ('client phone:', 'client_phone'),
            ('client email:', 'client_email'),
            ('emergency contact name:', 'emergency_contact'),
            ('emergency contact:', 'emergency_contact'),
            ('emergency phone:', 'emergency_phone'),
            ('date of birth:', 'date_of_birth'),
            ('dob:', 'date_of_birth'),
            ('home phone:', 'client_phone'),
            ('work phone:', 'work_phone'),
            ('cell phone:', 'client_phone'),
            ('social security:', 'ssn'),
            ('ssn:', 'ssn'),
            ('medicaid #:', 'medicaid_number'),
            ('medicaid number:', 'medicaid_number'),
            ('medicare #:', 'medicare_number'),
            ('medicare number:', 'medicare_number'),
            ('physician name:', 'physician_name'),
            ('physician phone:', 'physician_phone'),
            ('physician:', 'physician_name'),

            # Agency fields
            ('agency name:', 'agency_name'),
            ('agency address:', 'agency_address'),
            ('agency phone:', 'agency_phone'),
            ('agency email:', 'agency_email'),
            ('agency rep name:', 'agency_name'),
            ('agency rep signature:', 'agency_name'),

            # Rate & billing fields (specific first)
            ('hourly rate:', 'hourly_rate_value'),
            ('hourly:', 'hourly_rate_value'),
            ('weekday rate:', 'hourly_rate_value'),
            ('weekday:', 'hourly_rate_value'),
            ('weekend rate:', 'weekend_rate'),
            ('weekend:', 'weekend_rate'),
            ('holiday rate:', 'holiday_rate'),
            ('holiday:', 'holiday_rate'),
            ('client rate:', 'hourly_rate_value'),
            ('monthly package:', 'monthly_cost'),
            ('monthly cost:', 'monthly_cost'),
            ('monthly estimate:', 'monthly_cost'),
            ('weekly cost:', 'weekly_cost'),
            ('weekly estimate:', 'weekly_cost'),
            ('administrative fee:', 'admin_fee'),
            ('admin fee:', 'admin_fee'),
            ('deposit:', 'deposit'),
            ('prepayment:', 'prepayment'),
            ('total:', 'monthly_cost'),

            # Schedule fields
            ('weekly hours:', 'weekly_hours'),
            ('hours per week:', 'weekly_hours'),
            ('days of service:', 'schedule_days'),
            ('schedule:', 'schedule_days'),
            ('frequency:', 'frequency'),
            ('start time:', 'time'),
            ('end time:', 'time'),

            # Date fields
            ('effective date:', 'effective_date'),
            ('contract date:', 'contract_date'),
            ('start date:', 'effective_date'),
            ('end date:', 'effective_date'),
            ('agreement date:', 'contract_date'),

            # Services
            ('services to be provided:', 'services'),
            ('services provided:', 'services'),
            ('services:', 'services'),

            # Policy / Procedure / Terms fields
            ('cancellation policy:', 'cancellation_policy'),
            ('termination policy:', 'cancellation_policy'),
            ('policies and procedures:', 'policies_and_procedures'),
            ('policy and procedures:', 'policies_and_procedures'),
            ('terms and conditions:', 'terms_and_conditions'),
            ('terms of service:', 'terms_and_conditions'),
            ('special requirements:', 'special_requirements'),
            ('special instructions:', 'special_requirements'),
            ('safety considerations:', 'safety_concerns'),
            ('safety concerns:', 'safety_concerns'),
            ('living situation:', 'living_situation'),
            ('bill to:', 'client_name'),

            # Generic fallbacks (must be last)
            ('zip code:', 'client_zip'),
            ('zip:', 'client_zip'),
            ('name:', 'client_name'),
            ('address:', 'client_address'),
            ('city:', 'client_city'),
            ('state:', 'client_state'),
            ('phone:', 'client_phone'),
            ('email:', 'client_email'),
            ('date:', 'date'),
        ]

        def replace_placeholder_text(text: str, ph: Dict[str, str]) -> str:
            """Replace {{key}}, {key}, [[key]], [key], (key) patterns."""
            if not text:
                return text
            result = text
            for key, value in ph.items():
                for pat in [
                    rf'\{{\{{\s*{re.escape(key)}\s*\}}\}}',
                    rf'\{{\s*{re.escape(key)}\s*\}}',
                    rf'\[\[\s*{re.escape(key)}\s*\]\]',
                    rf'\[\s*{re.escape(key)}\s*\]',
                ]:
                    result = re.sub(pat, str(value or ''), result, flags=re.IGNORECASE)

            # Replace (Label) patterns like $ (Client Rate) -> $29.00
            # Also remove the leading "$ " to avoid "$ $29.00"
            result = re.sub(r'\$\s*\(\s*Client\s+Rate\s*\)', ph.get('hourly_rate', ''), result, flags=re.IGNORECASE)
            result = re.sub(r'\(\s*Client\s+Rate\s*\)', ph.get('hourly_rate_value', ''), result, flags=re.IGNORECASE)
            result = re.sub(r'\(\s*Client\s+Name\s*\)', ph.get('client_name', ''), result, flags=re.IGNORECASE)
            result = re.sub(r'\(\s*Agency\s+Name\s*\)', ph.get('agency_name', ''), result, flags=re.IGNORECASE)
            return result

        def fill_labels_in_text(text: str, ph: Dict[str, str]) -> str:
            """
            Replace "Label:" patterns with values.  Handles:
            - "Label: ___________"  (underscores)
            - "Label:  "           (just spaces)
            - "Label:"             (nothing after, end of line)
            - "City: State: Zip:"  (multiple labels on one line)
            """
            if not text:
                return text

            result = text
            text_lower = text.lower()
            # Track replaced ranges [start, end) to prevent overlapping fills
            replaced_ranges: list = []

            def overlaps(start: int, end: int) -> bool:
                for rs, re_ in replaced_ranges:
                    if start < re_ and end > rs:
                        return True
                return False

            for label, pk in LABEL_MAPPINGS:
                idx = text_lower.find(label)
                if idx == -1:
                    continue

                # Word boundary check: if char before match is a letter, skip
                # (prevents "services:" matching inside "Person to Receive Services:")
                if idx > 0 and result[idx - 1].isalpha():
                    continue

                # Check this position hasn't already been filled
                label_end = idx + len(label)
                if overlaps(idx, label_end):
                    continue

                value = ph.get(pk, '')
                if not value:
                    continue

                # Find the span of underscores/spaces/$ after the label
                span_end = label_end
                while span_end < len(result) and result[span_end] in '_ \t$':
                    span_end += 1

                span_content = result[label_end:span_end]

                should_fill = False
                if span_content and any(c == '_' for c in span_content):
                    should_fill = True
                elif span_end >= len(result):
                    should_fill = True
                elif span_end == label_end:
                    rest = result[label_end:].strip()
                    if not rest:
                        should_fill = True
                else:
                    if not span_content.strip() or span_content.strip() == '$':
                        should_fill = True

                if should_fill:
                    insertion = ' ' + str(value)
                    new_result = result[:label_end] + insertion + result[span_end:]
                    replaced_ranges.append((idx, label_end + len(insertion)))
                    shift = len(new_result) - len(result)
                    # Adjust existing ranges after this point
                    result = new_result
                    text_lower = result.lower()

            return result

        def set_paragraph_text(paragraph, new_text: str):
            """Overwrite a paragraph's text while keeping first run's formatting."""
            for run in paragraph.runs:
                run.text = ''
            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)

        def process_paragraph(paragraph, ph: Dict[str, str]):
            full_text = paragraph.text
            if not full_text:
                return
            new_text = replace_placeholder_text(full_text, ph)
            new_text = fill_labels_in_text(new_text, ph)
            if new_text != full_text:
                set_paragraph_text(paragraph, new_text)

        def process_cell_inline(cell, ph: Dict[str, str]):
            """
            Fill placeholders and label patterns within a cell.
            Returns the placeholder_key if this cell is a pure label cell
            (so the caller can fill the adjacent cell).
            """
            cell_text = cell.text.strip()
            if not cell_text:
                return None

            # Strip trailing asterisks/required markers for matching
            cell_clean = re.sub(r'[\s*]+$', '', cell_text)
            cell_lower = cell_clean.lower().strip()

            # Check if this cell is a pure label like "Name:", "Address", "Date of Birth:"
            for label, pk in LABEL_MAPPINGS:
                label_bare = label.rstrip(':').strip()
                if cell_lower == label or cell_lower == label_bare or cell_lower == label_bare + ':':
                    return pk

            # Also check without trailing colon for labels like "Home Phone" (no colon)
            if cell_lower.endswith(':'):
                cell_no_colon = cell_lower[:-1].strip()
            else:
                cell_no_colon = cell_lower
            for label, pk in LABEL_MAPPINGS:
                label_bare = label.rstrip(':').strip()
                if cell_no_colon == label_bare:
                    return pk

            # Fill {{placeholder}} patterns
            new_text = replace_placeholder_text(cell_text, ph)

            # Fill "Label: ___" patterns inside the cell
            new_text = fill_labels_in_text(new_text, ph)

            if new_text != cell_text:
                for paragraph in cell.paragraphs:
                    old = paragraph.text
                    if old.strip():
                        p_new = replace_placeholder_text(old, ph)
                        p_new = fill_labels_in_text(p_new, ph)
                        if p_new != old:
                            set_paragraph_text(paragraph, p_new)

            return None

        def process_table(table, ph: Dict[str, str]):
            """Process all rows/cells in a table, including nested tables."""
            for row in table.rows:
                cells = list(row.cells)
                pending_pk = None

                for i, cell in enumerate(cells):
                    if pending_pk:
                        value = ph.get(pending_pk, '')
                        if value:
                            cell_text = cell.text.strip()
                            # Fill if cell is blank, has only underscores/spaces, or has placeholder patterns
                            is_blank = (
                                not cell_text
                                or all(c in '_ \t\n*' for c in cell_text)
                                or cell_text in ('*', '**')
                            )
                            if is_blank:
                                for paragraph in cell.paragraphs:
                                    set_paragraph_text(paragraph, str(value))
                                    break
                            else:
                                # Cell has content — still try placeholder/label fill
                                process_cell_inline(cell, ph)
                        pending_pk = None
                    else:
                        pending_pk = process_cell_inline(cell, ph)

                    # Process nested tables inside this cell
                    for nested_table in cell.tables:
                        process_table(nested_table, ph)

        # --- Process document body paragraphs ---
        for paragraph in doc.paragraphs:
            process_paragraph(paragraph, placeholders)

        # --- Process tables (including nested) ---
        for table in doc.tables:
            process_table(table, placeholders)

        # --- Process headers / footers ---
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    process_paragraph(paragraph, placeholders)
                for table in section.header.tables:
                    process_table(table, placeholders)
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    process_paragraph(paragraph, placeholders)
                for table in section.footer.tables:
                    process_table(table, placeholders)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    except Exception as e:
        logger.error(f"Failed to fill DOCX template: {e}")
        raise


def docx_to_html(template_bytes: bytes, placeholders: Dict[str, str]) -> str:
    """
    Convert a DOCX template to styled HTML with placeholders filled in.
    Renders the full document for in-browser preview.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        filled_bytes = fill_docx_template(template_bytes, placeholders)
        doc = Document(io.BytesIO(filled_bytes))

        # Label -> placeholder key for second-pass smart fill in the HTML
        LABEL_MAP = {}
        for key in placeholders:
            nice = key.replace('_', ' ').title()
            LABEL_MAP[key.lower()] = key
            LABEL_MAP[nice.lower()] = key
            LABEL_MAP[nice.lower() + ':'] = key
        # Explicit common labels
        for lbl, pk in [
            ('name', 'client_name'), ('address', 'client_address'),
            ('city', 'client_city'), ('state', 'client_state'),
            ('zip', 'client_zip'), ('zip code', 'client_zip'),
            ('phone', 'client_phone'), ('home phone', 'client_phone'),
            ('work phone', 'work_phone'), ('cell phone', 'client_phone'),
            ('email', 'client_email'), ('date of birth', 'date_of_birth'),
            ('dob', 'date_of_birth'), ('date', 'date'),
            ('hourly rate', 'hourly_rate'), ('hourly', 'hourly_rate_value'),
            ('weekday', 'hourly_rate_value'), ('weekend', 'weekend_rate'),
            ('holiday', 'holiday_rate'), ('deposit', 'deposit'),
            ('prepayment', 'prepayment'), ('total', 'monthly_cost'),
            ('monthly package', 'monthly_cost'), ('administrative fee', 'admin_fee'),
            ('admin fee', 'admin_fee'), ('weekly hours', 'weekly_hours'),
            ('hours per week', 'weekly_hours'), ('effective date', 'effective_date'),
            ('start date', 'effective_date'), ('contract date', 'contract_date'),
            ('emergency contact', 'emergency_contact'), ('emergency phone', 'emergency_phone'),
            ('signature', ''), ('client signature', ''), ('agency representative', ''),
        ]:
            LABEL_MAP[lbl] = pk
            LABEL_MAP[lbl + ':'] = pk

        def smart_fill_cell(text: str) -> str:
            """Try to resolve a cell's content — if it's blank/underscores, keep as-is.
            If it contains a label with blanks, fill the blanks."""
            if not text:
                return text
            clean = text.strip()
            # Already has meaningful content (not just underscores)
            if clean and not all(c in '_ \t\n*$' for c in clean):
                return text
            return text

        def resolve_label(label_text: str) -> str:
            """Given a label like 'Name:' or 'Home Phone *', find the value."""
            clean = re.sub(r'[\s:*]+$', '', label_text).strip().lower()
            pk = LABEL_MAP.get(clean, '') or LABEL_MAP.get(clean + ':', '')
            if pk:
                return placeholders.get(pk, '')
            # Try partial
            for key, pkey in LABEL_MAP.items():
                if key and clean.endswith(key):
                    val = placeholders.get(pkey, '')
                    if val:
                        return val
            return ''

        def esc(text: str) -> str:
            return (text or '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

        html = []
        html.append(
            '<div style="font-family:Calibri,Arial,sans-serif; color:#1a1a1a; '
            'line-height:1.6; max-width:100%;">'
        )

        def runs_to_html(para) -> str:
            parts = []
            for run in para.runs:
                t = run.text
                if not t:
                    continue
                t = esc(t)
                if run.bold and run.italic:
                    t = f'<strong><em>{t}</em></strong>'
                elif run.bold:
                    t = f'<strong>{t}</strong>'
                elif run.italic:
                    t = f'<em>{t}</em>'
                if run.underline:
                    t = f'<u>{t}</u>'
                parts.append(t)
            return ''.join(parts) or esc(para.text)

        def get_align(para) -> str:
            try:
                if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    return 'text-align:center;'
                elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    return 'text-align:right;'
                elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                    return 'text-align:justify;'
            except Exception:
                pass
            return ''

        # Known multi-word labels for paragraph smart-fill (ordered long-first)
        PARA_LABELS = [
            ('person to receive services', None),  # skip — not a fillable field
            ('date of birth', 'date_of_birth'),
            ('home phone', 'client_phone'),
            ('work phone', 'work_phone'),
            ('cell phone', 'client_phone'),
            ('emergency contact', 'emergency_contact'),
            ('emergency phone', 'emergency_phone'),
            ('zip code', 'client_zip'),
            ('administrative fee', 'admin_fee'),
            ('admin fee', 'admin_fee'),
            ('monthly package', 'monthly_cost'),
            ('agency visits', None),
            ('hourly rate', 'hourly_rate'),
            ('weekly hours', 'weekly_hours'),
            ('bill to', None),  # section header
            ('signature', None),  # render as line
            ('name', 'client_name'),
            ('address', 'client_address'),
            ('city', 'client_city'),
            ('state', 'client_state'),
            ('zip', 'client_zip'),
            ('phone', 'client_phone'),
            ('email', 'client_email'),
            ('date', 'date'),
            ('hourly', 'hourly_rate'),
            ('weekday', 'hourly_rate_value'),
            ('weekend', 'weekend_rate'),
            ('holiday', 'holiday_rate'),
            ('prepayment', 'prepayment'),
            ('deposit', 'deposit'),
            ('total', 'monthly_cost'),
        ]

        def smart_fill_paragraph(text: str) -> str:
            """
            Parse paragraph with 'Label:' patterns using known label list.
            Returns HTML with bold labels and filled/editable values.
            """
            if not text or ':' not in text:
                return esc(text)

            # Find all label positions using the known label list
            text_lower = text.lower()
            found_labels = []  # [(start, end_of_label_colon, label_text, pk)]

            for lbl, pk in PARA_LABELS:
                search_from = 0
                while True:
                    idx = text_lower.find(lbl, search_from)
                    if idx == -1:
                        break
                    # Must be preceded by start-of-string or non-alpha
                    if idx > 0 and text_lower[idx - 1].isalpha():
                        search_from = idx + 1
                        continue
                    # Find the colon after the label
                    colon_pos = idx + len(lbl)
                    while colon_pos < len(text) and text[colon_pos] in ' \t':
                        colon_pos += 1
                    if colon_pos < len(text) and text[colon_pos] == ':':
                        found_labels.append((idx, colon_pos + 1, lbl, pk))
                    search_from = idx + 1

            if not found_labels:
                return esc(text)

            # Sort by position, remove overlapping
            found_labels.sort(key=lambda x: x[0])
            filtered = []
            last_end = -1
            for start, end, lbl, pk in found_labels:
                if start >= last_end:
                    filtered.append((start, end, lbl, pk))
                    last_end = end

            # Build HTML segments
            parts = []
            for i, (start, end, lbl, pk) in enumerate(filtered):
                # Get value text between this label's colon and next label's start
                val_start = end
                val_end = filtered[i + 1][0] if i + 1 < len(filtered) else len(text)
                raw_value = text[val_start:val_end].strip()

                # Clean value: strip leading $, underscores, asterisks
                clean_val = raw_value.strip(' \t_*$')

                # Try to fill if empty
                if not clean_val and pk:
                    clean_val = placeholders.get(pk, '')

                label_display = text[start:end - 1].strip()  # original case
                label_esc = esc(label_display)
                val_esc = esc(clean_val)

                if pk is None and lbl == 'signature':
                    parts.append(
                        f'<strong>{label_esc}:</strong> '
                        f'<span style="display:inline-block; width:200px; '
                        f'border-bottom:1px solid #374151; margin:0 8px;">&nbsp;</span>'
                    )
                elif pk is None:
                    # Non-fillable label (section header like "Bill To", "Person to Receive Services")
                    if clean_val:
                        parts.append(
                            f'<strong style="color:#374151;">{label_esc}:</strong> '
                            f'<span style="color:#111827;">{val_esc}</span>'
                        )
                    else:
                        parts.append(f'<strong style="color:#374151;">{label_esc}:</strong>')
                elif clean_val:
                    parts.append(
                        f'<strong style="color:#374151;">{label_esc}:</strong> '
                        f'<span contenteditable="true" '
                        f'style="color:#1e40af; border-bottom:1px dashed #93c5fd; '
                        f'padding:1px 4px;">{val_esc}</span>'
                    )
                else:
                    parts.append(
                        f'<strong style="color:#374151;">{label_esc}:</strong> '
                        f'<span contenteditable="true" '
                        f'style="color:#9ca3af; border-bottom:1px dashed #d1d5db; '
                        f'padding:1px 4px; min-width:80px; display:inline-block;">'
                        f'&nbsp;</span>'
                    )

            return '&emsp;'.join(parts) if parts else esc(text)

        def para_to_html(para) -> str:
            text = para.text.strip()
            if not text:
                return '<div style="height:6px;"></div>'

            style_name = (para.style.name or '').lower() if para.style else ''
            align = get_align(para)
            inline = runs_to_html(para)

            # --- Hide lines that are only underscores ---
            if all(c in '_ \t\n' for c in text):
                return '<hr style="border:none; border-top:1px solid #e5e7eb; margin:4px 0;"/>'

            # --- Hide standalone "$" or "$ 25" lines (duplicate rate echoes) ---
            stripped = text.strip()
            if stripped == '$' or (stripped.startswith('$') and len(stripped) < 10
                                  and all(c in '$ 0123456789.,\t' for c in stripped)):
                return ''

            # --- Section heading detection ---
            is_heading = 'heading' in style_name or 'title' in style_name
            if not is_heading and text.isupper() and len(text) > 3:
                is_heading = True

            if is_heading or 'title' in style_name:
                level = '1' if 'title' in style_name or 'heading 1' in style_name else '2'
                return (
                    f'<div style="margin-top:28px; margin-bottom:10px; padding:8px 12px; '
                    f'background:linear-gradient(135deg,#1e3a8a 0%,#3b82f6 100%); '
                    f'color:white; font-size:{"16px" if level == "1" else "14px"}; '
                    f'font-weight:700; letter-spacing:0.5px; border-radius:4px; {align}">'
                    f'{inline}</div>'
                )
            elif 'heading 3' in style_name:
                return (
                    f'<div style="margin-top:16px; margin-bottom:6px; font-weight:600; '
                    f'color:#374151; font-size:13px; {align}">{inline}</div>'
                )

            # --- Determine if this is a short form-field line vs. a content paragraph ---
            colon_count = text.count(':')
            text_after_last_colon = text.rsplit(':', 1)[-1].strip() if ':' in text else ''
            # It's a form line if: has colons, is short, and content after last colon is short
            is_form_line = (
                colon_count > 0
                and len(text) < 120
                and len(text_after_last_colon) < 60
                and not text.startswith('(')
            )

            if is_form_line:
                filled = smart_fill_paragraph(text)
                if filled and filled != esc(text):
                    return (
                        f'<div style="margin:4px 0; padding:6px 8px; font-size:12px; '
                        f'background:#fafafa; border-radius:4px; border-left:3px solid #e5e7eb; '
                        f'{align}">{filled}</div>'
                    )

            # --- Default: regular paragraph ---
            return (
                f'<p style="margin:3px 0; font-size:12px; {align}">{inline}</p>'
            )

        def table_to_html(table) -> str:
            rows_html = []
            num_cols = max((len(row.cells) for row in table.rows), default=2)
            is_form_table = num_cols == 2

            for row_idx, row in enumerate(table.rows):
                cells = list(row.cells)
                cells_html = []

                if is_form_table and len(cells) == 2:
                    label_text = cells[0].text.strip()
                    value_text = cells[1].text.strip()

                    value_is_blank = (
                        not value_text
                        or all(c in '_ \t\n*$' for c in value_text)
                    )
                    if value_is_blank:
                        resolved = resolve_label(label_text)
                        if resolved:
                            value_text = resolved

                    # Signature rows get a line, not "click to edit"
                    is_signature = 'signature' in label_text.lower()

                    label_esc = esc(label_text)
                    value_esc = esc(value_text)

                    cells_html.append(
                        f'<td style="padding:8px 14px; border:1px solid #e5e7eb; '
                        f'font-weight:600; font-size:12px; color:#374151; '
                        f'background:#f9fafb; width:40%; vertical-align:top;">'
                        f'{label_esc}</td>'
                    )
                    if is_signature and not value_text:
                        cells_html.append(
                            f'<td style="padding:8px 14px; border:1px solid #e5e7eb; font-size:12px;">'
                            f'<span style="display:inline-block; width:200px; '
                            f'border-bottom:1px solid #374151;">&nbsp;</span></td>'
                        )
                    elif value_text:
                        cells_html.append(
                            f'<td contenteditable="true" style="padding:8px 14px; border:1px solid #e5e7eb; '
                            f'font-size:12px; color:#1e40af; font-weight:500;">'
                            f'{value_esc}</td>'
                        )
                    else:
                        cells_html.append(
                            f'<td contenteditable="true" style="padding:8px 14px; border:1px solid #e5e7eb; '
                            f'font-size:12px; color:#9ca3af;">'
                            f'&nbsp;</td>'
                        )
                else:
                    for ci, cell in enumerate(cells):
                        ct = cell.text.strip()
                        # Skip cells that are only $ or underscores
                        if all(c in '_ \t\n$' for c in ct) and ct.strip() in ('$', ''):
                            ct = ''
                        ct_esc = esc(ct)
                        bg = '#f9fafb' if row_idx == 0 else ('#ffffff' if row_idx % 2 == 1 else '#fafafa')
                        fw = '600' if row_idx == 0 else '400'
                        cells_html.append(
                            f'<td contenteditable="true" style="padding:8px 14px; border:1px solid #e5e7eb; '
                            f'font-size:12px; font-weight:{fw}; background:{bg};">'
                            f'{ct_esc or "&nbsp;"}</td>'
                        )

                rows_html.append(f'<tr>{"".join(cells_html)}</tr>')

            return (
                '<table style="width:100%; border-collapse:collapse; '
                'margin:8px 0; border-radius:6px; overflow:hidden; '
                'border:1px solid #e5e7eb;">'
                f'{"".join(rows_html)}</table>'
            )

        body = doc.element.body
        for child in body:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'p':
                for para in doc.paragraphs:
                    if para._element is child:
                        html.append(para_to_html(para))
                        break
            elif tag == 'tbl':
                for table in doc.tables:
                    if table._tbl is child:
                        html.append(table_to_html(table))
                        break

        html.append('</div>')
        return '\n'.join(html)

    except Exception as e:
        logger.error(f"DOCX to HTML conversion failed: {e}")
        return ""


def generate_contract_from_uploaded_template(
    client: Any, 
    contract: Any, 
    template_base64: str,
    agency_settings: Optional[Any] = None
) -> bytes:
    """
    Generate a contract DOCX by filling in an uploaded template.
    
    Args:
        client: Client model with name, address, etc.
        contract: Contract model with services, schedule, rates
        template_base64: Base64-encoded DOCX template
        agency_settings: Optional agency settings for agency info
    
    Returns:
        DOCX file as bytes with placeholders replaced
    """
    # Decode the template
    template_bytes = base64.b64decode(template_base64)
    
    # Get all placeholder values
    placeholders = get_template_placeholders(client, contract, agency_settings)
    
    # Fill the template
    filled_doc = fill_docx_template(template_bytes, placeholders)
    
    return filled_doc


def get_custom_styles():
    """Create custom paragraph styles for the PDF."""
    styles = getSampleStyleSheet()
    
    # Title style
    styles.add(ParagraphStyle(
        name='ContractTitle',
        parent=styles['Heading1'],
        fontSize=18,
        alignment=TA_CENTER,
        spaceAfter=6,
        textColor=colors.HexColor('#1a365d'),
        fontName='Helvetica-Bold'
    ))
    
    # Section header
    styles.add(ParagraphStyle(
        name='SectionHeader',
        parent=styles['Heading2'],
        fontSize=12,
        spaceBefore=16,
        spaceAfter=8,
        textColor=colors.HexColor('#2c5282'),
        fontName='Helvetica-Bold',
        borderPadding=4,
    ))
    
    # Body text - modify existing instead of adding
    if 'BodyText' in styles.byName:
        styles['BodyText'].fontSize = 10
        styles['BodyText'].alignment = TA_JUSTIFY
        styles['BodyText'].spaceAfter = 6
        styles['BodyText'].leading = 14
    else:
        styles.add(ParagraphStyle(
            name='BodyText',
            parent=styles['Normal'],
            fontSize=10,
            alignment=TA_JUSTIFY,
            spaceAfter=6,
            leading=14
        ))
    
    # Bullet item
    styles.add(ParagraphStyle(
        name='BulletItem',
        parent=styles['Normal'],
        fontSize=10,
        leftIndent=20,
        spaceAfter=4,
        bulletIndent=10,
        leading=12
    ))
    
    # Label style (bold)
    styles.add(ParagraphStyle(
        name='Label',
        parent=styles['Normal'],
        fontSize=10,
        fontName='Helvetica-Bold',
        spaceAfter=2
    ))
    
    # Small text
    styles.add(ParagraphStyle(
        name='SmallText',
        parent=styles['Normal'],
        fontSize=8,
        textColor=colors.gray,
        alignment=TA_CENTER
    ))
    
    # Signature line
    styles.add(ParagraphStyle(
        name='SignatureLine',
        parent=styles['Normal'],
        fontSize=10,
        spaceBefore=20,
        spaceAfter=4
    ))
    
    return styles


def generate_contract_pdf(client: Any, contract: Any) -> bytes:
    """
    Generate a professional one-page PDF contract for home care services.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.5*inch,
        leftMargin=0.5*inch,
        topMargin=0.5*inch,
        bottomMargin=0.5*inch
    )
    
    styles = get_custom_styles()
    story = []
    
    schedule = contract.schedule or {}
    services = contract.services or []
    
    # === HEADER ===
    story.append(Paragraph("HOME CARE SERVICE AGREEMENT", styles['ContractTitle']))
    story.append(Paragraph(f"Effective Date: {date.today().strftime('%B %d, %Y')}", styles['SmallText']))
    story.append(Spacer(1, 12))
    
    # === PARTIES SECTION ===
    # Create two-column layout for parties
    parties_data = [
        [
            Paragraph("<b>SERVICE PROVIDER</b><br/>Home Care Services Agency", styles['BodyText']),
            Paragraph(f"<b>CLIENT</b><br/>{client.full_name}<br/>{client.address or ''}<br/>Phone: {client.phone or 'N/A'}", styles['BodyText'])
        ]
    ]
    parties_table = Table(parties_data, colWidths=[3.5*inch, 3.5*inch])
    parties_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
    ]))
    story.append(parties_table)
    story.append(Spacer(1, 8))
    
    # Divider line
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e2e8f0')))
    
    # === CARE ASSESSMENT ===
    story.append(Paragraph("CARE ASSESSMENT", styles['SectionHeader']))
    
    care_need_level = schedule.get("care_need_level", "MODERATE")
    client_profile = schedule.get("client_profile", {})
    
    assessment_text = f"<b>Care Need Level:</b> {care_need_level}"
    if client_profile.get("primary_diagnosis"):
        assessment_text += f"&nbsp;&nbsp;|&nbsp;&nbsp;<b>Primary Condition:</b> {client_profile['primary_diagnosis']}"
    if client_profile.get("mobility_status"):
        assessment_text += f"&nbsp;&nbsp;|&nbsp;&nbsp;<b>Mobility:</b> {client_profile['mobility_status']}"
    
    story.append(Paragraph(assessment_text, styles['BodyText']))
    
    # === SERVICES ===
    story.append(Paragraph("SERVICES TO BE PROVIDED", styles['SectionHeader']))
    
    if services:
        services_text = ""
        for i, service in enumerate(services):
            if isinstance(service, str):
                services_text += f"• {service}<br/>"
            else:
                name = service.get("name", "Service")
                desc = service.get("description", "")
                freq = service.get("frequency", "")
                line = f"• <b>{name}</b>"
                if desc:
                    line += f": {desc}"
                if freq:
                    line += f" ({freq})"
                services_text += line + "<br/>"
        story.append(Paragraph(services_text, styles['BulletItem']))
    else:
        story.append(Paragraph("• General home care services as determined by care plan", styles['BulletItem']))
    
    # === SCHEDULE & RATES (Side by side) ===
    # Get schedule info
    days = schedule.get("preferred_days", [])
    if isinstance(days, list):
        days_str = ", ".join([d if isinstance(d, str) else (d.get('day', str(d)) if isinstance(d, dict) else str(d)) for d in days]) if days else "TBD"
    else:
        days_str = str(days) if days else "TBD"
    
    hourly_rate = float(contract.hourly_rate or 0)
    weekly_hours = float(contract.weekly_hours or 0)
    weekly_cost = hourly_rate * weekly_hours
    monthly_cost = weekly_cost * 4.33
    
    schedule_rates_data = [
        [
            Paragraph("<b>SCHEDULE</b>", styles['Label']),
            Paragraph("<b>RATES</b>", styles['Label'])
        ],
        [
            Paragraph(f"Frequency: {schedule.get('frequency', 'As scheduled')}<br/>"
                     f"Days: {days_str}<br/>"
                     f"Time: {schedule.get('preferred_times', 'Flexible')}<br/>"
                     f"Hours/Week: {weekly_hours:.0f}", styles['BodyText']),
            Paragraph(f"Hourly Rate: ${hourly_rate:.2f}<br/>"
                     f"Weekly Est: ${weekly_cost:.2f}<br/>"
                     f"Monthly Est: ${monthly_cost:.2f}", styles['BodyText'])
        ]
    ]
    
    schedule_table = Table(schedule_rates_data, colWidths=[4*inch, 3*inch])
    schedule_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f7fafc')),
    ]))
    story.append(Spacer(1, 8))
    story.append(schedule_table)
    
    # === SPECIAL REQUIREMENTS & SAFETY (if any) ===
    special_reqs = schedule.get("special_requirements", [])
    safety = schedule.get("safety_concerns", [])
    
    if special_reqs or safety:
        story.append(Spacer(1, 8))
        
        reqs_safety_data = []
        
        # Build requirements text
        reqs_text = ""
        if special_reqs:
            for req in special_reqs[:4]:  # Limit to 4 items
                if isinstance(req, str):
                    reqs_text += f"• {req}<br/>"
                elif isinstance(req, dict):
                    reqs_text += f"• {req.get('name', req.get('requirement', str(req)))}<br/>"
        
        # Build safety text
        safety_text = ""
        if safety:
            for concern in safety[:4]:  # Limit to 4 items
                if isinstance(concern, str):
                    safety_text += f"• {concern}<br/>"
                elif isinstance(concern, dict):
                    text = concern.get("concern", str(concern))
                    severity = concern.get("severity", "")
                    if severity:
                        safety_text += f"• {text} [{severity}]<br/>"
                    else:
                        safety_text += f"• {text}<br/>"
        
        if reqs_text or safety_text:
            reqs_safety_data = [
                [
                    Paragraph("<b>SPECIAL REQUIREMENTS</b>", styles['Label']) if reqs_text else Paragraph("", styles['Label']),
                    Paragraph("<b>SAFETY CONSIDERATIONS</b>", styles['Label']) if safety_text else Paragraph("", styles['Label'])
                ],
                [
                    Paragraph(reqs_text or "None specified", styles['BulletItem']),
                    Paragraph(safety_text or "None noted", styles['BulletItem'])
                ]
            ]
            
            reqs_table = Table(reqs_safety_data, colWidths=[3.5*inch, 3.5*inch])
            reqs_table.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
            ]))
            story.append(reqs_table)
    
    # === POLICIES (Compact) ===
    story.append(Paragraph("TERMS & POLICIES", styles['SectionHeader']))
    
    policies_text = """<b>Cancellation:</b> 24-hour notice required. Late cancellations may incur 50% charge. 
    <b>Confidentiality:</b> All client information protected under HIPAA. 
    <b>Termination:</b> Either party may terminate with 14 days written notice. 
    <b>Payment:</b> Due within 30 days of invoice date."""
    
    story.append(Paragraph(policies_text, styles['BodyText']))
    
    # === SIGNATURES ===
    story.append(Spacer(1, 16))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e2e8f0')))
    story.append(Spacer(1, 8))
    
    sig_data = [
        [
            Paragraph("<b>CLIENT / AUTHORIZED REPRESENTATIVE</b>", styles['Label']),
            Paragraph("<b>AGENCY REPRESENTATIVE</b>", styles['Label'])
        ],
        [
            Paragraph("_" * 40 + "<br/>Signature", styles['SignatureLine']),
            Paragraph("_" * 40 + "<br/>Signature", styles['SignatureLine'])
        ],
        [
            Paragraph(f"Name: {client.full_name}", styles['BodyText']),
            Paragraph("Name: _________________________", styles['BodyText'])
        ],
        [
            Paragraph("Date: _______________", styles['BodyText']),
            Paragraph("Date: _______________", styles['BodyText'])
        ]
    ]
    
    sig_table = Table(sig_data, colWidths=[3.5*inch, 3.5*inch])
    sig_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    story.append(sig_table)
    
    # Footer
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"Contract ID: {contract.id}", styles['SmallText']))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def generate_note_pdf(visit: Any, note: Any) -> bytes:
    """
    Generate a professional PDF visit note.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.5*inch,
        leftMargin=0.5*inch,
        topMargin=0.5*inch,
        bottomMargin=0.5*inch
    )
    
    styles = get_custom_styles()
    story = []
    
    # === HEADER ===
    story.append(Paragraph("HOME CARE VISIT NOTE", styles['ContractTitle']))
    story.append(Spacer(1, 8))
    
    # === VISIT INFO ===
    visit_info = note.structured_data.get("visit_info", {}) if note.structured_data else {}
    
    info_data = [
        [
            Paragraph(f"<b>Client:</b> {visit.client.full_name if visit.client else 'N/A'}", styles['BodyText']),
            Paragraph(f"<b>Date:</b> {str(visit.actual_start or visit.scheduled_start or 'N/A')[:10]}", styles['BodyText'])
        ],
        [
            Paragraph(f"<b>Caregiver:</b> {visit.caregiver.full_name if visit.caregiver else 'N/A'}", styles['BodyText']),
            Paragraph(f"<b>Duration:</b> {visit_info.get('duration_minutes', 0)} minutes", styles['BodyText'])
        ]
    ]
    
    info_table = Table(info_data, colWidths=[3.5*inch, 3.5*inch])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f7fafc')),
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 12))
    
    # === SERVICES PROVIDED ===
    story.append(Paragraph("SERVICES PROVIDED", styles['SectionHeader']))
    
    tasks = note.structured_data.get("tasks_performed", []) if note.structured_data else []
    if tasks:
        tasks_text = ""
        for task in tasks:
            desc = task.get('description', 'Service provided')
            duration = task.get('duration_minutes', 0)
            tasks_text += f"• {desc} ({duration} min)<br/>"
        story.append(Paragraph(tasks_text, styles['BulletItem']))
    else:
        story.append(Paragraph("• Care services provided as scheduled", styles['BulletItem']))
    
    # === OBSERVATIONS ===
    story.append(Paragraph("OBSERVATIONS", styles['SectionHeader']))
    
    structured = note.structured_data or {}
    observations = structured.get("observations", "No specific observations recorded.")
    story.append(Paragraph(observations, styles['BodyText']))
    
    # === CONCERNS ===
    concerns = structured.get("risks_concerns", "")
    if concerns and concerns != "None noted.":
        story.append(Paragraph("CONCERNS / FOLLOW-UP", styles['SectionHeader']))
        story.append(Paragraph(concerns, styles['BodyText']))
    
    # === NARRATIVE ===
    if note.narrative:
        story.append(Paragraph("NARRATIVE NOTE", styles['SectionHeader']))
        story.append(Paragraph(note.narrative, styles['BodyText']))
    
    # === FOOTER ===
    story.append(Spacer(1, 20))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e2e8f0')))
    story.append(Spacer(1, 8))
    
    footer_data = [
        [
            Paragraph("_" * 35 + "<br/>Caregiver Signature", styles['SignatureLine']),
            Paragraph(f"Date: {str(note.created_at)[:10]}", styles['BodyText'])
        ]
    ]
    footer_table = Table(footer_data, colWidths=[4*inch, 3*inch])
    story.append(footer_table)
    
    story.append(Spacer(1, 8))
    story.append(Paragraph(f"Note ID: {note.id} | Version: {note.version}", styles['SmallText']))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


# Keep DOCX functions for backward compatibility
def generate_note_docx(visit: Any, note: Any) -> bytes:
    """Legacy DOCX generation - use generate_note_pdf instead."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    title = doc.add_heading("Visit Note", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading("Visit Information", level=1)
    doc.add_paragraph(f"Client: {visit.client.full_name if visit.client else 'N/A'}")
    doc.add_paragraph(f"Caregiver: {visit.caregiver.full_name if visit.caregiver else 'N/A'}")
    doc.add_paragraph(f"Date: {str(visit.actual_start or visit.scheduled_start or 'N/A')}")
    
    if note.narrative:
        doc.add_heading("Narrative", level=1)
        doc.add_paragraph(note.narrative)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_contract_docx(client: Any, contract: Any) -> bytes:
    """
    Generate a comprehensive DOCX contract using all available assessment data.
    This is used as fallback when no custom template is uploaded.
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # ==================== HEADER ====================
    title = doc.add_heading("HOME CARE SERVICE AGREEMENT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Effective Date: {date.today().strftime('%B %d, %Y')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # ==================== PARTIES ====================
    doc.add_heading("1. PARTIES", level=1)
    
    # Client info
    p = doc.add_paragraph()
    p.add_run("CLIENT:\n").bold = True
    p.add_run(f"Name: {client.full_name or 'N/A'}\n")
    p.add_run(f"Address: {client.address or 'N/A'}\n")
    if hasattr(client, 'city') and client.city:
        p.add_run(f"City/State/Zip: {client.city or ''}, {client.state or ''} {client.zip_code or ''}\n")
    p.add_run(f"Phone: {client.phone or 'N/A'}\n")
    if client.emergency_contact_name:
        p.add_run(f"Emergency Contact: {client.emergency_contact_name} ({client.emergency_contact_phone or 'N/A'})\n")
    
    doc.add_paragraph()
    
    # ==================== CARE ASSESSMENT ====================
    schedule = contract.schedule or {}
    client_profile = schedule.get('client_profile', {})
    care_need_level = schedule.get('care_need_level', 'MODERATE')
    
    doc.add_heading("2. CARE ASSESSMENT SUMMARY", level=1)
    
    p = doc.add_paragraph()
    p.add_run(f"Care Need Level: ").bold = True
    p.add_run(f"{care_need_level}\n")
    
    if client_profile.get('primary_diagnosis'):
        p.add_run(f"Primary Diagnosis: ").bold = True
        p.add_run(f"{client_profile['primary_diagnosis']}\n")
    
    if client_profile.get('secondary_conditions'):
        conditions = client_profile['secondary_conditions']
        if isinstance(conditions, list):
            conditions = ', '.join(conditions)
        p.add_run(f"Secondary Conditions: ").bold = True
        p.add_run(f"{conditions}\n")
    
    if client_profile.get('mobility_status'):
        p.add_run(f"Mobility Status: ").bold = True
        p.add_run(f"{client_profile['mobility_status']}\n")
    
    if client_profile.get('cognitive_status'):
        p.add_run(f"Cognitive Status: ").bold = True
        p.add_run(f"{client_profile['cognitive_status']}\n")
    
    if client_profile.get('living_situation'):
        p.add_run(f"Living Situation: ").bold = True
        p.add_run(f"{client_profile['living_situation']}\n")
    
    # ==================== SERVICES ====================
    doc.add_heading("3. SERVICES TO BE PROVIDED", level=1)
    
    services = contract.services or []
    if services:
        for i, service in enumerate(services, 1):
            if isinstance(service, str):
                doc.add_paragraph(f"{i}. {service}", style='List Number')
            else:
                name = service.get('name', 'Service')
                desc = service.get('description', '')
                freq = service.get('frequency', '')
                priority = service.get('priority', '')
                
                p = doc.add_paragraph(style='List Number')
                p.add_run(f"{name}").bold = True
                if desc:
                    p.add_run(f"\n   Description: {desc}")
                if freq:
                    p.add_run(f"\n   Frequency: {freq}")
                if priority:
                    p.add_run(f"\n   Priority: {priority}")
    else:
        doc.add_paragraph("Services to be determined based on care plan.")
    
    # ==================== SCHEDULE ====================
    doc.add_heading("4. SCHEDULE", level=1)
    
    p = doc.add_paragraph()
    
    frequency = schedule.get('frequency', 'As scheduled')
    p.add_run(f"Frequency: ").bold = True
    p.add_run(f"{frequency}\n")
    
    days = schedule.get('preferred_days', [])
    if isinstance(days, list):
        days_str = ', '.join([d if isinstance(d, str) else d.get('day', str(d)) for d in days]) if days else 'To be determined'
    else:
        days_str = str(days) if days else 'To be determined'
    p.add_run(f"Days: ").bold = True
    p.add_run(f"{days_str}\n")
    
    times = schedule.get('preferred_times', 'Flexible')
    p.add_run(f"Preferred Time: ").bold = True
    p.add_run(f"{times}\n")
    
    weekly_hours = float(contract.weekly_hours or 0)
    p.add_run(f"Hours per Week: ").bold = True
    p.add_run(f"{weekly_hours:.0f}\n")
    
    # ==================== RATES ====================
    doc.add_heading("5. SERVICE RATES", level=1)
    
    hourly_rate = float(contract.hourly_rate or 0)
    weekly_cost = hourly_rate * weekly_hours
    monthly_cost = weekly_cost * 4.33
    
    p = doc.add_paragraph()
    p.add_run(f"Hourly Rate: ").bold = True
    p.add_run(f"${hourly_rate:.2f}/hour\n")
    p.add_run(f"Estimated Weekly Cost: ").bold = True
    p.add_run(f"${weekly_cost:.2f}\n")
    p.add_run(f"Estimated Monthly Cost: ").bold = True
    p.add_run(f"${monthly_cost:.2f}\n")
    p.add_run(f"\nPayment Terms: ").bold = True
    p.add_run("Payment due within 7 days of invoice.\n")
    
    # ==================== SPECIAL REQUIREMENTS ====================
    special_reqs = schedule.get('special_requirements', [])
    if special_reqs:
        doc.add_heading("6. SPECIAL REQUIREMENTS", level=1)
        for req in special_reqs:
            if isinstance(req, str):
                doc.add_paragraph(f"• {req}", style='List Bullet')
            elif isinstance(req, dict):
                req_text = req.get('requirement') or req.get('name', str(req))
                doc.add_paragraph(f"• {req_text}", style='List Bullet')
    
    # ==================== SAFETY CONCERNS ====================
    safety = schedule.get('safety_concerns', [])
    if safety:
        doc.add_heading("7. SAFETY CONSIDERATIONS", level=1)
        for concern in safety:
            if isinstance(concern, str):
                doc.add_paragraph(f"• {concern}", style='List Bullet')
            elif isinstance(concern, dict):
                text = concern.get('concern', str(concern))
                severity = concern.get('severity', '')
                if severity:
                    doc.add_paragraph(f"• {text} [{severity}]", style='List Bullet')
                else:
                    doc.add_paragraph(f"• {text}", style='List Bullet')
    
    # ==================== CARE PLAN GOALS ====================
    goals = schedule.get('care_plan_goals', {})
    if goals:
        doc.add_heading("8. CARE PLAN GOALS", level=1)
        
        short_term = goals.get('short_term', [])
        if short_term:
            p = doc.add_paragraph()
            p.add_run("Short-Term Goals (30 days):").bold = True
            for goal in short_term:
                doc.add_paragraph(f"• {goal}", style='List Bullet')
        
        long_term = goals.get('long_term', [])
        if long_term:
            p = doc.add_paragraph()
            p.add_run("Long-Term Goals (90+ days):").bold = True
            for goal in long_term:
                doc.add_paragraph(f"• {goal}", style='List Bullet')
    
    # ==================== TERMS & CONDITIONS ====================
    doc.add_heading("9. TERMS AND CONDITIONS", level=1)
    
    doc.add_paragraph("""CANCELLATION POLICY: Cancellations made more than 24 hours in advance incur no charge. Cancellations within 24 hours may incur 50% of scheduled visit fee. Emergency cancellations will be evaluated on a case-by-case basis.

CONFIDENTIALITY: All client information is kept strictly confidential in accordance with HIPAA regulations. Information will only be shared with healthcare providers directly involved in the client's care or as required by law.

TERMINATION: Either party may terminate this Agreement with 14 days written notice. Immediate termination may occur if there is a safety concern or non-payment of services.

LIABILITY: The Agency maintains comprehensive general liability and professional liability insurance. Caregivers are employees of the Agency and covered under workers' compensation insurance.""")
    
    # ==================== SIGNATURES ====================
    doc.add_heading("10. SIGNATURES", level=1)
    
    doc.add_paragraph("By signing below, both parties agree to the terms and conditions of this Agreement.\n")
    
    # Client signature
    p = doc.add_paragraph()
    p.add_run("CLIENT / AUTHORIZED REPRESENTATIVE:\n\n").bold = True
    p.add_run("Signature: _________________________________  Date: _______________\n\n")
    p.add_run(f"Printed Name: {client.full_name or '________________________________'}\n")
    
    doc.add_paragraph()
    
    # Agency signature
    p = doc.add_paragraph()
    p.add_run("AGENCY REPRESENTATIVE:\n\n").bold = True
    p.add_run("Signature: _________________________________  Date: _______________\n\n")
    p.add_run("Printed Name: ________________________________\n\n")
    p.add_run("Title: ________________________________\n")
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph(f"Contract ID: {contract.id}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
