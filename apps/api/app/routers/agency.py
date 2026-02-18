"""
Agency Settings Router

Manages agency-wide settings, branding, and templates.
"""

import os
import json
import re
import logging
from typing import Optional, List, Any
from fastapi import APIRouter, Depends, HTTPException, Request, status
from sqlalchemy.orm import Session
from pydantic import BaseModel

from app.core.deps import get_db, get_current_user
from app.models.user import User
from app.models.agency_settings import AgencySettings

router = APIRouter()
logger = logging.getLogger(__name__)


class UploadedDocument(BaseModel):
    id: str
    name: str
    type: str
    category: str
    content: str
    uploaded_at: str


class AgencySettingsUpdate(BaseModel):
    name: Optional[str] = None
    address: Optional[str] = None
    city: Optional[str] = None
    state: Optional[str] = None
    zip_code: Optional[str] = None
    phone: Optional[str] = None
    email: Optional[str] = None
    website: Optional[str] = None
    logo: Optional[str] = None
    primary_color: Optional[str] = None
    secondary_color: Optional[str] = None
    documents: Optional[List[UploadedDocument]] = None
    cancellation_policy: Optional[str] = None
    terms_and_conditions: Optional[str] = None
    tax_id: Optional[str] = None
    license_number: Optional[str] = None
    npi_number: Optional[str] = None
    contact_person: Optional[str] = None
    contact_title: Optional[str] = None


class AgencySettingsResponse(BaseModel):
    id: str
    name: str
    address: Optional[str] = None
    city: Optional[str] = None
    state: Optional[str] = None
    zip_code: Optional[str] = None
    phone: Optional[str] = None
    email: Optional[str] = None
    website: Optional[str] = None
    logo: Optional[str] = None
    primary_color: Optional[str] = None
    secondary_color: Optional[str] = None
    documents: Optional[List[dict]] = None
    cancellation_policy: Optional[str] = None
    terms_and_conditions: Optional[str] = None
    tax_id: Optional[str] = None
    license_number: Optional[str] = None
    npi_number: Optional[str] = None
    contact_person: Optional[str] = None
    contact_title: Optional[str] = None


class ExtractInfoRequest(BaseModel):
    content: str  # Base64 encoded document
    document_type: str


class ExtractedInfo(BaseModel):
    name: Optional[str] = None
    address: Optional[str] = None
    city: Optional[str] = None
    state: Optional[str] = None
    zip_code: Optional[str] = None
    phone: Optional[str] = None
    email: Optional[str] = None
    website: Optional[str] = None
    tax_id: Optional[str] = None
    license_number: Optional[str] = None
    npi_number: Optional[str] = None
    contact_person: Optional[str] = None
    contact_title: Optional[str] = None
    cancellation_policy: Optional[str] = None
    terms_and_conditions: Optional[str] = None


def get_or_create_settings(db: Session, user_id) -> AgencySettings:
    """Get or create agency settings for the current user (data isolation)."""
    settings = db.query(AgencySettings).filter(
        AgencySettings.user_id == user_id
    ).first()
    
    if not settings:
        settings = AgencySettings(user_id=user_id, settings_key=f"user_{user_id}")
        db.add(settings)
        db.commit()
        db.refresh(settings)
    
    return settings


@router.get("", response_model=AgencySettingsResponse)
async def get_agency_settings(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Get agency settings (data isolation enforced)."""
    settings = get_or_create_settings(db, current_user.id)
    
    # Parse documents JSON if it exists
    documents = []
    if hasattr(settings, 'documents') and settings.documents:
        try:
            documents = json.loads(settings.documents) if isinstance(settings.documents, str) else settings.documents
        except (json.JSONDecodeError, TypeError, ValueError):
            logger.warning("Failed to parse agency documents JSON")
            documents = []
    
    return AgencySettingsResponse(
        id=str(settings.id),
        name=settings.name,
        address=settings.address,
        city=settings.city,
        state=settings.state,
        zip_code=settings.zip_code,
        phone=settings.phone,
        email=settings.email,
        website=settings.website,
        logo=settings.logo,
        primary_color=settings.primary_color,
        secondary_color=settings.secondary_color,
        documents=documents,
        cancellation_policy=settings.cancellation_policy,
        terms_and_conditions=settings.terms_and_conditions,
        tax_id=getattr(settings, 'tax_id', None),
        license_number=getattr(settings, 'license_number', None),
        npi_number=getattr(settings, 'npi_number', None),
        contact_person=getattr(settings, 'contact_person', None),
        contact_title=getattr(settings, 'contact_title', None),
    )


@router.put("", response_model=AgencySettingsResponse)
async def update_agency_settings(
    settings_update: AgencySettingsUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Update agency settings (data isolation enforced)."""
    settings = get_or_create_settings(db, current_user.id)
    
    update_data = settings_update.model_dump(exclude_unset=True)
    
    # Handle documents as JSON
    if 'documents' in update_data:
        docs = update_data['documents']
        if docs is not None:
            docs_list = [d.model_dump() if hasattr(d, 'model_dump') else d for d in docs]
            update_data['documents'] = json.dumps(docs_list)
            
            # Also update legacy contract_template fields if a contract template is present
            contract_template_doc = None
            for doc in docs_list:
                if doc.get('category') == 'contract_template':
                    contract_template_doc = doc
                    break
            
            if contract_template_doc:
                settings.contract_template = contract_template_doc.get('content')
                settings.contract_template_name = contract_template_doc.get('name')
                settings.contract_template_type = contract_template_doc.get('type')
                logger.info(f"Updated contract template: {contract_template_doc.get('name')}")
            else:
                # Clear legacy fields if no contract template in documents
                settings.contract_template = None
                settings.contract_template_name = None
                settings.contract_template_type = None
        else:
            update_data['documents'] = '[]'
    
    for field, value in update_data.items():
        if hasattr(settings, field):
            setattr(settings, field, value)
    
    db.commit()
    db.refresh(settings)
    
    # Parse documents for response
    documents = []
    if settings.documents:
        try:
            documents = json.loads(settings.documents) if isinstance(settings.documents, str) else settings.documents
        except (json.JSONDecodeError, TypeError, ValueError):
            logger.warning("Failed to parse agency documents JSON")
            documents = []
    
    return AgencySettingsResponse(
        id=str(settings.id),
        name=settings.name,
        address=settings.address,
        city=settings.city,
        state=settings.state,
        zip_code=settings.zip_code,
        phone=settings.phone,
        email=settings.email,
        website=settings.website,
        logo=settings.logo,
        primary_color=settings.primary_color,
        secondary_color=settings.secondary_color,
        documents=documents,
        cancellation_policy=settings.cancellation_policy,
        terms_and_conditions=settings.terms_and_conditions,
        tax_id=getattr(settings, 'tax_id', None),
        license_number=getattr(settings, 'license_number', None),
        npi_number=getattr(settings, 'npi_number', None),
        contact_person=getattr(settings, 'contact_person', None),
        contact_title=getattr(settings, 'contact_title', None),
    )


@router.post("/extract-info", response_model=ExtractedInfo)
async def extract_company_info(
    request: ExtractInfoRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Extract company information from an uploaded document using AI.
    Supports letterheads, policy documents, contracts, etc.
    """
    openai_key = os.getenv("OPENAI_API_KEY")
    anthropic_key = os.getenv("ANTHROPIC_API_KEY")
    
    # Try OpenAI first (more reliable), then fall back to Anthropic
    if openai_key:
        result = await _extract_with_openai(request.content, request.document_type, openai_key)
        if result.name or result.address or result.phone or result.email:
            return result
    
    if anthropic_key:
        result = await _extract_with_claude(request.content, request.document_type, anthropic_key)
        if result.name or result.address or result.phone or result.email:
            return result
    
    if not openai_key and not anthropic_key:
        logger.warning("No LLM API key found for document extraction")
    
    return ExtractedInfo()


async def _extract_with_claude(content: str, doc_type: str, api_key: str) -> ExtractedInfo:
    """Extract company info using Claude."""
    try:
        import anthropic
        
        client = anthropic.Anthropic(api_key=api_key)
        
        # Prepare prompt based on document type
        system_prompt = """You are an expert at extracting business information from documents.
        
Extract the following information if present:
- Company/Agency name
- Street address
- City, State, ZIP code
- Phone number
- Email address
- Website
- Tax ID / EIN
- License number
- NPI number (for healthcare)
- Contact person name
- Contact person title
- Cancellation policy (if this is a policy document)
- Terms and conditions (if present)

Return ONLY a JSON object with these fields (use null for missing fields):
{
    "name": "Company Name",
    "address": "123 Main St",
    "city": "City",
    "state": "State",
    "zip_code": "12345",
    "phone": "(555) 123-4567",
    "email": "email@company.com",
    "website": "https://www.company.com",
    "tax_id": "XX-XXXXXXX",
    "license_number": "ABC123",
    "npi_number": "1234567890",
    "contact_person": "John Smith",
    "contact_title": "Administrator",
    "cancellation_policy": "...",
    "terms_and_conditions": "..."
}"""

        # For base64 content, we'll describe it
        user_prompt = f"""This is a {doc_type} document (base64 encoded). 
        
Please analyze the document content and extract any company/business information you can find.
Look for:
- Letterhead information
- Contact details
- Business identifiers
- Policy text (if applicable)

Document content (first 5000 chars of base64): {content[:5000]}

Extract and return the JSON object with company information."""

        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2000,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}],
            temperature=0,
        )
        
        # Parse response
        response_text = response.content[0].text
        
        # Try to extract JSON from response
        json_match = re.search(r'\{[^{}]*\}', response_text, re.DOTALL)
        if json_match:
            data = json.loads(json_match.group())
            return ExtractedInfo(**{k: v for k, v in data.items() if v is not None})
        
        return ExtractedInfo()
        
    except Exception as e:
        logger.error(f"Claude extraction failed: {e}")
        return ExtractedInfo()


def _extract_text_from_document(content: str) -> tuple[str, str]:
    """
    Extract text from various document formats.
    Returns (text_content, mime_type)
    """
    import base64
    import io
    
    text_content = ""
    mime_type = "text/plain"
    
    # Check if it's a data URL
    if content.startswith("data:"):
        # Parse data URL: data:mime/type;base64,<data>
        try:
            header, b64_data = content.split(",", 1)
            mime_type = header.split(":")[1].split(";")[0] if ":" in header else "application/octet-stream"
            raw_bytes = base64.b64decode(b64_data)
        except Exception as e:
            logger.error(f"Failed to decode base64: {e}")
            return content[:3000], mime_type
    else:
        # Assume it's plain text or raw base64
        try:
            raw_bytes = base64.b64decode(content)
        except (Exception,):
            return content[:3000], "text/plain"
    
    # Handle different mime types
    try:
        if "pdf" in mime_type.lower():
            # Try to extract text from PDF
            try:
                import pypdf
                pdf_reader = pypdf.PdfReader(io.BytesIO(raw_bytes))
                text_parts = []
                for page in pdf_reader.pages[:5]:  # First 5 pages
                    text_parts.append(page.extract_text() or "")
                text_content = "\n".join(text_parts)
            except ImportError:
                # pypdf not installed, try raw decode
                text_content = raw_bytes.decode('utf-8', errors='ignore')[:3000]
            except Exception as e:
                logger.error(f"PDF extraction failed: {e}")
                text_content = raw_bytes.decode('utf-8', errors='ignore')[:3000]
                
        elif "word" in mime_type.lower() or "docx" in mime_type.lower() or "msword" in mime_type.lower():
            # Try to extract text from DOCX
            try:
                from docx import Document
                doc = Document(io.BytesIO(raw_bytes))
                text_parts = []
                for para in doc.paragraphs[:50]:  # First 50 paragraphs
                    if para.text.strip():
                        text_parts.append(para.text)
                text_content = "\n".join(text_parts)
            except ImportError:
                text_content = raw_bytes.decode('utf-8', errors='ignore')[:3000]
            except Exception as e:
                logger.error(f"DOCX extraction failed: {e}")
                text_content = raw_bytes.decode('utf-8', errors='ignore')[:3000]
                
        elif "image" in mime_type.lower():
            # For images, we'll use GPT-4 Vision - return the original content
            return content, mime_type
            
        else:
            # Try to decode as text
            text_content = raw_bytes.decode('utf-8', errors='ignore')[:4000]
            
    except Exception as e:
        logger.error(f"Text extraction failed: {e}")
        text_content = str(raw_bytes[:2000])
    
    return text_content, mime_type


async def _extract_with_openai(content: str, doc_type: str, api_key: str) -> ExtractedInfo:
    """Extract company info using OpenAI."""
    try:
        from openai import OpenAI
        
        client = OpenAI(api_key=api_key)
        
        # Extract text from document
        text_content, mime_type = _extract_text_from_document(content)
        
        logger.info(f"Extracted text ({len(text_content)} chars) from {mime_type}")
        
        # For images, use GPT-4 Vision
        if "image" in mime_type.lower() and content.startswith("data:"):
            return await _extract_with_openai_vision(content, doc_type, api_key)
        
        system_prompt = """You are an expert at extracting business information from documents.
        
Extract ALL of the following information if present in the document:
- name: Company or agency name
- address: Street address
- city: City name
- state: State (abbreviation or full name)
- zip_code: ZIP or postal code
- phone: Phone number
- email: Email address  
- website: Website URL
- tax_id: Tax ID or EIN (format: XX-XXXXXXX)
- license_number: Business or healthcare license number
- npi_number: National Provider Identifier (10 digits)
- contact_person: Primary contact name
- contact_title: Contact's job title
- cancellation_policy: Any cancellation policy text
- terms_and_conditions: Any terms and conditions text

Return a JSON object with these exact field names. Use null for any field not found."""

        user_content = f"""Extract company/business information from this {doc_type} document:

{text_content[:4000]}

Return ONLY a valid JSON object with the extracted information."""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content}
            ],
            temperature=0,
            response_format={"type": "json_object"},
        )
        
        data = json.loads(response.choices[0].message.content)
        logger.info(f"OpenAI extracted: {data}")
        return ExtractedInfo(**{k: v for k, v in data.items() if v is not None and k in ExtractedInfo.model_fields})
        
    except Exception as e:
        logger.error(f"OpenAI extraction failed: {e}")
        return ExtractedInfo()


async def _extract_with_openai_vision(content: str, doc_type: str, api_key: str) -> ExtractedInfo:
    """Extract company info from images using OpenAI Vision."""
    try:
        from openai import OpenAI
        
        client = OpenAI(api_key=api_key)
        
        system_prompt = """You are an expert at extracting business information from document images.
        
Look at the image and extract ALL of the following information if visible:
- name: Company or agency name
- address: Street address
- city: City name
- state: State (abbreviation or full name)
- zip_code: ZIP or postal code
- phone: Phone number
- email: Email address  
- website: Website URL
- tax_id: Tax ID or EIN
- license_number: Business or healthcare license number
- npi_number: National Provider Identifier (10 digits)
- contact_person: Primary contact name
- contact_title: Contact's job title

Return ONLY a JSON object with these field names. Use null for any field not found."""

        response = client.chat.completions.create(
            model="gpt-4o",  # Use GPT-4o for vision
            messages=[
                {"role": "system", "content": system_prompt},
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": f"Extract company information from this {doc_type} image:"},
                        {"type": "image_url", "image_url": {"url": content}}
                    ]
                }
            ],
            temperature=0,
            max_tokens=1500,
        )
        
        response_text = response.choices[0].message.content
        
        # Try to extract JSON from response
        json_match = re.search(r'\{[^{}]*\}', response_text, re.DOTALL)
        if json_match:
            data = json.loads(json_match.group())
            logger.info(f"OpenAI Vision extracted: {data}")
            return ExtractedInfo(**{k: v for k, v in data.items() if v is not None and k in ExtractedInfo.model_fields})
        
        return ExtractedInfo()
        
    except Exception as e:
        logger.error(f"OpenAI Vision extraction failed: {e}")
        return ExtractedInfo()


# Internal endpoint for contract generation worker — protected by API key
@router.get("/public", response_model=AgencySettingsResponse)
async def get_public_agency_settings(
    request: Request,
    db: Session = Depends(get_db),
):
    """Get agency settings (internal - for worker access). Requires X-Internal-Key header."""
    expected_key = os.getenv("INTERNAL_API_KEY", "")
    provided_key = request.headers.get("X-Internal-Key", "")
    if not expected_key or provided_key != expected_key:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Invalid or missing internal API key")
    
    # For internal worker access, get the first available settings
    settings = db.query(AgencySettings).first()
    if not settings:
        raise HTTPException(status_code=404, detail="No agency settings found")
    
    documents = []
    if hasattr(settings, 'documents') and settings.documents:
        try:
            documents = json.loads(settings.documents) if isinstance(settings.documents, str) else settings.documents
        except (json.JSONDecodeError, TypeError, ValueError):
            logger.warning("Failed to parse agency documents JSON")
            documents = []
    
    return AgencySettingsResponse(
        id=str(settings.id),
        name=settings.name,
        address=settings.address,
        city=settings.city,
        state=settings.state,
        zip_code=settings.zip_code,
        phone=settings.phone,
        email=settings.email,
        website=settings.website,
        logo=settings.logo,
        primary_color=settings.primary_color,
        secondary_color=settings.secondary_color,
        documents=documents,
        cancellation_policy=settings.cancellation_policy,
        terms_and_conditions=settings.terms_and_conditions,
        tax_id=getattr(settings, 'tax_id', None),
        license_number=getattr(settings, 'license_number', None),
        npi_number=getattr(settings, 'npi_number', None),
        contact_person=getattr(settings, 'contact_person', None),
        contact_title=getattr(settings, 'contact_title', None),
    )
