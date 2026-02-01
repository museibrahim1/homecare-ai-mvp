"""
Template Parser Router

Extracts agency information from uploaded contract templates using Claude.
"""

import base64
import logging
from typing import Optional
from fastapi import APIRouter, Depends, HTTPException
from pydantic import BaseModel
from sqlalchemy.orm import Session

from app.core.deps import get_db, get_current_user
from app.models.user import User

logger = logging.getLogger(__name__)

router = APIRouter()


class TemplateParseRequest(BaseModel):
    template_data: str  # Base64 encoded file
    template_name: str
    template_type: str  # mime type


class ExtractedAgencyInfo(BaseModel):
    name: Optional[str] = None
    address: Optional[str] = None
    city: Optional[str] = None
    state: Optional[str] = None
    zip_code: Optional[str] = None
    phone: Optional[str] = None
    email: Optional[str] = None
    website: Optional[str] = None
    logo_found: bool = False


@router.post("/extract-agency-info", response_model=ExtractedAgencyInfo)
async def extract_agency_info(
    request: TemplateParseRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Extract agency information from an uploaded template using Claude.
    """
    import os
    
    try:
        # Decode base64 to get file content
        # Remove data URL prefix if present
        template_data = request.template_data
        if ',' in template_data:
            template_data = template_data.split(',')[1]
        
        file_bytes = base64.b64decode(template_data)
        
        # Extract text from document
        text_content = ""
        
        if 'pdf' in request.template_type.lower():
            # Try to extract text from PDF
            try:
                import io
                # Use a simple PDF text extraction
                text_content = extract_text_from_pdf(file_bytes)
            except Exception as e:
                logger.warning(f"PDF extraction failed: {e}")
                text_content = f"[PDF document: {request.template_name}]"
                
        elif 'word' in request.template_type.lower() or 'docx' in request.template_type.lower():
            # Extract text from DOCX
            try:
                text_content = extract_text_from_docx(file_bytes)
            except Exception as e:
                logger.warning(f"DOCX extraction failed: {e}")
                text_content = f"[Word document: {request.template_name}]"
        else:
            # Try to decode as text
            try:
                text_content = file_bytes.decode('utf-8')
            except:
                text_content = f"[Document: {request.template_name}]"
        
        # Use Claude to extract agency information
        extracted = await extract_with_llm(text_content, request.template_name)
        
        return extracted
        
    except Exception as e:
        logger.error(f"Failed to extract agency info: {e}")
        raise HTTPException(status_code=500, detail=str(e))


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes."""
    try:
        from reportlab.lib.pagesizes import letter
        import io
        
        # Try PyPDF2 or pdfplumber if available
        try:
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
            return text
        except ImportError:
            pass
        
        # Fallback - just indicate it's a PDF
        return "[PDF content - text extraction unavailable]"
        
    except Exception as e:
        logger.error(f"PDF text extraction error: {e}")
        return "[PDF content]"


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes."""
    try:
        from docx import Document
        import io
        
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        
        # Extract from headers first (important for letterhead)
        for section in doc.sections:
            try:
                header = section.header
                for para in header.paragraphs:
                    if para.text.strip():
                        text_parts.append(f"HEADER: {para.text.strip()}")
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_parts.append(f"HEADER: {cell.text.strip()}")
            except:
                pass
        
        # Extract from body paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        
        return "\n".join(text_parts)
        
    except Exception as e:
        logger.error(f"DOCX text extraction error: {e}")
        return "[Word document content]"


async def extract_with_llm(text_content: str, filename: str) -> ExtractedAgencyInfo:
    """Use Claude to extract agency information from document text."""
    import os
    import json
    
    # Try Anthropic first
    anthropic_key = os.getenv("ANTHROPIC_API_KEY")
    openai_key = os.getenv("OPENAI_API_KEY")
    
    system_prompt = """You are an expert at extracting business/agency information from documents.
    
Extract the HOME CARE AGENCY information from this contract template. 

CRITICAL: The agency info is usually at the TOP of the document - the FIRST 4-5 LINES typically contain:
- Line 1: Company/Agency Name (e.g., "Patron Senior Living LLC", "ABC Home Care")
- Line 2: Street Address (e.g., "1402 Jones Street Suite 211")
- Line 3: City, State ZIP (e.g., "Omaha, NE 68102")
- Line 4: Phone Number (e.g., "402-800-7759")

Look for:
- Agency/Company Name (the business providing care services) - FIRST LINE
- Street Address - SECOND LINE
- City, State, ZIP - THIRD LINE (parse city/state/zip separately)
- Phone Number - usually has format XXX-XXX-XXXX
- Email Address (if present)
- Website (if present)

IMPORTANT: 
- Extract the AGENCY information, not client/patient information (ignore "Person to Receive Services" section)
- The first few lines are almost always the agency letterhead
- Parse "Omaha, NE 68102" as city="Omaha", state="NE", zip_code="68102"

Return a JSON object with these exact keys:
{
    "name": "Agency Name or null",
    "address": "Street Address or null",
    "city": "City or null", 
    "state": "State (2 letter code) or null",
    "zip_code": "ZIP or null",
    "phone": "Phone or null",
    "email": "Email or null",
    "website": "Website or null",
    "logo_found": true/false
}"""

    user_prompt = f"""Document filename: {filename}

Document content:
{text_content[:8000]}

Extract the agency/company information and return as JSON."""

    try:
        if anthropic_key:
            import httpx
            
            async with httpx.AsyncClient() as client:
                response = await client.post(
                    "https://api.anthropic.com/v1/messages",
                    headers={
                        "x-api-key": anthropic_key,
                        "anthropic-version": "2023-06-01",
                        "content-type": "application/json",
                    },
                    json={
                        "model": "claude-3-haiku-20240307",
                        "max_tokens": 1024,
                        "temperature": 0,
                        "system": system_prompt,
                        "messages": [{"role": "user", "content": user_prompt}]
                    },
                    timeout=30.0
                )
                
                if response.status_code == 200:
                    result = response.json()
                    content = result.get("content", [{}])[0].get("text", "{}")
                    
                    # Parse JSON from response
                    try:
                        # Find JSON in response
                        if "{" in content and "}" in content:
                            json_start = content.index("{")
                            json_end = content.rindex("}") + 1
                            json_str = content[json_start:json_end]
                            data = json.loads(json_str)
                            
                            return ExtractedAgencyInfo(
                                name=data.get("name") if data.get("name") != "null" else None,
                                address=data.get("address") if data.get("address") != "null" else None,
                                city=data.get("city") if data.get("city") != "null" else None,
                                state=data.get("state") if data.get("state") != "null" else None,
                                zip_code=data.get("zip_code") if data.get("zip_code") != "null" else None,
                                phone=data.get("phone") if data.get("phone") != "null" else None,
                                email=data.get("email") if data.get("email") != "null" else None,
                                website=data.get("website") if data.get("website") != "null" else None,
                                logo_found=data.get("logo_found", False),
                            )
                    except json.JSONDecodeError:
                        logger.warning("Failed to parse LLM JSON response")
        
        # Fallback to OpenAI
        if openai_key:
            import httpx
            
            async with httpx.AsyncClient() as client:
                response = await client.post(
                    "https://api.openai.com/v1/chat/completions",
                    headers={
                        "Authorization": f"Bearer {openai_key}",
                        "Content-Type": "application/json",
                    },
                    json={
                        "model": "gpt-3.5-turbo",
                        "temperature": 0,
                        "messages": [
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_prompt}
                        ]
                    },
                    timeout=30.0
                )
                
                if response.status_code == 200:
                    result = response.json()
                    content = result.get("choices", [{}])[0].get("message", {}).get("content", "{}")
                    
                    try:
                        if "{" in content and "}" in content:
                            json_start = content.index("{")
                            json_end = content.rindex("}") + 1
                            json_str = content[json_start:json_end]
                            data = json.loads(json_str)
                            
                            return ExtractedAgencyInfo(
                                name=data.get("name") if data.get("name") != "null" else None,
                                address=data.get("address") if data.get("address") != "null" else None,
                                city=data.get("city") if data.get("city") != "null" else None,
                                state=data.get("state") if data.get("state") != "null" else None,
                                zip_code=data.get("zip_code") if data.get("zip_code") != "null" else None,
                                phone=data.get("phone") if data.get("phone") != "null" else None,
                                email=data.get("email") if data.get("email") != "null" else None,
                                website=data.get("website") if data.get("website") != "null" else None,
                                logo_found=data.get("logo_found", False),
                            )
                    except json.JSONDecodeError:
                        logger.warning("Failed to parse OpenAI JSON response")
                        
    except Exception as e:
        logger.error(f"LLM extraction failed: {e}")
    
    # Return empty if all else fails
    return ExtractedAgencyInfo()


@router.get("/placeholders")
async def get_template_placeholders_docs():
    """
    Return documentation of available template placeholders.
    """
    return {
        "description": "Use these placeholders in your DOCX template. Supported formats: {key}, {{key}}, [key], [[key]]",
        "categories": {
            "dates": {
                "description": "Date placeholders",
                "placeholders": {
                    "date": "Today's date (e.g., January 29, 2026)",
                    "contract_date": "Contract date",
                    "effective_date": "Effective date",
                    "today": "Today's date",
                    "current_date": "Current date",
                }
            },
            "agency": {
                "description": "Agency/Provider information",
                "placeholders": {
                    "agency_name": "Agency name",
                    "agency": "Agency name (alias)",
                    "company_name": "Company name",
                    "provider_name": "Provider name",
                    "agency_address": "Full agency address",
                    "agency_street": "Street address",
                    "agency_city": "City",
                    "agency_state": "State",
                    "agency_zip": "ZIP code",
                    "agency_phone": "Phone number",
                    "agency_email": "Email address",
                }
            },
            "client": {
                "description": "Client/Patient information",
                "placeholders": {
                    "client_name": "Client's full name",
                    "client": "Client's name (alias)",
                    "patient_name": "Patient name (alias)",
                    "client_first_name": "First name",
                    "client_last_name": "Last name",
                    "client_address": "Full address",
                    "client_city": "City",
                    "client_state": "State",
                    "client_zip": "ZIP code",
                    "client_phone": "Phone number",
                    "client_email": "Email address",
                    "date_of_birth": "Date of birth",
                    "dob": "Date of birth (alias)",
                    "emergency_contact": "Emergency contact name",
                    "emergency_phone": "Emergency contact phone",
                }
            },
            "care_assessment": {
                "description": "Care assessment data",
                "placeholders": {
                    "care_level": "Care need level",
                    "primary_diagnosis": "Primary diagnosis",
                    "mobility_status": "Mobility status",
                    "cognitive_status": "Cognitive status",
                    "living_situation": "Living situation",
                }
            },
            "services": {
                "description": "Service information",
                "placeholders": {
                    "services": "Formatted list of services",
                    "services_list": "Formatted list of services (alias)",
                }
            },
            "schedule": {
                "description": "Schedule information",
                "placeholders": {
                    "schedule_days": "Days of service",
                    "days": "Days (alias)",
                    "preferred_days": "Preferred days",
                    "schedule_time": "Preferred time",
                    "time": "Time (alias)",
                    "frequency": "Service frequency",
                    "weekly_hours": "Hours per week",
                }
            },
            "rates": {
                "description": "Billing rates",
                "placeholders": {
                    "hourly_rate": "Hourly rate with $ (e.g., $35.00)",
                    "hourly_rate_value": "Hourly rate number only (e.g., 35.00)",
                    "rate": "Hourly rate with $",
                    "weekly_cost": "Weekly cost estimate",
                    "monthly_cost": "Monthly cost estimate",
                    "weekday_rate": "Weekday rate (number)",
                    "weekend_rate": "Weekend rate (+25%)",
                    "holiday_rate": "Holiday rate (+50%)",
                }
            },
            "requirements": {
                "description": "Special requirements and safety",
                "placeholders": {
                    "special_requirements": "Special requirements list",
                    "requirements": "Requirements (alias)",
                    "safety_concerns": "Safety concerns list",
                    "safety": "Safety (alias)",
                }
            },
            "contract": {
                "description": "Contract metadata",
                "placeholders": {
                    "contract_id": "Contract UUID",
                }
            },
        },
        "label_auto_fill": {
            "description": "The system also auto-fills fields that follow a 'Label:' pattern",
            "examples": [
                "Name: ___________ → Auto-filled with client name",
                "Address: ________ → Auto-filled with client address",
                "Phone: __________ → Auto-filled with client phone",
                "Hourly Rate: ____ → Auto-filled with rate",
            ]
        },
        "tips": [
            "Use simple formatting - complex tables may not fill correctly",
            "Test with a sample client to verify placeholders work",
            "Use DOCX format - DOC, PDF, and other formats are not supported",
            "All placeholders are case-insensitive",
        ]
    }
