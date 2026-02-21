"""
Contract Templates Router

Upload, OCR-scan, and manage versioned contract templates.
Handles field detection, mapping reconciliation, template versioning,
and a gallery of pre-made starter templates users can clone.
"""

import base64
import hashlib
import json
import logging
from pathlib import Path
from typing import List, Optional
from uuid import UUID

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, status
from pydantic import BaseModel
from sqlalchemy.orm import Session

from app.core.deps import get_db, get_current_user
from app.models.user import User
from app.models.contract_template import ContractTemplate
from app.services.ocr_template_scanner import (
    extract_text,
    detect_fields_with_ai,
    build_field_mapping,
    reconcile_versions,
    compute_file_hash,
    DB_FIELD_REGISTRY,
)

logger = logging.getLogger(__name__)

router = APIRouter()


# ---------- Schemas ----------

class FieldInfo(BaseModel):
    field_id: str
    label: str = ""
    type: str = "text"
    required: bool = False
    section: str = ""
    mapped_to: Optional[str] = None
    is_filled: bool = False


class TemplateResponse(BaseModel):
    id: UUID
    name: str
    version: int
    description: Optional[str]
    is_active: bool
    file_type: str
    detected_fields: list
    field_mapping: dict
    unmapped_fields: list
    created_at: str
    updated_at: str

    class Config:
        from_attributes = True


class TemplateListItem(BaseModel):
    id: UUID
    name: str
    version: int
    is_active: bool
    file_type: str
    field_count: int
    unmapped_count: int
    created_at: str

    class Config:
        from_attributes = True


class FieldMappingUpdate(BaseModel):
    field_id: str
    mapped_to: str


class ReconciliationReport(BaseModel):
    added_fields: list
    removed_fields: list
    unchanged_fields: list
    total_old: int
    total_new: int
    summary: str


# ---------- Endpoints ----------

@router.post("/upload", response_model=TemplateResponse)
async def upload_and_scan_template(
    file: UploadFile = File(...),
    name: str = Form(...),
    description: Optional[str] = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Upload a contract template (PDF or DOCX), run OCR to extract text,
    detect form fields via AI, and build the field mapping.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")

    allowed_types = [
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ]
    content_type = file.content_type or ""
    is_pdf = file.filename.lower().endswith(".pdf") or "pdf" in content_type
    is_docx = file.filename.lower().endswith((".docx", ".doc")) or "word" in content_type

    if not (is_pdf or is_docx):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")

    file_bytes = await file.read()
    if len(file_bytes) > 20 * 1024 * 1024:  # 20MB limit
        raise HTTPException(status_code=400, detail="File too large (max 20MB)")

    file_hash = compute_file_hash(file_bytes)

    # Check for duplicate upload
    existing = db.query(ContractTemplate).filter(
        ContractTemplate.owner_id == current_user.id,
        ContractTemplate.file_hash == file_hash,
    ).first()
    if existing:
        raise HTTPException(
            status_code=409,
            detail=f"This file was already uploaded as template '{existing.name}' (v{existing.version})"
        )

    # Determine next version for this template name
    latest = db.query(ContractTemplate).filter(
        ContractTemplate.owner_id == current_user.id,
        ContractTemplate.name == name,
    ).order_by(ContractTemplate.version.desc()).first()
    next_version = (latest.version + 1) if latest else 1

    # Deactivate previous versions
    if latest:
        db.query(ContractTemplate).filter(
            ContractTemplate.owner_id == current_user.id,
            ContractTemplate.name == name,
            ContractTemplate.is_active == True,
        ).update({"is_active": False})

    # Step 1: Extract text via OCR
    ocr_text = await extract_text(file_bytes, file.filename, content_type)

    # Step 2: Detect fields with AI
    detected_fields = await detect_fields_with_ai(ocr_text, file.filename)

    # Step 3: Build field mapping
    field_mapping, unmapped_fields = build_field_mapping(detected_fields)

    # Store template with base64 file data
    file_url = f"data:{content_type};base64,{base64.b64encode(file_bytes).decode()}"

    template = ContractTemplate(
        owner_id=current_user.id,
        name=name,
        version=next_version,
        description=description,
        is_active=True,
        file_type="pdf" if is_pdf else "docx",
        file_url=file_url,
        file_hash=file_hash,
        ocr_text=ocr_text,
        detected_fields=detected_fields,
        field_mapping=field_mapping,
        unmapped_fields=unmapped_fields,
    )

    db.add(template)
    db.commit()
    db.refresh(template)

    return _template_to_response(template)


@router.get("/", response_model=List[TemplateListItem])
async def list_templates(
    active_only: bool = True,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """List all contract templates for the current user."""
    query = db.query(ContractTemplate).filter(
        ContractTemplate.owner_id == current_user.id,
    )
    if active_only:
        query = query.filter(ContractTemplate.is_active == True)

    templates = query.order_by(ContractTemplate.created_at.desc()).all()

    return [
        TemplateListItem(
            id=t.id,
            name=t.name,
            version=t.version,
            is_active=t.is_active,
            file_type=t.file_type,
            field_count=len(t.detected_fields or []),
            unmapped_count=len(t.unmapped_fields or []),
            created_at=str(t.created_at),
        )
        for t in templates
    ]


@router.get("/{template_id}", response_model=TemplateResponse)
async def get_template(
    template_id: UUID,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Get full template details including fields and mapping."""
    template = db.query(ContractTemplate).filter(
        ContractTemplate.id == template_id,
        ContractTemplate.owner_id == current_user.id,
    ).first()

    if not template:
        raise HTTPException(status_code=404, detail="Template not found")

    return _template_to_response(template)


@router.delete("/{template_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_template(
    template_id: UUID,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Delete a template."""
    template = db.query(ContractTemplate).filter(
        ContractTemplate.id == template_id,
        ContractTemplate.owner_id == current_user.id,
    ).first()

    if not template:
        raise HTTPException(status_code=404, detail="Template not found")

    db.delete(template)
    db.commit()


@router.post("/{template_id}/rescan")
async def rescan_template(
    template_id: UUID,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Re-run OCR and field detection on an existing template."""
    template = db.query(ContractTemplate).filter(
        ContractTemplate.id == template_id,
        ContractTemplate.owner_id == current_user.id,
    ).first()

    if not template:
        raise HTTPException(status_code=404, detail="Template not found")

    if not template.file_url:
        raise HTTPException(status_code=400, detail="No file stored for this template")

    # Decode stored file
    if template.file_url.startswith("data:"):
        _, encoded = template.file_url.split(",", 1)
        file_bytes = base64.b64decode(encoded)
    else:
        raise HTTPException(status_code=400, detail="Cannot rescan: file not stored inline")

    content_type = "application/pdf" if template.file_type == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    filename = f"{template.name}.{template.file_type}"

    ocr_text = await extract_text(file_bytes, filename, content_type)
    detected_fields = await detect_fields_with_ai(ocr_text, filename)
    field_mapping, unmapped_fields = build_field_mapping(detected_fields)

    template.ocr_text = ocr_text
    template.detected_fields = detected_fields
    template.field_mapping = field_mapping
    template.unmapped_fields = unmapped_fields

    db.commit()
    db.refresh(template)

    return _template_to_response(template)


@router.put("/{template_id}/mapping")
async def update_field_mapping(
    template_id: UUID,
    updates: List[FieldMappingUpdate],
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Manually update field mappings — map unmapped template fields to DB paths.
    This handles the case where a new template version has fields the DB doesn't
    recognize yet.
    """
    template = db.query(ContractTemplate).filter(
        ContractTemplate.id == template_id,
        ContractTemplate.owner_id == current_user.id,
    ).first()

    if not template:
        raise HTTPException(status_code=404, detail="Template not found")

    mapping = dict(template.field_mapping or {})
    remaining_unmapped = list(template.unmapped_fields or [])

    for update in updates:
        mapping[update.field_id] = update.mapped_to
        remaining_unmapped = [
            f for f in remaining_unmapped if f.get("field_id") != update.field_id
        ]

    template.field_mapping = mapping
    template.unmapped_fields = remaining_unmapped

    db.commit()
    db.refresh(template)

    return _template_to_response(template)


@router.post("/{template_id}/reconcile/{old_template_id}", response_model=ReconciliationReport)
async def reconcile_template_versions(
    template_id: UUID,
    old_template_id: UUID,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Compare two template versions and report field differences.
    Shows what fields were added, removed, or unchanged between versions.
    """
    new_template = db.query(ContractTemplate).filter(
        ContractTemplate.id == template_id,
        ContractTemplate.owner_id == current_user.id,
    ).first()

    old_template = db.query(ContractTemplate).filter(
        ContractTemplate.id == old_template_id,
        ContractTemplate.owner_id == current_user.id,
    ).first()

    if not new_template or not old_template:
        raise HTTPException(status_code=404, detail="Template not found")

    report = reconcile_versions(
        old_template.detected_fields or [],
        new_template.detected_fields or [],
    )

    return ReconciliationReport(**report)


@router.get("/registry/fields")
async def get_field_registry():
    """
    Return all known database fields that templates can map to.
    Used by the frontend to let users manually map unmapped fields.
    """
    return {
        "fields": [
            {
                "field_id": field_id,
                "path": info["path"],
                "type": info["type"],
                "category": info["category"],
            }
            for field_id, info in DB_FIELD_REGISTRY.items()
        ],
        "categories": sorted(set(info["category"] for info in DB_FIELD_REGISTRY.values())),
    }


@router.get("/preview/{contract_id}")
async def preview_template_with_data(
    contract_id: UUID,
    template_id: UUID = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Return the active OCR template's detected fields populated with
    actual contract + client + agency data.  The frontend renders this
    as a live preview so the user can see exactly what the exported
    document will contain.
    """
    from app.models.contract import Contract
    from app.models.client import Client
    from app.models.agency_settings import AgencySettings
    from app.services.document_generation import get_template_placeholders

    contract = db.query(Contract).join(Client, Contract.client_id == Client.id).filter(
        Contract.id == contract_id,
        Client.created_by == current_user.id,
    ).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    client = db.query(Client).filter(Client.id == contract.client_id).first()
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")

    if template_id:
        template = db.query(ContractTemplate).filter(
            ContractTemplate.id == template_id,
            ContractTemplate.owner_id == current_user.id,
        ).first()
    else:
        template = db.query(ContractTemplate).filter(
            ContractTemplate.owner_id == current_user.id,
            ContractTemplate.is_active == True,
        ).order_by(ContractTemplate.version.desc()).first()

    if not template:
        return {"has_template": False, "fields": [], "template_name": None}

    agency_settings = db.query(AgencySettings).filter(
        AgencySettings.user_id == current_user.id,
    ).first()

    placeholders = get_template_placeholders(client, contract, agency_settings)

    filled_fields = []
    for field in (template.detected_fields or []):
        fid = field.get("field_id", "")
        value = placeholders.get(fid, "")
        if not value and fid in (template.field_mapping or {}):
            mapped_path = template.field_mapping[fid]
            short_key = mapped_path.rsplit(".", 1)[-1] if "." in mapped_path else mapped_path
            value = placeholders.get(short_key, "")

        filled_fields.append({
            "field_id": fid,
            "label": field.get("label", fid),
            "section": field.get("section", ""),
            "type": field.get("type", "text"),
            "required": field.get("required", False),
            "value": str(value) if value else "",
            "is_mapped": fid not in [u.get("field_id") for u in (template.unmapped_fields or [])],
        })

    return {
        "has_template": True,
        "template_name": template.name,
        "template_version": template.version,
        "file_type": template.file_type,
        "fields": filled_fields,
    }


# ---------- Gallery: Pre-made starter templates ----------

# Gallery templates live inside the API package at apps/api/templates/contracts/
# In Docker (WORKDIR /app) they land at /app/templates/contracts/
_API_ROOT = Path(__file__).resolve().parent.parent.parent  # apps/api
_WORKDIR_ROOT = Path("/app")

if (_API_ROOT / "templates" / "contracts").is_dir():
    GALLERY_DIR = _API_ROOT / "templates" / "contracts"
elif (_WORKDIR_ROOT / "templates" / "contracts").is_dir():
    GALLERY_DIR = _WORKDIR_ROOT / "templates" / "contracts"
else:
    GALLERY_DIR = _API_ROOT / "templates" / "contracts"


class GalleryItem(BaseModel):
    slug: str
    name: str
    description: str
    file_type: str
    field_count: int
    mapped_count: int
    unmapped_count: int
    sections: List[str]


@router.get("/gallery/list", response_model=List[GalleryItem])
async def list_gallery_templates():
    """
    List all pre-made starter templates available in the gallery.
    These are shipped with the app — users can clone one into their account.
    No authentication required to browse.
    """
    items = []
    if not GALLERY_DIR.is_dir():
        return items
    for meta_path in sorted(GALLERY_DIR.glob("*_meta.json")):
        try:
            with open(meta_path) as f:
                meta = json.load(f)

            sections = sorted(set(
                f.get("section", "unknown") for f in meta.get("detected_fields", [])
            ))

            items.append(GalleryItem(
                slug=meta_path.stem.replace("_meta", ""),
                name=meta.get("name", meta_path.stem),
                description=meta.get("description", ""),
                file_type=meta.get("file_type", "docx"),
                field_count=len(meta.get("detected_fields", [])),
                mapped_count=len(meta.get("field_mapping", {})),
                unmapped_count=len(meta.get("unmapped_fields", [])),
                sections=sections,
            ))
        except Exception as e:
            logger.warning(f"Failed to read gallery template {meta_path}: {e}")

    return items


@router.post("/gallery/clone/{slug}")
async def clone_gallery_template(
    slug: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Clone a pre-made gallery template into the current user's account.
    This creates a personal copy — the user can then customize the field mapping.
    """
    meta_path = GALLERY_DIR / f"{slug}_meta.json"
    docx_path = GALLERY_DIR / f"{slug}.docx"

    if not meta_path.exists():
        raise HTTPException(status_code=404, detail="Gallery template not found")

    with open(meta_path) as f:
        meta = json.load(f)

    # Check if user already has this template (by name)
    existing = db.query(ContractTemplate).filter(
        ContractTemplate.owner_id == current_user.id,
        ContractTemplate.name == meta["name"],
    ).first()
    if existing:
        raise HTTPException(
            status_code=409,
            detail=f"You already have a template named '{meta['name']}'. Delete it first or upload a new version."
        )

    # Read the DOCX file
    if not docx_path.exists():
        raise HTTPException(status_code=404, detail="Template file missing from gallery")

    with open(docx_path, "rb") as f:
        docx_bytes = f.read()

    file_hash = hashlib.sha256(docx_bytes).hexdigest()
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    file_url = f"data:{content_type};base64,{base64.b64encode(docx_bytes).decode()}"

    # Extract text for display
    ocr_text = ""
    try:
        from docx import Document
        import io
        doc = Document(io.BytesIO(docx_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        ocr_text = "\n".join(parts)
    except Exception:
        ocr_text = ""

    template = ContractTemplate(
        owner_id=current_user.id,
        name=meta["name"],
        version=1,
        description=meta.get("description", ""),
        is_active=True,
        file_type=meta.get("file_type", "docx"),
        file_url=file_url,
        file_hash=file_hash,
        ocr_text=ocr_text,
        detected_fields=meta.get("detected_fields", []),
        field_mapping=meta.get("field_mapping", {}),
        unmapped_fields=meta.get("unmapped_fields", []),
    )

    db.add(template)
    db.commit()
    db.refresh(template)

    return _template_to_response(template)


def _template_to_response(template: ContractTemplate) -> TemplateResponse:
    return TemplateResponse(
        id=template.id,
        name=template.name,
        version=template.version,
        description=template.description,
        is_active=template.is_active,
        file_type=template.file_type,
        detected_fields=template.detected_fields or [],
        field_mapping=template.field_mapping or {},
        unmapped_fields=template.unmapped_fields or [],
        created_at=str(template.created_at),
        updated_at=str(template.updated_at),
    )
