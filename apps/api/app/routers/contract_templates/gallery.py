import base64
import hashlib
import json
import logging
from pathlib import Path
from typing import List, Optional
from uuid import UUID

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, status
from sqlalchemy.orm import Session

from app.core.deps import get_db, get_current_user
from app.models.user import User
from app.models.contract_template import ContractTemplate
from app.services.ocr_template_scanner import (
    extract_text,
    detect_fields_with_ai,
    build_field_mapping,
    reconcile_versions,
    compute_file_hash,
    DB_FIELD_REGISTRY,
)

from .common import GALLERY_DIR, _template_to_response
from .schemas import (
    FieldInfo, TemplateResponse, TemplateListItem,
    FieldMappingUpdate, ReconciliationReport, GalleryItem,
)

logger = logging.getLogger(__name__)

router = APIRouter()

@router.get("/gallery/list", response_model=List[GalleryItem])
async def list_gallery_templates():
    """
    List all pre-made starter templates available in the gallery.
    These are shipped with the app — users can clone one into their account.
    No authentication required to browse.
    """
    items = []
    if not GALLERY_DIR.is_dir():
        return items
    for meta_path in sorted(GALLERY_DIR.glob("*_meta.json")):
        try:
            with open(meta_path) as f:
                meta = json.load(f)

            sections = sorted(set(
                f.get("section", "unknown") for f in meta.get("detected_fields", [])
            ))

            items.append(GalleryItem(
                slug=meta_path.stem.replace("_meta", ""),
                name=meta.get("name", meta_path.stem),
                description=meta.get("description", ""),
                file_type=meta.get("file_type", "docx"),
                field_count=len(meta.get("detected_fields", [])),
                mapped_count=len(meta.get("field_mapping", {})),
                unmapped_count=len(meta.get("unmapped_fields", [])),
                sections=sections,
            ))
        except Exception as e:
            logger.warning(f"Failed to read gallery template {meta_path}: {e}")

    return items


@router.post("/gallery/clone/{slug}")
async def clone_gallery_template(
    slug: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Clone a pre-made gallery template into the current user's account.
    This creates a personal copy — the user can then customize the field mapping.
    """
    meta_path = GALLERY_DIR / f"{slug}_meta.json"
    docx_path = GALLERY_DIR / f"{slug}.docx"

    if not meta_path.exists():
        raise HTTPException(status_code=404, detail="Gallery template not found")

    with open(meta_path) as f:
        meta = json.load(f)

    existing = db.query(ContractTemplate).filter(
        ContractTemplate.owner_id == current_user.id,
        ContractTemplate.name == meta["name"],
    ).first()
    if existing:
        raise HTTPException(
            status_code=409,
            detail=f"You already have a template named '{meta['name']}'. Delete it first or upload a new version."
        )

    if not docx_path.exists():
        raise HTTPException(status_code=404, detail="Template file missing from gallery")

    with open(docx_path, "rb") as f:
        docx_bytes = f.read()

    file_hash = hashlib.sha256(docx_bytes).hexdigest()
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    file_url = f"data:{content_type};base64,{base64.b64encode(docx_bytes).decode()}"

    ocr_text = ""
    try:
        from docx import Document
        import io
        doc = Document(io.BytesIO(docx_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        ocr_text = "\n".join(parts)
    except Exception:
        ocr_text = ""

    template = ContractTemplate(
        owner_id=current_user.id,
        name=meta["name"],
        version=1,
        description=meta.get("description", ""),
        is_active=True,
        file_type=meta.get("file_type", "docx"),
        file_url=file_url,
        file_hash=file_hash,
        ocr_text=ocr_text,
        detected_fields=meta.get("detected_fields", []),
        field_mapping=meta.get("field_mapping", {}),
        unmapped_fields=meta.get("unmapped_fields", []),
    )

    db.add(template)
    db.commit()
    db.refresh(template)

    return _template_to_response(template)


# ---------- Endpoints: catch-all /{template_id} routes LAST ----------

