"""
Contract Template System

Generates contracts by filling in agency templates with extracted data.
Supports DOCX templates with placeholder variables.
"""

import os
import re
import logging
from datetime import datetime, date
from typing import Dict, Any, List, Optional
from decimal import Decimal

logger = logging.getLogger(__name__)

# Default template content (used if no DOCX template exists)
DEFAULT_CONTRACT_TEMPLATE = """
HOME CARE SERVICE AGREEMENT

This Home Care Service Agreement ("Agreement") is entered into on {contract_date}

BETWEEN:

Service Provider: {agency_name}
Address: {agency_address}
Phone: {agency_phone}

AND

Client: {client_name}
Address: {client_address}
Phone: {client_phone}
Emergency Contact: {emergency_contact} ({emergency_phone})

1. SERVICES TO BE PROVIDED

The following home care services will be provided. Each service includes its
agreed frequency from the assessment:

{services_list}

1A. SERVICES NOT INCLUDED (DECLINED)

The client or authorized representative declined the following services. These
are not included in this Agreement unless added later in writing:

{declined_services_list}

2. CARE ASSESSMENT SUMMARY

Care Need Level: {care_need_level}
Primary Diagnosis: {primary_diagnosis}
Secondary Conditions: {secondary_conditions}

Client Profile:
{client_profile}

3. SCHEDULE OF SERVICES

Overall Frequency: {frequency}
Days: {schedule_days}
Preferred Time: {preferred_time}
Hours per Week: {weekly_hours}

Per-service schedule:
{per_service_schedule}

4. SERVICE RATES AND PAYMENT

Hourly Rate: ${hourly_rate}/hour
Estimated Weekly Cost: ${weekly_cost}
Estimated Monthly Cost: ${monthly_cost}

Payment Terms: {payment_terms}

5. SPECIAL REQUIREMENTS

{special_requirements}

6. SAFETY CONSIDERATIONS

{safety_concerns}

7. CANCELLATION POLICY

{cancellation_policy}

8. TERMS AND CONDITIONS

{terms_and_conditions}

9. SIGNATURES

By signing below, both parties agree to the terms and conditions of this Agreement.


Client/Authorized Representative:

Signature: _______________________________  Date: ______________

Printed Name: {client_name}


Agency Representative:

Signature: _______________________________  Date: ______________

Printed Name: _______________________________

Title: _______________________________


CARE PLAN GOALS

Short-Term Goals (30 days):
{short_term_goals}

Long-Term Goals (90+ days):
{long_term_goals}
"""


def format_services_list(services: List[Dict]) -> str:
    """Format services list with an explicit frequency line for every service."""
    if not services:
        return "• No home care services identified for this agreement"

    lines = []
    for svc in services:
        if isinstance(svc, str):
            lines.append(f"• {svc}\n  Frequency: As needed")
            continue
        name = (svc.get("name") or "Service").strip()
        desc = (svc.get("description") or "").strip()
        freq = (svc.get("frequency") or "").strip() or "As needed"
        line = f"• {name}"
        if desc:
            line += f"\n  Description: {desc}"
        line += f"\n  Frequency: {freq}"
        evidence = (svc.get("evidence") or "").strip()
        if evidence:
            if len(evidence) > 160:
                evidence = evidence[:157].rstrip() + "..."
            line += f'\n  Assessment quote: "{evidence}"'
        lines.append(line)

    return "\n\n".join(lines)


def format_per_service_schedule(services: List[Dict]) -> str:
    """Compact schedule table: one line per service with frequency."""
    if not services:
        return "• No per-service schedule (no services identified)"
    lines = []
    for svc in services:
        if isinstance(svc, str):
            lines.append(f"• {svc}: As needed")
            continue
        name = (svc.get("name") or "Service").strip()
        freq = (svc.get("frequency") or "").strip() or "As needed"
        lines.append(f"• {name}: {freq}")
    return "\n".join(lines)


def format_declined_services_list(declined: Optional[List[Dict]]) -> str:
    """Client-facing declined / not-included services section."""
    if not declined:
        return (
            "• None recorded. All services listed in Section 1 are included "
            "as stated."
        )
    lines = []
    for item in declined:
        if isinstance(item, str):
            lines.append(
                f"• {item} (declined / not included in this Agreement)"
            )
            continue
        if not isinstance(item, dict):
            continue
        name = (item.get("name") or "Service").strip()
        evidence = (item.get("evidence") or "").strip()
        name_l = name.lower()
        # Clarify maid/deep-clean decline vs included light housekeeping.
        if any(k in name_l for k in ("maid", "deep clean", "deep-clean")):
            line = (
                f"• {name} (declined / not included in this Agreement). "
                "This does not remove Light Housekeeping listed in Section 1, "
                "which covers light tidying related to the client's care areas only."
            )
        else:
            line = f"• {name} (declined / not included in this Agreement)"
        if evidence:
            if len(evidence) > 160:
                evidence = evidence[:157].rstrip() + "..."
            line += f'\n  Client statement: "{evidence}"'
        lines.append(line)
    return "\n\n".join(lines) if lines else (
        "• None recorded. All services listed in Section 1 are included as stated."
    )


def _as_str_list(value: Any) -> List[str]:
    """Normalize LLM list-or-string fields so join never character-splits."""
    if value is None:
        return []
    if isinstance(value, str):
        text = value.strip()
        return [text] if text else []
    if isinstance(value, (list, tuple, set)):
        out: List[str] = []
        for item in value:
            if item is None:
                continue
            if isinstance(item, str):
                text = item.strip()
                if text:
                    out.append(text)
            elif isinstance(item, dict):
                text = str(
                    item.get("name")
                    or item.get("condition")
                    or item.get("concern")
                    or item.get("requirement")
                    or ""
                ).strip()
                if text:
                    out.append(text)
            else:
                text = str(item).strip()
                if text:
                    out.append(text)
        return out
    text = str(value).strip()
    return [text] if text else []


def _safe_float(value: Any, default: float) -> float:
    try:
        if value is None or value == "":
            return float(default)
        return float(value)
    except (TypeError, ValueError):
        return float(default)


def format_list(items: List, prefix: str = "• ") -> str:
    """Format a list of items."""
    if isinstance(items, str):
        items = [items] if items.strip() else []
    if not items:
        return "None specified"
    
    lines = []
    for item in items:
        if isinstance(item, str):
            lines.append(f"{prefix}{item}")
        elif isinstance(item, dict):
            # Handle object items
            if 'concern' in item:
                text = str(item['concern'] or "").strip()
            elif 'name' in item:
                text = str(item['name'] or "").strip()
            elif 'requirement' in item:
                text = str(item['requirement'] or "").strip()
            else:
                text = str(item)
            # Drop internal metadata like "(Severity: High)" from client-facing copy.
            text = re.sub(r"\s*\(Severity:\s*[^)]+\)", "", text, flags=re.IGNORECASE).strip()
            if text:
                lines.append(f"{prefix}{text}")
        else:
            lines.append(f"{prefix}{str(item)}")
    
    return "\n".join(lines) if lines else "None specified"


def sanitize_contract_plain_text(text: str) -> str:
    """Strip ASCII banner lines and severity tags from stored contract bodies."""
    if not text:
        return ""
    cleaned_lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if stripped and set(stripped) <= {"=", " "} and len(stripped) >= 4:
            # Skip pure ==== banner lines (keep signature underscores)
            continue
        if stripped and set(stripped) <= {"-", " "} and len(stripped) >= 8:
            continue
        line = re.sub(r"\s*\(Severity:\s*[^)]+\)", "", line, flags=re.IGNORECASE)
        cleaned_lines.append(line.rstrip())
    # Collapse runs of blank lines left by removed banners
    out: List[str] = []
    blank = 0
    for line in cleaned_lines:
        if not line.strip():
            blank += 1
            if blank <= 1:
                out.append("")
            continue
        blank = 0
        out.append(line)
    return "\n".join(out).strip() + ("\n" if out else "")


def format_goals(goals: List) -> str:
    """Format goals list."""
    if not goals:
        return "• To be determined based on ongoing assessment"
    return format_list(goals)


def generate_contract_from_template(
    contract_data: Dict[str, Any],
    client_info: Dict[str, Any],
    assessment_data: Optional[Dict[str, Any]] = None,
    agency_info: Optional[Dict[str, Any]] = None,
) -> str:
    """
    Generate a contract document from template and extracted data.
    
    Args:
        contract_data: Contract model data (services, schedule, rates)
        client_info: Client information
        assessment_data: LLM-extracted assessment data
        agency_info: Agency information (optional)
    
    Returns:
        Formatted contract text
    """
    
    # Default agency info
    if not agency_info:
        agency_info = {
            "name": "Home Care Services Agency",
            "address": "123 Care Street, Suite 100",
            "phone": "(555) 123-4567",
        }
    
    # Extract schedule info
    schedule = contract_data.get('schedule', {}) or {}
    
    # Extract assessment data
    if assessment_data:
        client_profile_data = assessment_data.get('client_profile', {})
        care_goals = assessment_data.get('care_plan_goals', {})
        extracted_services = assessment_data.get('services_identified', [])
    else:
        client_profile_data = {}
        care_goals = {}
        extracted_services = []
    
    # Build client profile text
    profile_lines = []
    if client_profile_data:
        if client_profile_data.get('mobility_status'):
            profile_lines.append(f"Mobility: {client_profile_data['mobility_status']}")
        if client_profile_data.get('cognitive_status'):
            profile_lines.append(f"Cognitive Status: {client_profile_data['cognitive_status']}")
        if client_profile_data.get('living_situation'):
            profile_lines.append(f"Living Situation: {client_profile_data['living_situation']}")
        risk_factors = _as_str_list(client_profile_data.get('risk_factors'))
        if risk_factors:
            profile_lines.append(f"Risk Factors: {', '.join(risk_factors)}")
    
    client_profile = "\n".join(profile_lines) if profile_lines else "See care assessment documentation"
    
    # Use contract services or extracted services
    services = contract_data.get('services', []) or extracted_services

    # Enrich empty frequencies from overall schedule when possible
    schedule_days_raw = schedule.get('preferred_days', [])
    preferred_times = schedule.get('preferred_times') or ""
    default_freq_bits = []
    if isinstance(schedule_days_raw, list) and schedule_days_raw:
        day_bits = [
            d if isinstance(d, str) else d.get("day", str(d))
            for d in schedule_days_raw
        ]
        if day_bits:
            default_freq_bits.append(", ".join(day_bits))
    elif isinstance(schedule_days_raw, str) and schedule_days_raw.strip():
        default_freq_bits.append(schedule_days_raw.strip())
    if preferred_times and str(preferred_times).lower() not in ("flexible", "n/a", ""):
        default_freq_bits.append(str(preferred_times))
    default_freq = " ".join(default_freq_bits).strip()

    enriched_services = []
    for svc in services:
        if isinstance(svc, str):
            enriched_services.append(
                {"name": svc, "description": "", "frequency": default_freq or "As needed"}
            )
            continue
        if not isinstance(svc, dict):
            continue
        item = dict(svc)
        freq = (item.get("frequency") or "").strip()
        if not freq or freq.lower() in ("as needed", "as scheduled", "n/a"):
            # Keep explicit "As needed" for non-companion add-ons; fill companion-like
            # services from the overall schedule when we have one.
            name_l = (item.get("name") or "").lower()
            if default_freq and any(
                k in name_l for k in ("companion", "supervision", "personal care", "homemaking", "housekeep")
            ):
                item["frequency"] = default_freq
            elif not freq:
                item["frequency"] = "As needed"
        enriched_services.append(item)
    services = enriched_services

    declined = contract_data.get("declined_services")
    if not declined and assessment_data:
        declined = assessment_data.get("declined_services") or []
    
    # Calculate costs (tolerate None / blank LLM fields)
    hourly_rate = _safe_float(contract_data.get('hourly_rate', 25), 25.0)
    weekly_hours = _safe_float(contract_data.get('weekly_hours', 12), 12.0)
    weekly_cost = hourly_rate * weekly_hours
    monthly_cost = weekly_cost * 4.33
    
    # Format schedule days
    schedule_days = schedule.get('preferred_days', [])
    if isinstance(schedule_days, list):
        schedule_days = ', '.join([
            d if isinstance(d, str) else d.get('day', str(d))
            for d in schedule_days
        ]) if schedule_days else "To be determined"
    
    # Get special requirements and safety concerns
    special_reqs = schedule.get('special_requirements', [])
    if assessment_data:
        special_reqs = special_reqs or assessment_data.get('special_requirements', [])
    
    safety = schedule.get('safety_concerns', [])
    if assessment_data:
        safety = safety or assessment_data.get('safety_concerns', [])
    
    # Template variables
    template_vars = {
        # Dates
        'contract_date': date.today().strftime('%B %d, %Y'),
        
        # Agency info
        'agency_name': agency_info.get('name', 'Home Care Services'),
        'agency_address': agency_info.get('address', ''),
        'agency_phone': agency_info.get('phone', ''),
        
        # Client info
        'client_name': client_info.get('full_name', 'Client'),
        'client_address': client_info.get('address', 'Not provided'),
        'client_phone': client_info.get('phone', 'Not provided'),
        'emergency_contact': client_info.get('emergency_contact_name', 'Not provided'),
        'emergency_phone': client_info.get('emergency_contact_phone', 'Not provided'),
        
        # Services
        'services_list': format_services_list(services),
        'declined_services_list': format_declined_services_list(declined),
        'per_service_schedule': format_per_service_schedule(services),
        
        # Assessment summary
        'care_need_level': schedule.get('care_need_level', 'MODERATE'),
        'primary_diagnosis': client_profile_data.get('primary_diagnosis', 'See medical records'),
        'secondary_conditions': (
            ', '.join(_as_str_list(client_profile_data.get('secondary_conditions')))
            or 'None noted'
        ),
        'client_profile': client_profile,
        
        # Schedule
        'frequency': schedule.get('frequency') or default_freq or 'As scheduled',
        'schedule_days': schedule_days,
        'preferred_time': schedule.get('preferred_times', 'Flexible'),
        'weekly_hours': f"{weekly_hours:.1f}",
        
        # Rates
        'hourly_rate': f"{hourly_rate:.2f}",
        'weekly_cost': f"{weekly_cost:.2f}",
        'monthly_cost': f"{monthly_cost:.2f}",
        'payment_terms': 'Payment due within 7 days of invoice',
        
        # Special requirements and safety
        'special_requirements': format_list(special_reqs),
        'safety_concerns': format_list(safety),
        
        # Policies
        'cancellation_policy': contract_data.get('cancellation_policy', '') or DEFAULT_CANCELLATION_POLICY,
        'terms_and_conditions': contract_data.get('terms_and_conditions', '') or DEFAULT_TERMS,
        
        # Goals
        'short_term_goals': format_goals(care_goals.get('short_term', [])),
        'long_term_goals': format_goals(care_goals.get('long_term', [])),
    }
    
    # Fill in template
    contract_text = sanitize_contract_plain_text(
        DEFAULT_CONTRACT_TEMPLATE.format(**template_vars)
    )
    
    return contract_text


DEFAULT_CANCELLATION_POLICY = """• Cancellations made more than 24 hours in advance: No charge
• Cancellations made within 24 hours: 50% of scheduled visit fee
• No-show without notice: 100% of scheduled visit fee
• Emergency cancellations will be evaluated on a case-by-case basis
• The agency reserves the right to cancel services with 14 days written notice"""


DEFAULT_TERMS = """CONFIDENTIALITY: All client information is protected under HIPAA regulations. 
Information will only be shared with healthcare providers directly involved in 
the client's care or as required by law.

LIABILITY: The Agency maintains comprehensive general liability and professional 
liability insurance. Caregivers are employees of the Agency and covered under 
workers' compensation insurance.

TERMINATION: Either party may terminate this Agreement with 14 days written 
notice. Immediate termination may occur if there is a safety concern or 
non-payment of services.

CAREGIVER PROVISIONS: All caregivers undergo background checks, are trained 
in home care practices, and are supervised by agency management. Caregivers 
are not permitted to accept gifts or be named in client wills."""


def generate_docx_contract(
    contract_data: Dict[str, Any],
    client_info: Dict[str, Any],
    assessment_data: Optional[Dict[str, Any]] = None,
    template_path: Optional[str] = None,
) -> bytes:
    """
    Generate a DOCX contract document.
    
    If a template_path is provided, uses that template.
    Otherwise generates a basic DOCX from the text template.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
    except ImportError:
        logger.warning("python-docx not installed, returning text format")
        return generate_contract_from_template(contract_data, client_info, assessment_data).encode()
    
    # Generate the text content
    contract_text = generate_contract_from_template(contract_data, client_info, assessment_data)
    
    # Create DOCX document
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Parse and add content
    lines = contract_text.split('\n')
    
    for line in lines:
        line = line.rstrip()
        
        if line.startswith('====='):
            # Section separator - skip
            continue
        elif line.startswith('HOME CARE SERVICE AGREEMENT'):
            # Title
            p = doc.add_heading(line, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line and line[0].isdigit() and '. ' in line[:4]:
            # Section heading
            doc.add_heading(line, level=1)
        elif line.startswith('•'):
            # Bullet point
            doc.add_paragraph(line[1:].strip(), style='List Bullet')
        elif line.strip():
            # Regular paragraph
            doc.add_paragraph(line)
    
    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()
