"""Regression tests for DOCX template fill (no double-inserted values)."""

import io

from docx import Document

from app.services.document_generation.docx_templates import fill_docx_template


def _docx_with_paragraphs(lines: list[str]) -> bytes:
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _paragraph_texts(docx_bytes: bytes) -> list[str]:
    doc = Document(io.BytesIO(docx_bytes))
    return [p.text for p in doc.paragraphs if p.text.strip()]


def test_placeholder_then_label_does_not_double_fill():
    """The production bug: `{client_name}` fills, then `Client Name:` fills again."""
    raw = _docx_with_paragraphs(
        [
            "HOME CARE SERVICE AGREEMENT",
            "Service Provider: {agency_name}",
            "Client Name: {client_name}",
            "Address: {client_address}",
            "Indemnity: Agency original indemnity clause remains here.",
        ]
    )
    filled = fill_docx_template(
        raw,
        {
            "agency_name": "Verify Agency 1787677180",
            "client_name": "Ada Client",
            "client_address": "123 Pine St",
        },
    )
    texts = _paragraph_texts(filled)
    joined = "\n".join(texts)

    assert "Ada ClientAda Client" not in joined
    assert "123 Pine St123 Pine St" not in joined
    assert "Client Name: Ada Client" in joined
    assert "Address: 123 Pine St" in joined
    assert "Service Provider: Verify Agency 1787677180" in joined
    assert "Indemnity: Agency original indemnity clause remains here." in joined


def test_blank_label_still_fills():
    raw = _docx_with_paragraphs(
        [
            "Client Name: ___________",
            "Address: ",
            "City: ______  State: ______",
        ]
    )
    filled = fill_docx_template(
        raw,
        {
            "client_name": "Ada Client",
            "client_address": "123 Pine St",
            "client_city": "Miami",
            "client_state": "FL",
        },
    )
    texts = _paragraph_texts(filled)
    joined = "\n".join(texts)

    assert "Client Name: Ada Client" in joined
    assert "Address: 123 Pine St" in joined
    assert "City: Miami" in joined
    assert "State: FL" in joined
    assert "Ada ClientAda Client" not in joined


def test_already_filled_label_is_left_alone():
    raw = _docx_with_paragraphs(["Client Name: Ada Client", "Address: 123 Pine St"])
    filled = fill_docx_template(
        raw,
        {"client_name": "Ada Client", "client_address": "123 Pine St"},
    )
    texts = _paragraph_texts(filled)
    assert texts == ["Client Name: Ada Client", "Address: 123 Pine St"]
